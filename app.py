#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Analizador de Sentencias IPP/INSS - Aplicación FastAPI
Aplicación robusta para análisis de documentos legales con modelo de IA pre-entrenado
"""

import os
import json
import re
import uuid
import shutil
import logging
from pathlib import Path
from threading import Lock
from typing import Optional, Dict, List, Any
from datetime import datetime
from io import BytesIO

from fastapi import FastAPI, Request, File, UploadFile, Form, HTTPException
from fastapi.templating import Jinja2Templates
from fastapi.staticfiles import StaticFiles
from fastapi.responses import HTMLResponse, JSONResponse, RedirectResponse, PlainTextResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE

# Configurar logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Cache para análisis
ANALISIS_CACHE = {}
CACHE_TIMESTAMP = None
CACHE_DURATION = 300  # 5 minutos en segundos

# Crear aplicación FastAPI con configuración robusta
app = FastAPI(
    title="Analizador de Sentencias IPP/INSS",
    description="API robusta para análisis inteligente de documentos legales",
    version="2.0.0",
    docs_url="/docs",
    redoc_url="/redoc"
)

# Configurar CORS para desarrollo
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Configurar templates y archivos estáticos
templates = Jinja2Templates(directory="templates")

# Configurar archivos estáticos con tipos MIME correctos
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse
import mimetypes

# Configurar tipos MIME para archivos PDF
mimetypes.add_type('application/pdf', '.pdf')

# Montar archivos estáticos con configuración personalizada
app.mount("/static", StaticFiles(directory="static"), name="static")

# Montar directorio de sentencias con configuración específica para PDFs
class PDFStaticFiles(StaticFiles):
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
    
    def get_response(self, path: str, scope):
        response = super().get_response(path, scope)
        if path.endswith('.pdf'):
            response.headers["Content-Type"] = "application/pdf"
            response.headers["X-Content-Type-Options"] = "nosniff"
            response.headers["Cache-Control"] = "no-cache, no-store, must-revalidate"
            response.headers["Pragma"] = "no-cache"
            response.headers["Expires"] = "0"
        return response

app.mount("/sentencias", PDFStaticFiles(directory="sentencias"), name="sentencias")

# Configuración de directorios
BASE_DIR = Path(__file__).parent
SENTENCIAS_DIR = BASE_DIR / "sentencias"
UPLOADS_DIR = BASE_DIR / "uploads"
MODELS_DIR = BASE_DIR / "models"
LOGS_DIR = BASE_DIR / "logs"
FRASES_FILE = BASE_DIR / "models" / "frases_clave.json"
frases_lock = Lock()

# Crear directorios necesarios
for directory in [SENTENCIAS_DIR, UPLOADS_DIR, MODELS_DIR, LOGS_DIR]:
    directory.mkdir(exist_ok=True)

# Configuración de archivos permitidos
ALLOWED_EXTENSIONS = {'.pdf', '.txt', '.doc', '.docx'}
MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB

# Frases por defecto en caso de que no exista el archivo o sea inválido
DEFAULT_FRASES_CLAVE: Dict[str, List[str]] = {
    "procedimiento_legal": ["procedente", "desestimamos", "estimamos"],
}


def load_frases_clave() -> Dict[str, List[str]]:
    """Carga el JSON de frases clave de disco con validación básica."""
    try:
        if not FRASES_FILE.exists():
            logger.warning(f"Archivo de frases no encontrado en {FRASES_FILE}")
            return DEFAULT_FRASES_CLAVE
        with open(FRASES_FILE, "r", encoding="utf-8") as f:
            data = json.load(f)
        if not isinstance(data, dict):
            logger.warning("Formato inválido en frases_clave.json (no es dict)")
            return DEFAULT_FRASES_CLAVE
        # Normalizar valores a listas de strings únicas
        normalizado: Dict[str, List[str]] = {}
        for categoria, frases in data.items():
            if not isinstance(categoria, str):
                continue
            if isinstance(frases, list):
                solo_texto = [str(x).strip() for x in frases if str(x).strip()]
                # quitar duplicados preservando orden
                seen = set()
                unicos: List[str] = []
                for s in solo_texto:
                    if s.lower() not in seen:
                        seen.add(s.lower())
                        unicos.append(s)
                # Incluir también categorías vacías (permitimos crear primero y rellenar después)
                normalizado[categoria] = unicos
        return normalizado or DEFAULT_FRASES_CLAVE
    except Exception as e:
        logger.error(f"Error cargando frases_clave.json: {e}")
        return DEFAULT_FRASES_CLAVE


def save_frases_clave(data: Dict[str, List[str]]) -> None:
    """Guarda el JSON de frases clave de forma atómica y segura."""
    if not isinstance(data, dict):
        raise ValueError("El payload debe ser un objeto {categoria: [frases]}")
    for categoria, frases in data.items():
        if not isinstance(categoria, str) or not isinstance(frases, list):
            raise ValueError("Formato inválido: claves str, valores lista de str")
    MODELS_DIR.mkdir(exist_ok=True)
    temp_path = FRASES_FILE.with_suffix(".tmp")
    with frases_lock:
        with open(temp_path, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        os.replace(temp_path, FRASES_FILE)

# Importar el analizador de IA (asumiendo que ya está entrenado)
try:
    from backend.analisis import AnalizadorLegal
    ANALIZADOR_IA_DISPONIBLE = True
    logger.info("✅ Módulo de IA cargado correctamente")
except ImportError as e:
    ANALIZADOR_IA_DISPONIBLE = False
    logger.warning(f"⚠️ Módulo de IA no disponible: {e}")
    logger.info("Se usará análisis básico como fallback")


class AnalizadorBasico:
    """Analizador básico como fallback cuando no hay IA disponible"""
    
    def __init__(self):
        self.frases_clave = {}
        self.cargar_frases_desde_modelo()

    def cargar_frases_desde_modelo(self) -> None:
        """Carga frases clave desde models/frases_clave.json o usa un conjunto por defecto."""
        try:
            datos = load_frases_clave()
            if isinstance(datos, dict) and datos:
                self.frases_clave = datos
            else:
                self.frases_clave = DEFAULT_FRASES_CLAVE
                logger.warning("Usando frases por defecto; archivo vacío o inválido")
        except Exception as e:
            logger.warning(f"No se pudo cargar frases_clave.json: {e}. Usando valores por defecto")
            self.frases_clave = DEFAULT_FRASES_CLAVE
    
    def analizar_documento(self, ruta_archivo: str, nombre_original: str = "") -> Dict[str, Any]:
        """Analiza un documento usando métodos básicos"""
        try:
            # Leer contenido del archivo
            contenido = self._leer_archivo(ruta_archivo)
            if not contenido:
                return self._crear_resultado_error("No se pudo leer el contenido del archivo")
            
            # Análisis básico
            frases_encontradas = self._contar_frases_clave(contenido, nombre_original)
            prediccion = self._prediccion_basica(contenido)
            argumentos = self._extraer_argumentos(contenido)
            insights = self._generar_insights(prediccion, frases_encontradas)
            
            # Calcular tiempo de procesamiento
            tiempo_inicio = getattr(self, '_tiempo_inicio', None)
            tiempo_procesamiento = None
            if tiempo_inicio:
                tiempo_procesamiento = f"{(datetime.now() - tiempo_inicio).total_seconds():.2f}s"
            
            return {
                "archivo": ruta_archivo,
                "nombre_archivo": nombre_original or Path(ruta_archivo).name,
                "procesado": True,
                "texto_extraido": contenido[:1000] + "..." if len(contenido) > 1000 else contenido,
                "longitud_texto": len(contenido),
                "prediccion": prediccion,
                "argumentos": argumentos,
                "frases_clave": frases_encontradas,
                "resumen_inteligente": self._generar_resumen(prediccion, frases_encontradas),
                "insights_juridicos": insights,
                "archivo_id": Path(ruta_archivo).name,
                "tipo_documento": "sentencia",
                "extract_entities": False,
                "analyze_arguments": False,
                "modelo_ia": False,
                "total_frases_clave": sum(datos["total"] for datos in frases_encontradas.values()),
                "timestamp": datetime.now().isoformat(),
                "tiempo_procesamiento": tiempo_procesamiento or "N/A",
                "ruta_archivo": ruta_archivo
            }
            
        except Exception as e:
            logger.error(f"Error en análisis básico: {e}")
            return self._crear_resultado_error(f"Error en análisis: {str(e)}")
    
    def _leer_archivo(self, ruta: str) -> Optional[str]:
        """Lee el contenido de un archivo con manejo de errores"""
        try:
            if ruta.endswith('.txt'):
                # Intentar diferentes encodings
                for encoding in ['utf-8', 'latin-1', 'cp1252']:
                    try:
                        with open(ruta, 'r', encoding=encoding) as f:
                            return f.read().strip()
                    except UnicodeDecodeError:
                        continue
                return None
            else:
                # Para otros formatos, usar función de lectura genérica
                return self._leer_archivo_generico(ruta)
        except Exception as e:
            logger.error(f"Error leyendo archivo {ruta}: {e}")
            return None
    
    def _leer_archivo_generico(self, ruta: str) -> Optional[str]:
        """Lee archivos de diferentes formatos"""
        try:
            # Por ahora solo manejamos texto, pero aquí se podría extender
            # para PDF, DOC, etc. usando librerías como PyPDF2, python-docx
            return "Contenido del archivo no disponible en formato de texto"
        except Exception as e:
            logger.error(f"Error leyendo archivo genérico {ruta}: {e}")
            return None
    
    def _contar_frases_clave(self, texto: str, nombre_archivo: str) -> Dict[str, Any]:
        """Cuenta las ocurrencias de frases clave"""
        if not texto:
            return {}
        
        resultados = {}
        for categoria, variantes in self.frases_clave.items():
            total = 0
            ocurrencias = []
            
            for variante in variantes:
                # Permitir espacios/guiones/underscores intercambiables entre palabras
                flexible = re.escape(variante)
                flexible = flexible.replace("\\ ", "\\s+")
                flexible = flexible.replace("\\_", "[\\s_\\-_]+")
                flexible = flexible.replace("\\-", "[\\s_\\-_]+")
                patron = re.compile(flexible, re.IGNORECASE)
                matches = patron.finditer(texto)
                
                for match in matches:
                    total += 1
                    start_pos = match.start()
                    end_pos = match.end()
                    
                    # Obtener contexto (100 caracteres antes y después)
                    context_start = max(0, start_pos - 100)
                    context_end = min(len(texto), end_pos + 100)
                    contexto = texto[context_start:context_end]
                    
                    # Marcar la frase encontrada
                    frase_encontrada = texto[start_pos:end_pos]
                    contexto_marcado = contexto.replace(frase_encontrada, f"**{frase_encontrada}**")
                    
                    ocurrencias.append({
                        "frase": variante,
                        "posicion": start_pos,
                        "contexto": contexto_marcado,
                        "linea": texto[:start_pos].count('\n') + 1,
                        "archivo": nombre_archivo
                    })
            
            if total > 0:
                # Obtener frases únicas encontradas
                frases_encontradas = list(set([oc["frase"] for oc in ocurrencias]))
                
                resultados[categoria] = {
                    "total": total,
                    "ocurrencias": ocurrencias,
                    "frases": frases_encontradas
                }
        
        return resultados
    
    def _prediccion_basica(self, texto: str) -> Dict[str, Any]:
        """Predicción básica basada en palabras clave"""
        texto_lower = texto.lower()
        
        palabras_positivas = [
            "estimamos", "estimamos procedente", "procedente", "favorable",
            "accedemos", "concedemos", "reconocemos", "declaramos procedente",
            "estimamos fundada", "fundada", "estimamos parcialmente"
        ]
        
        palabras_negativas = [
            "desestimamos", "desestimamos la reclamación", "desfavorable",
            "no procedente", "rechazamos", "denegamos", "estimamos infundada",
            "infundada", "desestimamos parcialmente"
        ]
        
        positivas = sum(1 for palabra in palabras_positivas if palabra in texto_lower)
        negativas = sum(1 for palabra in palabras_negativas if palabra in texto_lower)
        
        total = positivas + negativas
        if total == 0:
            return {
                "es_favorable": True,
                "confianza": 0.5,
                "interpretacion": "Neutral - No hay indicadores claros"
            }
        
        confianza = min(0.95, max(0.1, (positivas + negativas) / (total + 1)))
        es_favorable = positivas > negativas
        
        if es_favorable:
            interpretacion = f"Favorable - {positivas} indicadores positivos encontrados"
        else:
            interpretacion = f"Desfavorable - {negativas} indicadores negativos encontrados"
        
        return {
            "es_favorable": es_favorable,
            "confianza": confianza,
            "interpretacion": interpretacion
        }
    
    def _extraer_argumentos(self, texto: str) -> List[Dict[str, Any]]:
        """Extrae argumentos legales del texto"""
        argumentos = []
        
        patrones = [
            r"por\s+(?:lo\s+)?que\s+([^.]*?\.)",
            r"fundamentos?\s+(?:de\s+)?(?:derecho|derecho\s+por\s+lo\s+que)\s+([^.]*?\.)",
            r"considerando\s+que\s+([^.]*?\.)",
            r"vistos\s+([^.]*?\.)",
            r"resultando\s+([^.]*?\.)"
        ]
        
        for patron in patrones:
            matches = re.finditer(patron, texto, re.IGNORECASE)
            for match in matches:
                argumento = match.group(1).strip()
                if len(argumento) > 20:
                    argumentos.append({
                        "tipo": "argumento_legal",
                        "texto": argumento,
                        "posicion": match.start(),
                        "confianza": 0.7,
                        "categoria": "fundamento_juridico"
                    })
        
        return argumentos
    
    def _generar_insights(self, prediccion: Dict, frases_clave: Dict) -> List[str]:
        """Genera insights jurídicos básicos"""
        insights = []
        
        if prediccion["es_favorable"]:
            insights.append("El documento presenta argumentos sólidos a favor del caso.")
            insights.append("La resolución favorece al reclamante.")
        else:
            insights.append("El documento presenta argumentos que pueden ser desfavorables.")
            insights.append("La resolución no favorece al reclamante.")
        
        if frases_clave:
            categorias = list(frases_clave.keys())
            insights.append(f"Se identificaron {len(categorias)} categorías de frases clave.")
            insights.append(f"Las categorías más relevantes son: {', '.join(categorias[:3])}.")
        
        insights.append(f"Confianza del análisis: {prediccion['confianza']:.1%}")
        
        return insights
    
    def _generar_resumen(self, prediccion: Dict, frases_clave: Dict) -> str:
        """Genera un resumen inteligente del análisis"""
        if prediccion["es_favorable"]:
            base = "Análisis favorable del documento legal. "
        else:
            base = "Análisis desfavorable del documento legal. "
        
        if frases_clave:
            total_frases = sum(datos["total"] for datos in frases_clave.values())
            base += f"Se identificaron {total_frases} frases clave relevantes. "
        
        base += f"Confianza del análisis: {prediccion['confianza']:.1%}."
        
        return base
    
    def _crear_resultado_error(self, mensaje: str) -> Dict[str, Any]:
        """Crea un resultado de error estandarizado"""
        return {
            "error": mensaje,
            "procesado": False,
            "nombre_archivo": "archivo_desconocido",
            "prediccion": {"es_favorable": False, "confianza": 0.0},
            "frases_clave": {},
            "argumentos": [],
            "insights_juridicos": [f"Error: {mensaje}"],
            "longitud_texto": 0,
            "total_frases_clave": 0
        }


# Instanciar analizador básico
analizador_basico = AnalizadorBasico()


# ====== MODELOS Pydantic para CRUD ======
class FrasesPayload(BaseModel):
    categorias: Dict[str, List[str]]

class CategoriaPayload(BaseModel):
    nombre: str
    frases: List[str] = []

class FrasePayload(BaseModel):
    categoria: str
    frase: str

class RenameCategoryPayload(BaseModel):
    old_name: str
    new_name: str

class UpdatePhrasePayload(BaseModel):
    categoria: str
    old_frase: str
    new_frase: str


class DeleteDocumentPayload(BaseModel):
    nombre_archivo: str


# ====== ENDPOINTS CRUD DE FRASES CLAVE ======
@app.get("/api/frases")
async def listar_frases():
    """Lista todas las categorías y frases clave actuales."""
    datos = load_frases_clave()
    return {"categorias": datos}


@app.post("/api/frases")
async def reemplazar_frases(payload: FrasesPayload):
    """Reemplaza completamente el set de frases clave."""
    save_frases_clave(payload.categorias)
    analizador_basico.cargar_frases_desde_modelo()
    return {"status": "ok"}


@app.post("/api/frases/categoria")
async def crear_categoria(payload: CategoriaPayload):
    datos = load_frases_clave()
    nombre = payload.nombre.strip()
    if not nombre:
        raise HTTPException(status_code=400, detail="Nombre de categoría requerido")
    if nombre in datos:
        raise HTTPException(status_code=409, detail="La categoría ya existe")
    datos[nombre] = [s.strip() for s in payload.frases if s and s.strip()]
    save_frases_clave(datos)
    analizador_basico.cargar_frases_desde_modelo()
    return {"status": "ok", "categoria": nombre}


@app.delete("/api/frases/categoria/{nombre}")
async def eliminar_categoria(nombre: str):
    datos = load_frases_clave()
    if nombre not in datos:
        raise HTTPException(status_code=404, detail="Categoría no encontrada")
    datos.pop(nombre)
    save_frases_clave(datos)
    analizador_basico.cargar_frases_desde_modelo()
    return {"status": "ok"}


@app.patch("/api/frases/categoria")
async def renombrar_categoria(payload: RenameCategoryPayload):
    datos = load_frases_clave()
    old_name = payload.old_name.strip()
    new_name = payload.new_name.strip()
    if not old_name or not new_name:
        raise HTTPException(status_code=400, detail="Nombres requeridos")
    if old_name not in datos:
        raise HTTPException(status_code=404, detail="Categoría original no encontrada")
    if new_name in datos and new_name != old_name:
        raise HTTPException(status_code=409, detail="La categoría destino ya existe")
    if new_name == old_name:
        return {"status": "ok", "categoria": new_name}
    datos[new_name] = datos.pop(old_name)
    save_frases_clave(datos)
    analizador_basico.cargar_frases_desde_modelo()
    return {"status": "ok", "categoria": new_name}


@app.post("/api/frases/frase")
async def agregar_frase(payload: FrasePayload):
    datos = load_frases_clave()
    categoria = payload.categoria
    frase = payload.frase.strip()
    if not frase:
        raise HTTPException(status_code=400, detail="Frase requerida")
    if categoria not in datos:
        datos[categoria] = []
    # evitar duplicados case-insensitive
    if frase.lower() not in [f.lower() for f in datos[categoria]]:
        datos[categoria].append(frase)
    save_frases_clave(datos)
    analizador_basico.cargar_frases_desde_modelo()
    return {"status": "ok"}


@app.delete("/api/frases/frase")
async def eliminar_frase(payload: FrasePayload):
    datos = load_frases_clave()
    categoria = payload.categoria
    frase = payload.frase.strip()
    if categoria not in datos:
        raise HTTPException(status_code=404, detail="Categoría no encontrada")
    datos[categoria] = [f for f in datos[categoria] if f.lower() != frase.lower()]
    save_frases_clave(datos)
    analizador_basico.cargar_frases_desde_modelo()
    return {"status": "ok"}


@app.patch("/api/frases/frase")
async def actualizar_frase(payload: UpdatePhrasePayload):
    datos = load_frases_clave()
    categoria = payload.categoria
    old_frase = payload.old_frase.strip()
    new_frase = payload.new_frase.strip()
    if not new_frase:
        raise HTTPException(status_code=400, detail="Nueva frase requerida")
    if categoria not in datos:
        raise HTTPException(status_code=404, detail="Categoría no encontrada")
    frases = datos[categoria]
    indices = [i for i, f in enumerate(frases) if f.lower() == old_frase.lower()]
    if not indices:
        raise HTTPException(status_code=404, detail="Frase original no encontrada")
    # Evitar duplicado de destino
    if any(f.lower() == new_frase.lower() for f in frases):
        # si ya existe exacta, eliminar la vieja
        frases = [f for f in frases if f.lower() != old_frase.lower()]
    else:
        # reemplazar el primero y eliminar duplicados del resto
        idx = indices[0]
        frases[idx] = new_frase
        # limpiar duplicados case-insensitive preservando orden
        seen = set()
        dedup: List[str] = []
        for f in frases:
            key = f.lower()
            if key not in seen:
                seen.add(key)
                dedup.append(f)
        frases = dedup
    datos[categoria] = frases
    save_frases_clave(datos)
    analizador_basico.cargar_frases_desde_modelo()
    return {"status": "ok"}


def validar_archivo(archivo: UploadFile) -> Dict[str, Any]:
    """Valida un archivo subido"""
    errores = []
    
    # Validar nombre
    if not archivo.filename:
        errores.append("No se seleccionó ningún archivo")
    
    # Validar extensión
    if archivo.filename:
        extension = Path(archivo.filename).suffix.lower()
        if extension not in ALLOWED_EXTENSIONS:
            errores.append(f"Extensión no permitida: {extension}. Permitidas: {', '.join(ALLOWED_EXTENSIONS)}")
    
    # Validar tamaño
    if archivo.size and archivo.size > MAX_FILE_SIZE:
        errores.append(f"Archivo demasiado grande: {archivo.size / (1024*1024):.1f}MB. Máximo: {MAX_FILE_SIZE / (1024*1024)}MB")
    
    return {
        "valido": len(errores) == 0,
        "errores": errores
    }


@app.get("/", response_class=HTMLResponse)
async def pagina_principal(request: Request):
    """Página principal de la aplicación"""
    try:
        # Analizar sentencias existentes
        resultado = analizar_sentencias_existentes()
        
        # DEBUG: Log del resultado
        logger.info(f"Resultado de analizar_sentencias_existentes: {resultado}")
        logger.info(f"ranking_global keys: {list(resultado.get('ranking_global', {}).keys())}")
        logger.info(f"ranking_global length: {len(resultado.get('ranking_global', {}))}")
        
        return templates.TemplateResponse("index.html", {
            "request": request,
            "archivos_analizados": resultado.get("archivos_analizados", 0),
            "total_apariciones": resultado.get("total_apariciones", 0),
            "resultados_por_archivo": resultado.get("resultados_por_archivo", {}),
            "ranking_global": resultado.get("ranking_global", {}),
            "ia_disponible": ANALIZADOR_IA_DISPONIBLE
        })
    except Exception as e:
        logger.error(f"Error en página principal: {e}")
        return templates.TemplateResponse("index.html", {
            "request": request,
            "error": f"Error al cargar la página: {str(e)}",
            "ia_disponible": ANALIZADOR_IA_DISPONIBLE
        })


@app.get("/subir", response_class=HTMLResponse)
async def formulario_subida(request: Request):
    """Formulario para subir documentos"""
    return templates.TemplateResponse("subir.html", {
        "request": request,
        "ia_disponible": ANALIZADOR_IA_DISPONIBLE
    })


@app.post("/upload")
async def subir_documento(
    file: UploadFile = File(...),
    document_type: str = Form("sentencia"),
    extract_entities: bool = Form(True),
    analyze_arguments: bool = Form(True)
):
    """Procesa la subida de un documento"""
    try:
        # Validar archivo
        validacion = validar_archivo(file)
        if not validacion["valido"]:
            raise HTTPException(status_code=400, detail="; ".join(validacion["errores"]))
        
        # Generar nombre único
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        unique_id = str(uuid.uuid4())[:8]
        extension = Path(file.filename).suffix
        nuevo_nombre = f"{document_type}_{timestamp}_{unique_id}{extension}"
        
        # Guardar archivo
        ruta_archivo = UPLOADS_DIR / nuevo_nombre
        with open(ruta_archivo, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        
        logger.info(f"Archivo subido: {nuevo_nombre}")
        
        # Analizar documento
        if ANALIZADOR_IA_DISPONIBLE:
            try:
                analizador = AnalizadorLegal()
                resultado = analizador.analizar_documento(str(ruta_archivo))
                resultado["modelo_ia"] = True
                logger.info("Análisis con IA completado")
            except Exception as e:
                logger.warning(f"Fallback a análisis básico: {e}")
                resultado = analizador_basico.analizar_documento(str(ruta_archivo), file.filename)
                resultado["modelo_ia"] = False
        else:
            resultado = analizador_basico.analizar_documento(str(ruta_archivo), file.filename)
            resultado["modelo_ia"] = False
        
        # Agregar metadatos
        resultado.update({
            "nombre_archivo": file.filename,
            "archivo_id": nuevo_nombre,
            "tipo_documento": document_type,
            "extract_entities": extract_entities,
            "analyze_arguments": analyze_arguments,
            "timestamp": timestamp,
            "ruta_archivo": str(ruta_archivo)
        })
        
        # Mover a carpeta de sentencias si es apropiado
        if document_type == "sentencia":
            destino = SENTENCIAS_DIR / nuevo_nombre
            shutil.move(str(ruta_archivo), str(destino))
            resultado["ruta_archivo"] = str(destino)
        
        return JSONResponse(content=resultado)
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error procesando archivo: {e}")
        raise HTTPException(status_code=500, detail=f"Error interno: {str(e)}")


@app.get("/resultado/{archivo_id}", response_class=HTMLResponse)
async def mostrar_resultados(request: Request, archivo_id: str):
    """Muestra los resultados del análisis"""
    try:
        # Buscar archivo
        ruta_archivo = UPLOADS_DIR / archivo_id
        if not ruta_archivo.exists():
            ruta_archivo = SENTENCIAS_DIR / archivo_id
        
        if not ruta_archivo.exists():
            raise HTTPException(status_code=404, detail="Archivo no encontrado")
        
        # Analizar si es necesario
        if archivo_id.startswith(("sentencia_", "demanda_", "informe_")):
            if ANALIZADOR_IA_DISPONIBLE:
                try:
                    analizador = AnalizadorLegal()
                    resultado = analizador.analizar_documento(str(ruta_archivo))
                    resultado["modelo_ia"] = True
                except Exception as e:
                    logger.warning(f"Fallback a análisis básico: {e}")
                    resultado = analizador_basico.analizar_documento(str(ruta_archivo), archivo_id)
                    resultado["modelo_ia"] = False
            else:
                resultado = analizador_basico.analizar_documento(str(ruta_archivo), archivo_id)
                resultado["modelo_ia"] = False
        else:
            # Usar análisis existente
            resultado = {
                "nombre_archivo": archivo_id,
                "procesado": True,
                "prediccion": {"es_favorable": True, "confianza": 0.8},
                "resumen_inteligente": "Análisis del documento existente.",
                "argumentos": [],
                "frases_clave": {},
                "insights_juridicos": ["Documento analizado previamente."],
                "longitud_texto": 0,
                "total_frases_clave": 0,
                "modelo_ia": False
            }
        
        return templates.TemplateResponse("resultado.html", {
            "request": request,
            "resultado": resultado
        })
        
    except Exception as e:
        logger.error(f"Error mostrando resultados: {e}")
        raise HTTPException(status_code=500, detail=f"Error: {str(e)}")


@app.get("/archivo/{archivo_id}", response_class=HTMLResponse)
async def ver_archivo(request: Request, archivo_id: str, highlight: str = None, pos: int = None, index: int = None):
    """Muestra el contenido completo de un archivo con frases clave resaltadas y opcionalmente resalta una aparición específica"""
    try:
        logger.info(f"🔍 Procesando solicitud para archivo: {archivo_id}")
        logger.info(f"📋 Parámetros: highlight={highlight}, pos={pos}, index={index}")
        
        # Buscar el archivo en la carpeta de sentencias
        archivo_path = SENTENCIAS_DIR / archivo_id
        
        if not archivo_path.exists():
            logger.error(f"❌ Archivo no encontrado: {archivo_path}")
            raise HTTPException(status_code=404, detail=f"Archivo '{archivo_id}' no encontrado")
        
        logger.info(f"✅ Archivo encontrado: {archivo_path}")
        
        # Analizar el archivo para obtener frases clave
        if ANALIZADOR_IA_DISPONIBLE:
            try:
                analizador = AnalizadorLegal()
                resultado = analizador.analizar_documento(str(archivo_path))
            except Exception as e:
                logger.warning(f"Fallback a análisis básico: {e}")
                resultado = analizador_basico.analizar_documento(str(archivo_path), archivo_id)
        else:
            resultado = analizador_basico.analizar_documento(str(archivo_path), archivo_id)
        
        if not resultado.get("procesado"):
            raise HTTPException(status_code=500, detail="No se pudo procesar el archivo")
        
        # Función para limpiar contenido HTML mal formado
        def limpiar_contenido_html(texto: str) -> str:
            """Limpia contenido HTML mal formado y caracteres especiales"""
            if not texto:
                return ""
            
            import re
            
            # PRIMERA PASADA: Eliminar completamente todos los fragmentos HTML malformados
            # Patrones específicos que aparecen en el texto
            patrones_malformados = [
                r'categoria="lesiones_permanentes" title="Click para ver detalles de lesiones_permanentes">lesionesass="frase-resaltada frase-lesiones" data-',
                r'lesionesass="frase-resaltada frase-lesiones" data-categoria="lesiones_permanentes" title="Click para ver detalles de lesiones_permanentes">lesiones',
                r'frase-resaltada frase-lesiones"',
                r'Click para ver detalles de lesiones_permanentes',
                r'indemnizaciónfrase-resaltada frase-prestaciones" data-categoria="prestaciones" title="Click para ver detalles de prestaciones">indemnización',
                r'accidente de trabajoesaltada frase-accidente" data-categoria="accidente_laboral" title="Click para ver detalles de accidente_laboral">accidente de trabajo',
                r'INSSn class="frase-resaltada frase-inss" data-categoria="inss" title="Click para ver detalles de inss">INSS',
                r'reclamación="frase-resaltada frase-reclamacion" data-categoria="reclamacion_administrativa" title="Click para ver detalles de reclamacion_administrativa">reclamación',
                r'EVIan class="frase-resaltada frase-inss" data-categoria="inss" title="Click para ver detalles de inss">EVI',
                r'fundamento jurídicoresaltada frase-fundamentos" data-categoria="fundamentos_juridicos" title="Click para ver detalles de fundamentos_juridicos">fundamento jurídico',
                r'Seguridad Socialse-resaltada frase-inss" data-categoria="inss" title="Click para ver detalles de inss">Seguridad Social',
                r'Instituto Nacional-resaltada frase-inss" data-categoria="inss" title="Click para ver detalles de inss">Instituto Nacional',
                r'Instituto Nacional de la Seguridad Socialdata-categoria="inss" title="Click para ver detalles de inss">Instituto Nacional de la Seguridad Social',
                r'estimamosss="frase-resaltada frase-procedimiento" data-categoria="procedimiento_legal" title="Click para ver detalles de procedimiento_legal">estimamos'
            ]
            
            for patron in patrones_malformados:
                texto = re.sub(patron, '', texto, flags=re.IGNORECASE)
            
            # SEGUNDA PASADA: Eliminar todas las etiquetas HTML restantes
            texto = re.sub(r'<[^>]*>', '', texto)
            
            # TERCERA PASADA: Eliminar caracteres de escape HTML
            texto = re.sub(r'&[a-zA-Z0-9#]+;', ' ', texto)
            
            # CUARTA PASADA: Limpiar caracteres extraños y fragmentos de etiquetas
            texto = re.sub(r'[>]+', '', texto)
            texto = re.sub(r'[<]+', '', texto)
            texto = re.sub(r'frase-[a-zA-Z]+"', '', texto)
            texto = re.sub(r'onclick="[^"]*"', '', texto)
            texto = re.sub(r'title="[^"]*"', '', texto)
            texto = re.sub(r'resaltada', '', texto)
            texto = re.sub(r'data-categoria="[^"]*"', '', texto)
            texto = re.sub(r'class="[^"]*"', '', texto)
            texto = re.sub(r'style="[^"]*"', '', texto)
            
            # QUINTA PASADA: Limpiar fragmentos adicionales comunes
            texto = re.sub(r'esaltada frase-[a-zA-Z]+"', '', texto)
            texto = re.sub(r'lesionesass=', '', texto)
            texto = re.sub(r'frase-[a-zA-Z]+" data-categoria=', '', texto)
            texto = re.sub(r'Click para ver detalles de', '', texto)
            texto = re.sub(r'frase-resaltada frase-[a-zA-Z]+', '', texto)
            
            # SEXTA PASADA: Limpiar espacios múltiples y normalizar
            texto = re.sub(r'\s+', ' ', texto)
            
            return texto.strip()
        
        # Obtener el texto original SIN procesamiento de resaltado
        texto_original = resultado.get("texto_extraido", "")
        
        # Limpiar completamente el contenido - SOLO TEXTO PLANO
        contenido_limpio = limpiar_contenido_html(texto_original)
        
        # Limpieza adicional más agresiva para eliminar cualquier resto de HTML
        import re
        contenido_limpio = re.sub(r'<[^>]*>', '', contenido_limpio)  # Eliminar cualquier etiqueta HTML restante
        contenido_limpio = re.sub(r'&[a-zA-Z0-9#]+;', ' ', contenido_limpio)  # Eliminar entidades HTML
        contenido_limpio = re.sub(r'[<>]', '', contenido_limpio)  # Eliminar caracteres < y >
        contenido_limpio = re.sub(r'\s+', ' ', contenido_limpio)  # Normalizar espacios
        contenido_limpio = contenido_limpio.strip()
        
        logger.info(f"🧹 Contenido limpio (backend, primeros 500 chars): {contenido_limpio[:500]}...")
        
        # Preparar datos para el template - SIN resaltado automático
        datos_archivo = {
            "nombre": archivo_id,
            "contenido": contenido_limpio,
            "longitud": resultado.get("longitud_texto", 0),
            "frases_clave": {},  # DESHABILITAR resaltado automático
            "prediccion": resultado.get("prediccion", {}),
            "argumentos": resultado.get("argumentos", []),
            "insights": resultado.get("insights_juridicos", []),
            "total_frases": resultado.get("total_frases_clave", 0),
            "highlight_info": {
                "frase": highlight,
                "posicion": pos,
                "index": index
            } if highlight else None,
            "resaltado_deshabilitado": True  # Flag para indicar que el resaltado está deshabilitado
        }
        
        logger.info(f"📄 Devolviendo template archivo.html para {archivo_id}")
        logger.info(f"📊 Datos del archivo: procesado={datos_archivo.get('procesado')}, frases_clave={len(datos_archivo.get('frases_clave', {}))}")
        
        return templates.TemplateResponse("archivo.html", {
            "request": request,
            "archivo": datos_archivo
        })
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error mostrando archivo {archivo_id}: {e}")
        raise HTTPException(status_code=500, detail=f"Error interno: {str(e)}")


@app.get("/api/analizar")
async def api_analizar():
    """Endpoint API para análisis"""
    try:
        return analizar_sentencias_existentes()
    except Exception as e:
        logger.error(f"Error en API: {e}")
        return {"error": f"Error al analizar: {str(e)}"}

@app.post("/api/limpiar-cache")
async def limpiar_cache():
    """Endpoint para limpiar el caché de análisis"""
    global CACHE_TIMESTAMP, ANALISIS_CACHE
    ANALISIS_CACHE = {}
    CACHE_TIMESTAMP = None
    logger.info("🗑️ Caché limpiado")
    return JSONResponse(content={"mensaje": "Caché limpiado correctamente"})


@app.get("/api/analisis-predictivo")
async def api_analisis_predictivo():
    """Endpoint API para análisis predictivo e inteligente de resoluciones"""
    try:
        # Importar funciones del módulo de análisis predictivo
        from backend.analisis_predictivo import (
            realizar_analisis_predictivo,
            generar_insights_juridicos,
            identificar_patrones_favorables,
            extraer_factores_clave,
            generar_recomendaciones,
            calcular_confianza_analisis
        )
        
        # Obtener datos base
        resultado_base = analizar_sentencias_existentes()
        
        # Realizar análisis predictivo avanzado
        analisis_predictivo = realizar_analisis_predictivo(resultado_base)
        
        # Generar insights y recomendaciones
        insights = generar_insights_juridicos(resultado_base, analisis_predictivo)
        
        # Crear respuesta estructurada para UX mejorada
        return {
            "status": "success",
            "timestamp": datetime.now().isoformat(),
            "resumen_ejecutivo": {
                "total_documentos": resultado_base.get("archivos_analizados", 0),
                "total_frases_clave": resultado_base.get("total_apariciones", 0),
                "categorias_identificadas": len(resultado_base.get("ranking_global", {})),
                "confianza_analisis": calcular_confianza_analisis(resultado_base)
            },
            "analisis_predictivo": analisis_predictivo,
            "insights_juridicos": insights,
            "patrones_favorables": identificar_patrones_favorables(resultado_base),
            "factores_clave": extraer_factores_clave(resultado_base),
            "recomendaciones": generar_recomendaciones(resultado_base, analisis_predictivo),
            "datos_estadisticos": resultado_base,
            "metadata": {
                "modelo_ia": ANALIZADOR_IA_DISPONIBLE,
                "version_analisis": "2.0.0",
                "fecha_ultima_actualizacion": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            }
        }
        
    except Exception as e:
        logger.error(f"Error en análisis predictivo: {e}")
        return {
            "status": "error",
            "error": f"Error en análisis predictivo: {str(e)}",
            "timestamp": datetime.now().isoformat()
        }


@app.get("/analisis-predictivo")
async def pagina_analisis_predictivo(request: Request):
    """Página web para el análisis predictivo"""
    return templates.TemplateResponse("analisis_predictivo.html", {"request": request})


@app.get("/analisis-discrepancias/{archivo_id}")
async def pagina_analisis_discrepancias(request: Request, archivo_id: str):
    """Página web para mostrar análisis de discrepancias de un archivo específico"""
    try:
        logger.info(f"🔍 Buscando archivo para análisis de discrepancias: {archivo_id}")
        
        # Buscar el archivo en ambos directorios
        archivo_path = None
        
        # Buscar en directorio sentencias (PDF y otros formatos)
        for extension in ["*.pdf", "*.txt", "*.docx"]:
            for archivo in Path("sentencias").glob(extension):
                if archivo_id in archivo.name or archivo.name in archivo_id:
                    archivo_path = archivo
                    logger.info(f"✅ Archivo encontrado en sentencias/: {archivo}")
                    break
            if archivo_path:
                break
        
        # Si no se encuentra, buscar en directorio uploads
        if not archivo_path:
            for extension in ["*.pdf", "*.txt", "*.docx"]:
                for archivo in Path("uploads").glob(extension):
                    if archivo_id in archivo.name or archivo.name in archivo_id:
                        archivo_path = archivo
                        logger.info(f"✅ Archivo encontrado en uploads/: {archivo}")
                        break
                if archivo_path:
                    break
        
        # Búsqueda más flexible: buscar por partes del nombre
        if not archivo_path:
            logger.info(f"🔍 Búsqueda flexible para: {archivo_id}")
            # Extraer partes del ID que podrían ser el nombre real
            partes_id = archivo_id.split('_')
            posibles_nombres = [p for p in partes_id if len(p) > 5 and not p.isdigit()]
            
            for nombre_posible in posibles_nombres:
                logger.info(f"🔍 Buscando archivos que contengan: {nombre_posible}")
                for extension in ["*.pdf", "*.txt", "*.docx"]:
                    for archivo in Path("sentencias").glob(extension):
                        if nombre_posible in archivo.name:
                            archivo_path = archivo
                            logger.info(f"✅ Archivo encontrado por búsqueda flexible en sentencias/: {archivo}")
                            break
                    if archivo_path:
                        break
                    
                    for archivo in Path("uploads").glob(extension):
                        if nombre_posible in archivo.name:
                            archivo_path = archivo
                            logger.info(f"✅ Archivo encontrado por búsqueda flexible en uploads/: {archivo}")
                            break
                    if archivo_path:
                        break
                if archivo_path:
                    break
        
        # Búsqueda final: buscar archivos que contengan 'informe' y parte del hash
        if not archivo_path and 'informe' in archivo_id:
            logger.info("🔍 Búsqueda final por archivos de informe...")
            hash_parts = [p for p in archivo_id.split('_') if len(p) == 8 and p.isalnum()]
            for hash_part in hash_parts:
                for archivo in Path("sentencias").glob("*informe*.pdf"):
                    if hash_part in archivo.name:
                        archivo_path = archivo
                        logger.info(f"✅ Archivo encontrado por búsqueda de informe en sentencias/: {archivo}")
                        break
                if archivo_path:
                    break
                    
                for archivo in Path("uploads").glob("*informe*.pdf"):
                    if hash_part in archivo.name:
                        archivo_path = archivo
                        logger.info(f"✅ Archivo encontrado por búsqueda de informe en uploads/: {archivo}")
                        break
                if archivo_path:
                    break
        
        if not archivo_path:
            logger.error(f"❌ Archivo no encontrado: {archivo_id}")
            # Listar archivos disponibles para debug
            archivos_sentencias = list(Path("sentencias").glob("*"))
            archivos_uploads = list(Path("uploads").glob("*"))
            logger.info(f"Archivos en sentencias/: {[f.name for f in archivos_sentencias]}")
            logger.info(f"Archivos en uploads/: {[f.name for f in archivos_uploads]}")
            raise HTTPException(status_code=404, detail=f"Archivo no encontrado: {archivo_id}")
        
        # Realizar análisis del archivo
        logger.info(f"🔬 Iniciando análisis de discrepancias para: {archivo_path}")
        try:
            if ANALIZADOR_IA_DISPONIBLE:
                from backend.analisis import AnalizadorLegal
                analizador = AnalizadorLegal()
                resultado = analizador.analizar_documento(str(archivo_path))
                logger.info("✅ Análisis con IA completado")
            else:
                resultado = analizador_basico.analizar_documento(str(archivo_path), archivo_path.name)
                logger.info("✅ Análisis básico completado")
        except Exception as e:
            logger.error(f"❌ Error en análisis: {e}")
            raise HTTPException(status_code=500, detail=f"Error en análisis: {str(e)}")
        
        return templates.TemplateResponse("analisis_discrepancias.html", {
            "request": request,
            "resultado": resultado
        })
        
    except Exception as e:
        logger.error(f"Error en análisis de discrepancias: {e}")
        raise HTTPException(status_code=500, detail=f"Error procesando archivo: {str(e)}")


@app.get("/test-analisis-discrepancias/{archivo_id}")
async def test_analisis_discrepancias(archivo_id: str):
    """Endpoint de prueba para verificar que el análisis de discrepancias funciona"""
    try:
        logger.info(f"🧪 Test endpoint - Buscando archivo: {archivo_id}")
        
        # Buscar archivo en ambos directorios
        archivo_path = None
        
        # Buscar en sentencias
        for extension in ["*.pdf", "*.txt", "*.docx"]:
            for archivo in Path("sentencias").glob(extension):
                if archivo_id in archivo.name:
                    archivo_path = archivo
                    break
            if archivo_path:
                break
        
        # Si no se encuentra, buscar en uploads
        if not archivo_path:
            for extension in ["*.pdf", "*.txt", "*.docx"]:
                for archivo in Path("uploads").glob(extension):
                    if archivo_id in archivo.name:
                        archivo_path = archivo
                        break
                if archivo_path:
                    break
        
        if not archivo_path:
            archivos_sentencias = [f.name for f in Path("sentencias").glob("*")]
            archivos_uploads = [f.name for f in Path("uploads").glob("*")]
            return {
                "error": f"Archivo no encontrado: {archivo_id}", 
                "archivos_en_sentencias": archivos_sentencias,
                "archivos_en_uploads": archivos_uploads,
                "archivo_id_buscado": archivo_id
            }
        
        # Realizar análisis completo
        if ANALIZADOR_IA_DISPONIBLE:
            from backend.analisis import AnalizadorLegal
            analizador = AnalizadorLegal()
            resultado = analizador.analizar_documento(str(archivo_path))
        else:
            resultado = analizador_basico.analizar_documento(str(archivo_path), archivo_path.name)
        
        return {
            "archivo_encontrado": str(archivo_path),
            "archivo_id": archivo_id,
            "status": "ok",
            "analisis_completo": resultado,
            "tipo_documento": resultado.get("analisis_discrepancias", {}).get("tipo_documento", "NO_DETECTADO"),
            "resumen_ejecutivo": resultado.get("analisis_discrepancias", {}).get("resumen_ejecutivo", "NO_DISPONIBLE")
        }
        
    except Exception as e:
        return {"error": str(e)}


@app.get("/listar-archivos")
async def listar_archivos_disponibles():
    """Endpoint para listar todos los archivos disponibles para análisis"""
    try:
        archivos_sentencias = [f.name for f in Path("sentencias").glob("*") if f.is_file()]
        archivos_uploads = [f.name for f in Path("uploads").glob("*") if f.is_file()]
        
        return {
            "archivos_sentencias": archivos_sentencias,
            "archivos_uploads": archivos_uploads,
            "total_sentencias": len(archivos_sentencias),
            "total_uploads": len(archivos_uploads),
            "directorio_actual": str(Path.cwd()),
            "directorio_sentencias": str(Path("sentencias").resolve()),
            "directorio_uploads": str(Path("uploads").resolve()),
            "existe_sentencias": Path("sentencias").exists(),
            "existe_uploads": Path("uploads").exists()
        }
        
    except Exception as e:
        return {"error": str(e)}


@app.get("/test-sts2384")
async def test_sts2384():
    """Endpoint específico para probar STS_2384_2025"""
    try:
        archivo_id = "STS_2384_2025"
        
        # Buscar archivo
        archivo_path = None
        for extension in ["*.pdf", "*.txt", "*.docx"]:
            for archivo in Path("sentencias").glob(extension):
                if archivo_id in archivo.name:
                    archivo_path = archivo
                    break
            if archivo_path:
                break
        
        if not archivo_path:
            for extension in ["*.pdf", "*.txt", "*.docx"]:
                for archivo in Path("uploads").glob(extension):
                    if archivo_id in archivo.name:
                        archivo_path = archivo
                        break
                if archivo_path:
                    break
        
        return {
            "archivo_id": archivo_id,
            "archivo_encontrado": str(archivo_path) if archivo_path else None,
            "existe": archivo_path is not None,
            "archivos_sentencias": [f.name for f in Path("sentencias").glob("STS*")],
            "archivos_uploads": [f.name for f in Path("uploads").glob("STS*")]
        }
        
    except Exception as e:
        return {"error": str(e)}


@app.get("/diagnostico")
async def pagina_diagnostico(request: Request):
    """Página de diagnóstico del sistema"""
    return templates.TemplateResponse("diagnostico.html", {"request": request})


@app.post("/api/descargar-informe-discrepancias")
async def descargar_informe_discrepancias(request: Request):
    """Genera y descarga un informe completo de discrepancias en formato Word"""
    try:
        logger.info("🔧 Iniciando generación de informe Word...")
        
        # Obtener datos del request
        datos = await request.json()
        logger.info(f"📋 Datos recibidos: {list(datos.keys())}")
        
        nombre_archivo = datos.get("nombre_archivo", "archivo_desconocido")
        analisis = datos.get("analisis_discrepancias", {})
        timestamp = datos.get("timestamp", "")
        
        logger.info(f"📄 Archivo: {nombre_archivo}")
        logger.info(f"📊 Análisis keys: {list(analisis.keys())}")
        
        # Crear documento Word con manejo de errores
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.enum.style import WD_STYLE_TYPE
            
            doc = Document()
            logger.info("✅ Documento Word creado")
            
            # Título principal
            titulo = doc.add_heading('ANÁLISIS DE DISCREPANCIAS MÉDICAS-LEGALES', 0)
            titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Información del archivo
            doc.add_heading('Información del Archivo', level=1)
            doc.add_paragraph(f"Archivo: {nombre_archivo}")
            doc.add_paragraph(f"Fecha de análisis: {timestamp}")
            doc.add_paragraph(f"Método: Análisis automático con IA")
            
            # Resumen ejecutivo
            doc.add_heading('Resumen Ejecutivo', level=1)
            discrepancias = analisis.get("discrepancias_detectadas", [])
            evidencia = analisis.get("evidencia_favorable", [])
            puntuacion = analisis.get("puntuacion_discrepancia", 0)
            probabilidad = analisis.get("probabilidad_ipp", 0)
            
            doc.add_paragraph(f"• Discrepancias detectadas: {len(discrepancias)}")
            doc.add_paragraph(f"• Evidencia favorable: {len(evidencia)} elementos")
            doc.add_paragraph(f"• Puntuación discrepancia: {puntuacion}/100")
            doc.add_paragraph(f"• Probabilidad IPP: {probabilidad:.1%}")
            
            # Conclusión
            if probabilidad >= 0.7:
                conclusion = "ALTA PROBABILIDAD DE IPP"
            elif probabilidad >= 0.5:
                conclusion = "PROBABILIDAD MEDIA DE IPP"
            else:
                conclusion = "BAJA PROBABILIDAD DE IPP"
            
            doc.add_paragraph(f"Conclusión: {conclusion}")
            
            # Discrepancias detectadas
            if discrepancias:
                doc.add_heading('Discrepancias Detectadas', level=1)
                for i, disc in enumerate(discrepancias, 1):
                    try:
                        doc.add_heading(f"{i}. {disc.get('tipo', '').replace('_', ' ').title()}", level=2)
                        doc.add_paragraph(f"Descripción: {disc.get('descripcion', '')}")
                        doc.add_paragraph(f"Severidad: {disc.get('severidad', '')}")
                        doc.add_paragraph(f"Argumento jurídico: {disc.get('argumento_juridico', '')}")
                    except Exception as e:
                        logger.warning(f"Error procesando discrepancia {i}: {e}")
                        doc.add_paragraph(f"Error procesando discrepancia {i}")
            
            # Evidencia favorable
            if evidencia:
                doc.add_heading('Evidencia Favorable para IPP', level=1)
                for i, ev in enumerate(evidencia, 1):
                    try:
                        doc.add_heading(f"{i}. {ev.get('tipo', '').replace('_', ' ').title()}", level=2)
                        doc.add_paragraph(f"Descripción: {ev.get('descripcion', '')}")
                        doc.add_paragraph(f"Relevancia: {ev.get('relevancia', '')}")
                        doc.add_paragraph(f"Argumento: {ev.get('argumento', '')}")
                    except Exception as e:
                        logger.warning(f"Error procesando evidencia {i}: {e}")
                        doc.add_paragraph(f"Error procesando evidencia {i}")
            
            # Argumentos jurídicos
            argumentos = analisis.get("argumentos_juridicos", [])
            if argumentos:
                doc.add_heading('Argumentos Jurídicos Generados', level=1)
                for i, arg in enumerate(argumentos, 1):
                    try:
                        doc.add_heading(f"{i}. {arg.get('titulo', '')}", level=2)
                        doc.add_paragraph(f"Contenido: {arg.get('contenido', '')}")
                        doc.add_paragraph(f"Fuerza: {arg.get('fuerza', '')}")
                    except Exception as e:
                        logger.warning(f"Error procesando argumento {i}: {e}")
                        doc.add_paragraph(f"Error procesando argumento {i}")
            
            # Recomendaciones de defensa
            recomendaciones = analisis.get("recomendaciones_defensa", [])
            if recomendaciones:
                doc.add_heading('Recomendaciones de Defensa', level=1)
                for i, rec in enumerate(recomendaciones, 1):
                    try:
                        doc.add_heading(f"{i}. {rec.get('titulo', '')}", level=2)
                        doc.add_paragraph(f"Contenido: {rec.get('contenido', '')}")
                        doc.add_paragraph(f"Prioridad: {rec.get('prioridad', '')}")
                        
                        acciones = rec.get('acciones', [])
                        if acciones:
                            doc.add_paragraph("Acciones recomendadas:")
                            for accion in acciones:
                                doc.add_paragraph(f"• {accion}", style='List Bullet')
                    except Exception as e:
                        logger.warning(f"Error procesando recomendación {i}: {e}")
                        doc.add_paragraph(f"Error procesando recomendación {i}")
            
            # Contradicciones internas
            contradicciones = analisis.get("contradicciones_internas", [])
            if contradicciones:
                doc.add_heading('Contradicciones Internas Detectadas', level=1)
                for i, cont in enumerate(contradicciones, 1):
                    try:
                        doc.add_heading(f"{i}. Contradicción Interna", level=2)
                        doc.add_paragraph(f"Descripción: {cont.get('descripcion', '')}")
                        doc.add_paragraph(f"Texto detectado: {cont.get('texto', '')}")
                        doc.add_paragraph(f"Argumento: {cont.get('argumento', '')}")
                    except Exception as e:
                        logger.warning(f"Error procesando contradicción {i}: {e}")
                        doc.add_paragraph(f"Error procesando contradicción {i}")
            
            logger.info("✅ Contenido del documento generado")
            
        except Exception as doc_error:
            logger.error(f"❌ Error creando documento Word: {doc_error}")
            raise doc_error
        
        # Guardar en memoria con manejo de errores
        try:
            from io import BytesIO
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            # Verificar que el buffer tiene contenido
            content = buffer.getvalue()
            logger.info(f"📦 Tamaño del archivo generado: {len(content)} bytes")
            
            if len(content) == 0:
                raise Exception("El archivo generado está vacío")
            
        except Exception as save_error:
            logger.error(f"❌ Error guardando documento: {save_error}")
            raise save_error
        
        # Preparar respuesta
        try:
            from fastapi.responses import Response
            
            # Limpiar nombre de archivo para evitar caracteres problemáticos
            nombre_limpio = re.sub(r'[<>:"/\\|?*]', '_', nombre_archivo)
            nombre_limpio = nombre_limpio.replace('.pdf', '').replace('.txt', '').replace('.docx', '')
            
            logger.info(f"📤 Enviando archivo: informe_discrepancias_{nombre_limpio}.docx")
            
            return Response(
                content=content,
                media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                headers={
                    "Content-Disposition": f"attachment; filename=informe_discrepancias_{nombre_limpio}.docx",
                    "Content-Length": str(len(content))
                }
            )
            
        except Exception as response_error:
            logger.error(f"❌ Error preparando respuesta: {response_error}")
            raise response_error
        
    except Exception as e:
        logger.error(f"❌ Error general generando informe Word: {e}")
        logger.error(f"❌ Tipo de error: {type(e).__name__}")
        import traceback
        logger.error(f"❌ Traceback: {traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=f"Error generando informe: {str(e)}")


@app.get("/api/test-word")
async def test_word_generation():
    """Endpoint de prueba para verificar que la generación de Word funciona"""
    try:
        logger.info("🧪 Probando generación de Word...")
        
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        # Crear documento de prueba
        doc = Document()
        
        # Título
        titulo = doc.add_heading('PRUEBA DE GENERACIÓN WORD', 0)
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Contenido simple
        doc.add_heading('Información de Prueba', level=1)
        doc.add_paragraph("Este es un documento de prueba para verificar que la generación de archivos Word funciona correctamente.")
        doc.add_paragraph("Fecha de prueba: " + datetime.now().strftime("%Y-%m-%d %H:%M:%S"))
        doc.add_paragraph("Estado: ✅ FUNCIONANDO")
        
        # Guardar en memoria
        from io import BytesIO
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        
        content = buffer.getvalue()
        logger.info(f"✅ Documento de prueba generado: {len(content)} bytes")
        
        from fastapi.responses import Response
        return Response(
            content=content,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": "attachment; filename=test_word.docx",
                "Content-Length": str(len(content))
            }
        )
        
    except Exception as e:
        logger.error(f"❌ Error en prueba Word: {e}")
        import traceback
        logger.error(f"❌ Traceback: {traceback.format_exc()}")
        raise HTTPException(status_code=500, detail=f"Error en prueba: {str(e)}")


@app.post("/api/descargar-resumen-pdf")
async def descargar_resumen_pdf(request: Request):
    """Genera y descarga un resumen en formato PDF"""
    try:
        from reportlab.lib.pagesizes import letter, A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from io import BytesIO
        
        # Obtener datos del request
        datos = await request.json()
        contenido = datos.get("contenido", "")
        nombre_archivo = datos.get("nombre_archivo", "archivo_desconocido")
        
        # Crear documento PDF en memoria
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=18)
        
        # Estilos
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=16,
            spaceAfter=30,
            alignment=1  # Centrado
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            spaceAfter=12,
            spaceBefore=20
        )
        
        # Contenido
        story = []
        
        # Convertir HTML a texto plano para PDF
        import re
        contenido_texto = re.sub(r'<[^>]+>', '', contenido)
        contenido_texto = contenido_texto.replace('&nbsp;', ' ')
        
        # Dividir en párrafos
        parrafos = contenido_texto.split('\n')
        
        for parrafo in parrafos:
            parrafo = parrafo.strip()
            if not parrafo:
                continue
                
            if parrafo.startswith('ANÁLISIS DE DISCREPANCIAS'):
                story.append(Paragraph(parrafo, title_style))
            elif parrafo.startswith('Archivo:') or parrafo.startswith('Fecha de análisis:'):
                story.append(Paragraph(parrafo, styles['Normal']))
            elif parrafo.startswith('RESUMEN EJECUTIVO') or parrafo.startswith('DISCREPANCIAS') or parrafo.startswith('EVIDENCIA') or parrafo.startswith('ARGUMENTOS') or parrafo.startswith('RECOMENDACIONES'):
                story.append(Spacer(1, 12))
                story.append(Paragraph(parrafo, heading_style))
            elif parrafo.startswith('•'):
                story.append(Paragraph(parrafo, styles['Normal']))
            else:
                story.append(Paragraph(parrafo, styles['Normal']))
            
            story.append(Spacer(1, 6))
        
        # Construir PDF
        doc.build(story)
        buffer.seek(0)
        
        # Preparar respuesta
        from fastapi.responses import Response
        return Response(
            content=buffer.getvalue(),
            media_type="application/pdf",
            headers={
                "Content-Disposition": f"attachment; filename=resumen_discrepancias_{nombre_archivo.replace('.pdf', '')}.pdf"
            }
        )
        
    except Exception as e:
        logger.error(f"Error generando PDF: {e}")
        raise HTTPException(status_code=500, detail=f"Error generando PDF: {str(e)}")


# ====== DEMANDA BASE: generación desde fallos y fundamentos ======
def _leer_texto_archivo_simple(path: Path) -> str:
    try:
        if path.suffix.lower() == '.txt':
            for enc in ['utf-8', 'latin-1', 'cp1252']:
                try:
                    return path.read_text(encoding=enc)
                except UnicodeDecodeError:
                    continue
            return path.read_text(errors='ignore')
        else:
            # Delegar a analizador para extraer texto
            if ANALIZADOR_IA_DISPONIBLE:
                from backend.analisis import AnalizadorLegal
                analizador = AnalizadorLegal()
                res = analizador.analizar_documento(str(path))
            else:
                res = analizador_basico.analizar_documento(str(path), path.name)
            return res.get('texto_extraido', '') or ''
    except Exception:
        return ''


def _extraer_seccion(texto: str, encabezados: List[str], fin_encabezados: List[str]) -> str:
    if not texto:
        return ''
    t = texto
    import re as _re
    # Construir regex de inicio y fin
    ini = r"|".join([_re.escape(h) for h in encabezados])
    fin = r"|".join([_re.escape(h) for h in fin_encabezados])
    patron_ini = _re.compile(rf"(?i)(?:^|\n)\s*(?:{ini})\b[:\-\s]*", _re.MULTILINE)
    patron_fin = _re.compile(rf"(?i)(?:^|\n)\s*(?:{fin})\b", _re.MULTILINE)
    m = patron_ini.search(t)
    if not m:
        return ''
    start = m.end()
    m2 = patron_fin.search(t, start)
    end = m2.start() if m2 else len(t)
    seccion = t[start:end].strip()
    # Limpiar artefactos HTML simples
    seccion = _re.sub(r"<[^>]+>", " ", seccion)
    seccion = _re.sub(r"\s+", " ", seccion).strip()
    return seccion


def _generar_demanda_base_para(paths: List[Path], meta: Dict[str, Any] = None) -> Dict[str, Any]:
    documentos: List[Dict[str, Any]] = []
    for p in paths:
        texto = _leer_texto_archivo_simple(p)
        fallo = _extraer_seccion(
            texto,
            encabezados=["FALLO", "PARTE DISPOSITIVA", "RESUELVO", "RESOLVEMOS"],
            fin_encabezados=["FUNDAMENTOS", "FUNDAMENTOS DE HECHO", "HECHOS", "FUNDAMENTOS DE DERECHO", "ANTECEDENTES", "SEGUNDO", "TERCERO"]
        )
        fundamentos = _extraer_seccion(
            texto,
            encabezados=["FUNDAMENTOS DE HECHO", "HECHOS PROBADOS", "ANTECEDENTES DE HECHO"],
            fin_encabezados=["FUNDAMENTOS DE DERECHO", "PARTE DISPOSITIVA", "FALLO", "RESUELVO", "RESOLVEMOS"]
        )
        # Resumenes breves
        fundamentos_resumen = _resumir_fundamentos(texto)
        fallo_breve = (fallo or "").strip()
        if len(fallo_breve) > 400:
            fallo_breve = fallo_breve[:400].rstrip() + "…"
        documentos.append({
            "nombre": p.name,
            "fallo": fallo,
            "fundamentos": fundamentos,
            "fallo_breve": fallo_breve,
            "fundamentos_resumen": fundamentos_resumen
        })

    meta = meta or {}
    nombre = meta.get("nombre", "[NOMBRE DEMANDANTE]")
    dni = meta.get("dni", "[DNI]")
    domicilio = meta.get("domicilio", "[DOMICILIO A EFECTOS]")
    letrado = meta.get("letrado", "[LETRADO/A]")
    empresa = meta.get("empresa", "[EMPRESA]")
    profesion = meta.get("profesion", "[PROFESIÓN HABITUAL]")
    grado_principal = meta.get("grado_principal", "Incapacidad Permanente Total")
    grado_subsidiario = meta.get("grado_subsidiario", "Incapacidad Permanente Parcial")
    base_reguladora = meta.get("base_reguladora", "[BASE REGULADORA]")
    indemnizacion_parcial = meta.get("indemnizacion_parcial", "24 mensualidades")
    mutua = meta.get("mutua", "[MUTUA]")
    
    # Campos de HECHOS
    relacion_laboral = meta.get("relacion_laboral", f"La demandante presta servicios como {profesion} para la empresa {empresa}.")
    contingencia_evolucion = meta.get("contingencia_evolucion", "[Accidente de trabajo / enfermedad común], con periodos de IT y secuelas actuales.")
    actuaciones_administrativas = meta.get("actuaciones_administrativas", "(EVI, inicio IP, audiencia, resolución del INSS y Reclamación Previa).")
    cuadro_clinico = meta.get("cuadro_clinico", "[Describir secuelas relevantes y su impacto en las tareas fundamentales].")
    
    # Plantilla base de demanda (formal)
    cuerpo = {
        "encabezado": "AL JUZGADO DE LO SOCIAL QUE POR TURNO CORRESPONDA\n\n",
        "parte": (
            f"D./Dña. {nombre}, con DNI {dni}, y domicilio a efectos de notificaciones en {domicilio}, "
            f"representado/a por Letrado/a {letrado}, ante el Juzgado comparece y DICE:\n\n"
        ),
        "hechos": None,  # se compone más abajo
        "fundamentos_derecho": (
            "FUNDAMENTOS DE DERECHO\n"
            "I. Jurisdicción y competencia (arts. 2 y 6 LRJS).\n"
            "II. Legitimación activa y pasiva (LRJS).\n"
            "III. Fondo del asunto. Arts. 193 y 194 LGSS (grados de incapacidad). Art. 194.2 LGSS (IPP = 24 mensualidades).\n"
            "IV. Doctrina jurisprudencial aplicable (STS 04-07-2025, rec. 1096/2024, sobre IPP subsidiaria; TSJ Castilla y León sobre limitaciones en limpiadoras).\n"
            "V. Principios y derechos constitucionales (art. 24 CE tutela judicial efectiva; 9.3 CE interdicción de la arbitrariedad).\n\n"
        ),
        "peticion": (
            "SUPLICO AL JUZGADO:\n"
            f"Primero.- Que, estimando la demanda, se declare a la actora afecta a {grado_principal} para su profesión habitual de {profesion}, derivada de [contingencia], con derecho a la prestación correspondiente sobre una base reguladora de {base_reguladora}.\n"
            f"Subsidiariamente.- Que, para el caso de no apreciarse lo anterior, se declare la {grado_subsidiario}, con derecho a la indemnización de {indemnizacion_parcial} de la base reguladora, a cargo de la Mutua {mutua} en caso de accidente de trabajo.\n"
            "Con expresa condena en costas a la parte demandada en los términos legalmente procedentes.\n\n"
        ),
    }

    # HECHOS (limpios y propios)
    cuerpo["hechos"] = (
        "HECHOS\n"
        f"1. Relación laboral. {relacion_laboral}\n"
        f"2. Contingencia y evolución. {contingencia_evolucion}\n"
        f"3. Actuaciones administrativas. {actuaciones_administrativas}\n"
        f"4. Cuadro clínico y limitaciones. {cuadro_clinico}\n\n"
    )

    # Resumen de fallos extraídos
    resumen_fallos = []
    for d in documentos:
        if d["fallo"]:
            resumen_fallos.append(f"— {d['nombre']}: {d['fallo']}")

    anexos = []
    for i, d in enumerate(documentos, start=1):
        if d["fallo"]:
            anexos.append(f"Documento nº {i}: Parte dispositiva de {d['nombre']}")

    # Jurisprudencia de apoyo (extractada)
    juris = []
    for d in documentos:
        linea = f"— {d['nombre']}: {d.get('fallo_breve','').strip()}"
        if d.get("fundamentos_resumen"):
            linea += "\n   Fundamentos: " + "; ".join(d["fundamentos_resumen"][:2])
        juris.append(linea.strip())

    doc_txt = (
        cuerpo["encabezado"] +
        cuerpo["parte"] +
        ("JURISPRUDENCIA DE APOYO (extracto de fallos)\n" + "\n\n".join(juris) + "\n\n" if juris else "") +
        cuerpo["hechos"] +
        cuerpo["fundamentos_derecho"] +
        cuerpo["peticion"] +
        ("ANEXO: Relación de documentos adjuntos\n" + "\n".join(anexos) + "\n" if anexos else "")
    )

    return {
        "documentos": documentos,
        "texto": doc_txt
    }


def _generar_demanda_docx(paths: List[Path], meta: Dict[str, Any] = None) -> BytesIO:
    """Genera un documento DOCX con formato profesional para la demanda"""
    # Obtener el texto de la demanda
    demanda_data = _generar_demanda_base_para(paths, meta)
    texto_demanda = demanda_data.get("texto", "")
    
    # Crear documento Word
    doc = Document()
    
    # Configurar estilos del documento
    styles = doc.styles
    
    # Estilo para el título principal
    title_style = styles.add_style('DemandaTitle', WD_STYLE_TYPE.PARAGRAPH)
    title_font = title_style.font
    title_font.name = 'Times New Roman'
    title_font.size = Pt(16)
    title_font.bold = True
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_style.paragraph_format.space_after = Pt(12)
    
    # Estilo para encabezados de sección
    heading_style = styles.add_style('DemandaHeading', WD_STYLE_TYPE.PARAGRAPH)
    heading_font = heading_style.font
    heading_font.name = 'Times New Roman'
    heading_font.size = Pt(14)
    heading_font.bold = True
    heading_style.paragraph_format.space_before = Pt(12)
    heading_style.paragraph_format.space_after = Pt(6)
    
    # Estilo para texto normal
    normal_style = styles.add_style('DemandaNormal', WD_STYLE_TYPE.PARAGRAPH)
    normal_font = normal_style.font
    normal_font.name = 'Times New Roman'
    normal_font.size = Pt(12)
    normal_style.paragraph_format.space_after = Pt(6)
    normal_style.paragraph_format.line_spacing = 1.5
    
    # Estilo para texto con sangría
    indent_style = styles.add_style('DemandaIndent', WD_STYLE_TYPE.PARAGRAPH)
    indent_font = indent_style.font
    indent_font.name = 'Times New Roman'
    indent_font.size = Pt(12)
    indent_style.paragraph_format.left_indent = Inches(0.5)
    indent_style.paragraph_format.space_after = Pt(6)
    indent_style.paragraph_format.line_spacing = 1.5
    
    # Configurar márgenes del documento
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)
    
    # Procesar el texto línea por línea
    lines = texto_demanda.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            # Línea vacía - agregar espacio
            doc.add_paragraph('', style='DemandaNormal')
            continue
            
        # Detectar tipo de contenido
        if line.upper().startswith('AL JUZGADO'):
            # Título principal
            p = doc.add_paragraph(line, style='DemandaTitle')
        elif any(line.upper().startswith(keyword) for keyword in ['HECHOS', 'FUNDAMENTOS', 'SUPLICO', 'JURISPRUDENCIA', 'ANEXO']):
            # Encabezado de sección
            p = doc.add_paragraph(line, style='DemandaHeading')
        elif line.startswith('D./Dña.') or line.startswith('D./Dña'):
            # Información de la parte
            p = doc.add_paragraph(line, style='DemandaNormal')
        elif line.startswith(('Primero.-', 'Segundo.-', 'Tercero.-', 'Subsidiariamente.-')):
            # Peticiones numeradas
            p = doc.add_paragraph(line, style='DemandaIndent')
        elif line.startswith(('—', '•', '1.', '2.', '3.', '4.')):
            # Lista con viñetas
            p = doc.add_paragraph(line, style='DemandaIndent')
        elif line.startswith('I.') or line.startswith('II.') or line.startswith('III.') or line.startswith('IV.') or line.startswith('V.'):
            # Numeración romana
            p = doc.add_paragraph(line, style='DemandaIndent')
        else:
            # Texto normal
            p = doc.add_paragraph(line, style='DemandaNormal')
    
    # Guardar en memoria
    docx_buffer = BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)
    
    return docx_buffer


# ====== EXTRACTOR ESTRUCTURADO PARA DEMANDA ======
def _inferir_instancia_desde_texto(texto: str) -> str:
    t = (texto or '').lower()
    if 'tribunal supremo' in t:
        return 'TS'
    if 'tribunal superior de justicia' in t or 'tsj' in t:
        return 'TSJ'
    return 'otra'


def _extraer_primera_fecha(texto: str) -> Optional[str]:
    import re as _re
    for patron in [r"\b\d{1,2}[\-/\.\s]\d{1,2}[\-/\.\s]\d{2,4}\b", r"\b\d{4}[\-/\.]\d{1,2}[\-/\.]\d{1,2}\b"]:
        m = _re.search(patron, texto)
        if m:
            return m.group(0)
    return None


def _extraer_por_regex(texto: str, regex: str, group: int = 1) -> Optional[str]:
    import re as _re
    m = _re.search(regex, texto, _re.IGNORECASE)
    return m.group(group).strip() if m else None


def _resumir_fundamentos(texto: str) -> List[str]:
    seccion = _extraer_seccion(
        texto,
        ["FUNDAMENTOS DE DERECHO", "FUNDAMENTOS"],
        ["FALLO", "PARTE DISPOSITIVA", "RESUELVO", "RESOLVEMOS", "SUPLICO", "HECHOS"]
    )
    if not seccion:
        return []
    # dividir en frases y tomar las 3 primeras relevantes
    import re as _re
    frases = [f.strip() for f in _re.split(r"(?<=[\.!?])\s+", seccion) if len(f.strip()) > 30]
    return frases[:3]


@app.post("/api/extract/demanda")
async def api_extract_demanda(payload: Dict[str, Any]):
    try:
        nombres: List[str] = payload.get("nombres_archivo") or []
        if not isinstance(nombres, list) or not nombres:
            raise HTTPException(status_code=400, detail="Debe indicar 'nombres_archivo' (lista)")
        paths = [SENTENCIAS_DIR / n for n in nombres if (SENTENCIAS_DIR / n).exists()]
        if not paths:
            raise HTTPException(status_code=404, detail="No se encontraron los archivos indicados")

        docs_out: List[Dict[str, Any]] = []
        # Sugerencias extraídas dinámicamente del documento
        sugerencias: Dict[str, Any] = {
            "profesion": None, 
            "empresa": None, 
            "mutua": None, 
            "base_reguladora": None
        }

        for p in paths:
            texto = _leer_texto_archivo_simple(p)
            fallo = _extraer_seccion(texto, ["FALLO", "PARTE DISPOSITIVA", "RESUELVO", "RESOLVEMOS"], ["FUNDAMENTOS", "HECHOS", "ANTECEDENTES"]) or ""
            fundamentos_resumen = _resumir_fundamentos(texto)
            instancia = _inferir_instancia_desde_texto(texto)
            fecha = _extraer_primera_fecha(texto)
            organo = None
            # aproximación al órgano
            for k in ["TRIBUNAL SUPREMO", "TRIBUNAL SUPERIOR DE JUSTICIA", "AUDIENCIA", "JUZGADO DE LO SOCIAL"]:
                if k.lower() in (texto or '').lower():
                    organo = k.title()
                    break

            # Extraer sugerencias del documento
            if sugerencias["profesion"] is None:
                profesion = _extraer_por_regex(texto, r"profesi[oó]n\s+habitual\s+(de|:)?\s*([A-Za-zÁÉÍÓÚÜÑáéíóúüñ ]+)", 2)
                if profesion:
                    sugerencias["profesion"] = profesion
            if sugerencias["empresa"] is None:
                empresa = _extraer_por_regex(texto, r"emplead[oa]\s+por\s+([A-ZÁÉÍÓÚÜÑa-záéíóúüñ0-9 .,&;-]+)")
                if empresa:
                    sugerencias["empresa"] = empresa
            if sugerencias["mutua"] is None:
                mutua = _extraer_por_regex(texto, r"mutua\s+([A-ZÁÉÍÓÚÜÑ][A-Za-zÁÉÍÓÚÜÑ ]+)")
                if mutua:
                    sugerencias["mutua"] = mutua
            if sugerencias["base_reguladora"] is None:
                br = _extraer_por_regex(texto, r"base\s+reguladora[^0-9]*([0-9\.,]+)")
                if br:
                    sugerencias["base_reguladora"] = br

            docs_out.append({
                "archivo": p.name,
                "instancia": instancia,
                "fecha": fecha,
                "organo": organo,
                "fallo": fallo,
                "fundamentos_resumen": fundamentos_resumen
            })

        return {"documentos": docs_out, "sugerencias_meta": sugerencias}

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error en extracción de demanda: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/sugerencias-test")
async def api_sugerencias_test():
    """Endpoint de prueba para verificar las sugerencias extraídas"""
    # Usar el primer documento disponible para extraer sugerencias
    archivos = list(SENTENCIAS_DIR.glob("*.pdf")) + list(SENTENCIAS_DIR.glob("*.txt"))
    if archivos:
        texto = _leer_texto_archivo_simple(archivos[0])
        sugerencias = {
            "profesion": None, 
            "empresa": None, 
            "mutua": None, 
            "base_reguladora": None
        }
        
        # Extraer sugerencias del documento
        profesion = _extraer_por_regex(texto, r"profesi[oó]n\s+habitual\s+(de|:)?\s*([A-Za-zÁÉÍÓÚÜÑáéíóúüñ ]+)", 2)
        if profesion:
            sugerencias["profesion"] = profesion
        empresa = _extraer_por_regex(texto, r"emplead[oa]\s+por\s+([A-ZÁÉÍÓÚÜÑa-záéíóúüñ0-9 .,&;-]+)")
        if empresa:
            sugerencias["empresa"] = empresa
        mutua = _extraer_por_regex(texto, r"mutua\s+([A-ZÁÉÍÓÚÜÑ][A-Za-zÁÉÍÓÚÜÑ ]+)")
        if mutua:
            sugerencias["mutua"] = mutua
        br = _extraer_por_regex(texto, r"base\s+reguladora[^0-9]*([0-9\.,]+)")
        if br:
            sugerencias["base_reguladora"] = br
    else:
        sugerencias = {
            "profesion": None, 
            "empresa": None, 
            "mutua": None, 
            "base_reguladora": None
        }
    
    return {"sugerencias_meta": sugerencias}


@app.get("/api/diagnostico/ia")
async def api_diagnostico_ia():
    """Endpoint para diagnóstico detallado del modelo IA"""
    try:
        # Verificar modelos disponibles
        modelos = {}
        
        # Verificar modelo TF-IDF original
        tfidf_path = Path("models/modelo_legal.pkl")
        if tfidf_path.exists():
            try:
                with open(tfidf_path, 'rb') as f:
                    import pickle
                    data = pickle.load(f)
                    modelos["tfidf_original"] = {
                        "existe": True,
                        "tipo": "TF-IDF",
                        "componentes": list(data.keys()) if isinstance(data, dict) else ["modelo"]
                    }
            except Exception as e:
                modelos["tfidf_original"] = {
                    "existe": True,
                    "error": str(e)
                }
        else:
            modelos["tfidf_original"] = {"existe": False}
        
        # Verificar modelo SBERT
        sbert_path = Path("models/modelo_legal_sbert.pkl")
        if sbert_path.exists():
            try:
                with open(sbert_path, 'rb') as f:
                    import pickle
                    data = pickle.load(f)
                    modelos["sbert"] = {
                        "existe": True,
                        "encoder_name": data.get("encoder_name", "desconocido"),
                        "tipo": "SBERT" if data.get("encoder_name") != "tfidf" else "TF-IDF Fallback",
                        "componentes": list(data.keys()) if isinstance(data, dict) else ["modelo"]
                    }
            except Exception as e:
                modelos["sbert"] = {
                    "existe": True,
                    "error": str(e)
                }
        else:
            modelos["sbert"] = {"existe": False}
        
        # Verificar analizador
        try:
            from backend.analisis import AnalizadorLegal
            analizador = AnalizadorLegal()
            
            estado_analizador = {
                "sbert_encoder": analizador.sbert_encoder is not None,
                "sbert_clf": analizador.sbert_clf is not None,
                "vectorizador": analizador.vectorizador is not None,
                "clasificador": analizador.clasificador is not None,
                "modelo": analizador.modelo is not None
            }
            
            # Determinar qué modelo se está usando
            if analizador.sbert_encoder is not None and analizador.sbert_clf is not None:
                if hasattr(analizador.sbert_encoder, 'encode'):
                    modelo_activo = "SBERT Real"
                else:
                    modelo_activo = "TF-IDF Fallback"
            elif analizador.vectorizador is not None and analizador.clasificador is not None:
                modelo_activo = "TF-IDF"
            else:
                modelo_activo = "Análisis por Reglas"
                
        except Exception as e:
            estado_analizador = {"error": str(e)}
            modelo_activo = "Error"
        
        return {
            "modelos": modelos,
            "estado_analizador": estado_analizador,
            "modelo_activo": modelo_activo,
            "ia_disponible": ANALIZADOR_IA_DISPONIBLE,
            "timestamp": datetime.now().isoformat()
        }
        
    except Exception as e:
        return {
            "error": str(e),
            "timestamp": datetime.now().isoformat()
        }

@app.post("/api/demanda-base")
async def api_demanda_base(payload: Dict[str, Any]):
    try:
        nombres: List[str] = payload.get("nombres_archivo") or []
        if not isinstance(nombres, list) or not nombres:
            raise HTTPException(status_code=400, detail="Debe indicar 'nombres_archivo' (lista)")
        paths: List[Path] = []
        for n in nombres:
            p = SENTENCIAS_DIR / n
            if p.exists():
                paths.append(p)
        if not paths:
            raise HTTPException(status_code=404, detail="No se encontraron los archivos indicados")
        meta = payload.get("meta") if isinstance(payload.get("meta"), dict) else {}
        doc = _generar_demanda_base_para(paths, meta)
        return doc
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error generando demanda base: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/demanda-base/txt")
async def api_demanda_base_txt(payload: Dict[str, Any]):
    try:
        nombres: List[str] = payload.get("nombres_archivo") or []
        if not isinstance(nombres, list) or not nombres:
            raise HTTPException(status_code=400, detail="Debe indicar 'nombres_archivo' (lista)")
        paths = [SENTENCIAS_DIR / n for n in nombres if (SENTENCIAS_DIR / n).exists()]
        if not paths:
            raise HTTPException(status_code=404, detail="No se encontraron los archivos indicados")
        meta = payload.get("meta") if isinstance(payload.get("meta"), dict) else {}
        doc = _generar_demanda_base_para(paths, meta)
        filename = f"demanda_base_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
        headers = {"Content-Disposition": f"attachment; filename={filename}"}
        return PlainTextResponse(content=doc.get("texto", ""), headers=headers)
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error generando TXT demanda base: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/demanda-base/docx")
async def api_demanda_base_docx(payload: Dict[str, Any]):
    """Genera una demanda en formato DOCX con formato profesional"""
    try:
        nombres: List[str] = payload.get("nombres_archivo") or []
        if not isinstance(nombres, list) or not nombres:
            raise HTTPException(status_code=400, detail="Debe indicar 'nombres_archivo' (lista)")
        paths = [SENTENCIAS_DIR / n for n in nombres if (SENTENCIAS_DIR / n).exists()]
        if not paths:
            raise HTTPException(status_code=404, detail="No se encontraron los archivos indicados")
        meta = payload.get("meta") if isinstance(payload.get("meta"), dict) else {}
        
        # Generar documento DOCX
        docx_buffer = _generar_demanda_docx(paths, meta)
        
        # Crear nombre de archivo
        nombre_demandante = meta.get("nombre", "demandante").replace(" ", "_")
        filename = f"demanda_{nombre_demandante}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        
        # Configurar headers para descarga
        headers = {
            "Content-Disposition": f"attachment; filename={filename}",
            "Content-Type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        }
        
        # Devolver el documento DOCX
        from fastapi.responses import Response
        return Response(
            content=docx_buffer.getvalue(),
            headers=headers,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error generando DOCX demanda base: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# ====== LISTAR / ELIMINAR DOCUMENTOS ======
@app.get("/api/documentos")
async def listar_documentos():
    """Lista documentos disponibles en sentencias/ y uploads/."""
    docs = []
    for carpeta in [SENTENCIAS_DIR, UPLOADS_DIR]:
        if carpeta.exists():
            for f in carpeta.iterdir():
                if f.is_file() and f.suffix.lower() in ['.pdf', '.txt']:
                    docs.append({
                        "nombre": f.name,
                        "ruta": str(f),
                        "carpeta": 'sentencias' if carpeta == SENTENCIAS_DIR else 'uploads',
                        "tamaño": f.stat().st_size,
                        "modificado": datetime.fromtimestamp(f.stat().st_mtime).isoformat()
                    })
    return {"documentos": sorted(docs, key=lambda x: x["modificado"], reverse=True)}


@app.delete("/api/documentos")
async def eliminar_documento_api(payload: DeleteDocumentPayload):
    """Elimina un documento por nombre desde sentencias/ o uploads/."""
    nombre = payload.nombre_archivo
    if not nombre:
        raise HTTPException(status_code=400, detail="Nombre de archivo requerido")
    candidatos = [SENTENCIAS_DIR / nombre, UPLOADS_DIR / nombre]
    encontrado = None
    for p in candidatos:
        if p.exists() and p.is_file():
            encontrado = p
            break
    if not encontrado:
        raise HTTPException(status_code=404, detail="Documento no encontrado")
    try:
        encontrado.unlink()
        return {"status": "ok", "eliminado": nombre}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"No se pudo eliminar: {e}")


@app.get("/health")
async def health_check():
    """Endpoint de salud del sistema"""
    return {
        "status": "ok",
        "version": "2.0.0",
        "timestamp": datetime.now().isoformat(),
        "ia_disponible": ANALIZADOR_IA_DISPONIBLE,
        "directorios": {
            "sentencias": str(SENTENCIAS_DIR),
            "uploads": str(UPLOADS_DIR),
            "models": str(MODELS_DIR)
        }
    }


@app.get("/api/documento/{nombre_archivo}")
async def obtener_documento(nombre_archivo: str):
    """Obtiene detalles de un documento específico"""
    try:
        # Decodificar el nombre del archivo
        nombre_decodificado = nombre_archivo
        
        # Buscar el archivo en el directorio de sentencias
        archivo_path = SENTENCIAS_DIR / nombre_decodificado
        
        if not archivo_path.exists():
            raise HTTPException(status_code=404, detail=f"Documento '{nombre_decodificado}' no encontrado")
        
        # Obtener información del archivo
        stat = archivo_path.stat()
        tamaño = f"{stat.st_size / 1024:.1f} KB"
        fecha_modificacion = datetime.fromtimestamp(stat.st_mtime).strftime("%Y-%m-%d %H:%M:%S")
        
        # Intentar analizar el documento si no ha sido analizado
        try:
            if ANALIZADOR_IA_DISPONIBLE:
                from backend.analisis import AnalizadorLegal
                analizador = AnalizadorLegal()
                resultado = analizador.analizar_documento(str(archivo_path))
            else:
                resultado = analizador_basico.analizar_documento(str(archivo_path), nombre_decodificado)
            
            if resultado.get("procesado"):
                return {
                    "nombre": nombre_decodificado,
                    "procesado": True,
                    "tamaño": tamaño,
                    "fecha_procesamiento": fecha_modificacion,
                    "frases_clave": resultado.get("frases_clave", {}),
                    "total_frases": sum(info["total"] for info in resultado.get("frases_clave", {}).values()),
                    "resumen": resultado.get("resumen", ""),
                    "tipo_documento": resultado.get("tipo_documento", "desconocido")
                }
            else:
                return {
                    "nombre": nombre_decodificado,
                    "procesado": False,
                    "tamaño": tamaño,
                    "fecha_procesamiento": fecha_modificacion,
                    "error": resultado.get("error", "Error desconocido en el procesamiento"),
                    "frases_clave": {},
                    "total_frases": 0
                }
                
        except Exception as e:
            logger.error(f"Error analizando documento {nombre_decodificado}: {e}")
            return {
                "nombre": nombre_decodificado,
                "procesado": False,
                "tamaño": tamaño,
                "fecha_procesamiento": fecha_modificacion,
                "error": f"Error en el análisis: {str(e)}",
                "frases_clave": {},
                "total_frases": 0
            }
            
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error obteniendo documento {nombre_archivo}: {e}")
        raise HTTPException(status_code=500, detail=f"Error interno: {str(e)}")


# Ruta de compatibilidad: redirige a la vista de archivo
@app.get("/documento/{nombre_archivo}")
async def redirigir_documento(nombre_archivo: str):
    return RedirectResponse(url=f"/archivo/{nombre_archivo}")



def analizar_sentencias_existentes() -> Dict[str, Any]:
    """Analiza las sentencias existentes en la carpeta con caché"""
    global CACHE_TIMESTAMP, ANALISIS_CACHE
    
    # Verificar si el caché es válido
    if (CACHE_TIMESTAMP and 
        (datetime.now() - CACHE_TIMESTAMP).total_seconds() < CACHE_DURATION and 
        ANALISIS_CACHE):
        logger.info("📋 Usando análisis desde caché")
        return ANALISIS_CACHE
    
    try:
        logger.info(f"🔍 Iniciando análisis de sentencias existentes en: {SENTENCIAS_DIR}")
        
        if not SENTENCIAS_DIR.exists():
            logger.warning(f"❌ La carpeta '{SENTENCIAS_DIR}' no existe")
            resultado = {
                "error": f"La carpeta '{SENTENCIAS_DIR}' no existe",
                "archivos_analizados": 0,
                "total_apariciones": 0,
                "resultados_por_archivo": {},
                "ranking_global": {}
            }
            ANALISIS_CACHE = resultado
            CACHE_TIMESTAMP = datetime.now()
            return resultado
        
        # Buscar archivos de texto y PDF en ambos directorios
        archivos_soportados = []
        
        # Buscar en directorio sentencias
        if SENTENCIAS_DIR.exists():
            archivos_soportados.extend([f for f in SENTENCIAS_DIR.iterdir() 
                                      if f.suffix.lower() in ['.txt', '.pdf']])
        
        # Buscar en directorio uploads
        if UPLOADS_DIR.exists():
            archivos_soportados.extend([f for f in UPLOADS_DIR.iterdir() 
                                      if f.suffix.lower() in ['.txt', '.pdf']])
        
        logger.info(f"📁 Archivos encontrados: {[f.name for f in archivos_soportados]}")
        
        if not archivos_soportados:
            logger.warning(f"❌ No se encontraron archivos .txt o .pdf en '{SENTENCIAS_DIR}' ni '{UPLOADS_DIR}'")
            return {
                "error": f"No se encontraron archivos .txt o .pdf en '{SENTENCIAS_DIR}' ni '{UPLOADS_DIR}'",
                "archivos_analizados": 0,
                "total_apariciones": 0,
                "resultados_por_archivo": {},
                "ranking_global": {}
            }
        
        resultados_por_archivo = {}
        ranking_global = {}
        total_apariciones = 0
        
        for archivo in archivos_soportados:
            try:
                logger.info(f"🔍 Analizando archivo: {archivo.name}")
                
                # Marcar tiempo de inicio para este archivo
                tiempo_inicio = datetime.now()
                
                # Usar el analizador de IA si está disponible, sino el básico
                if ANALIZADOR_IA_DISPONIBLE:
                    logger.info(f"🤖 Usando analizador de IA para: {archivo.name}")
                    from backend.analisis import AnalizadorLegal
                    analizador = AnalizadorLegal()
                    analizador._tiempo_inicio = tiempo_inicio
                    resultado = analizador.analizar_documento(str(archivo))
                else:
                    logger.info(f"🔧 Usando analizador básico para: {archivo.name}")
                    analizador_basico._tiempo_inicio = tiempo_inicio
                    resultado = analizador_basico.analizar_documento(str(archivo), archivo.name)
                
                logger.info(f"📊 Resultado para {archivo.name}: procesado={resultado.get('procesado')}")
                
                if resultado.get("procesado"):
                    # Calcular total de frases clave
                    frases_clave = resultado.get("frases_clave", {})
                    total_frases = sum(datos.get("total", 0) for datos in frases_clave.values())
                    
                    # Agregar campos adicionales al resultado
                    resultado["total_frases"] = total_frases
                    resultado["timestamp"] = tiempo_inicio.strftime("%Y-%m-%d %H:%M:%S")
                    resultado["categorias_encontradas"] = len(frases_clave)
                    
                    resultados_por_archivo[archivo.name] = resultado
                    
                    logger.info(f"🔑 Frases clave encontradas en {archivo.name}: {list(frases_clave.keys())}")
                    logger.info(f"📊 Total frases en {archivo.name}: {total_frases}")
                    
                    for categoria, datos in frases_clave.items():
                        if categoria in ranking_global:
                            ranking_global[categoria]["total"] += datos["total"]
                            ranking_global[categoria]["ocurrencias"].extend(datos["ocurrencias"])
                            logger.info(f"📈 Actualizando {categoria}: total={ranking_global[categoria]['total']}")
                        else:
                            ranking_global[categoria] = {
                                "total": datos["total"],
                                "ocurrencias": datos["ocurrencias"]
                            }
                            logger.info(f"🆕 Nueva categoría {categoria}: total={datos['total']}")
                        total_apariciones += datos["total"]
                else:
                    logger.warning(f"⚠️ Archivo {archivo.name} no se pudo procesar")
                    resultados_por_archivo[archivo.name] = {"error": "No se pudo procesar"}
                    
            except Exception as e:
                logger.error(f"❌ Error analizando {archivo}: {e}")
                resultados_por_archivo[archivo.name] = {"error": f"Error: {str(e)}"}
        
        # Ordenar ranking
        ranking_ordenado = dict(sorted(ranking_global.items(), key=lambda x: x[1]["total"], reverse=True))
        
        logger.info(f"📊 RESUMEN FINAL:")
        logger.info(f"  - Archivos analizados: {len(archivos_soportados)}")
        logger.info(f"  - Total apariciones: {total_apariciones}")
        logger.info(f"  - Categorías encontradas: {list(ranking_ordenado.keys())}")
        logger.info(f"  - Ranking global: {ranking_ordenado}")
        
        resultado = {
            "archivos_analizados": len(archivos_soportados),
            "total_apariciones": total_apariciones,
            "resultados_por_archivo": resultados_por_archivo,
            "ranking_global": ranking_ordenado
        }
        
        # Guardar en caché
        ANALISIS_CACHE = resultado
        CACHE_TIMESTAMP = datetime.now()
        logger.info("💾 Análisis guardado en caché")
        
        return resultado
        
    except Exception as e:
        logger.error(f"❌ Error analizando sentencias: {e}")
        return {
            "error": f"Error al analizar sentencias: {str(e)}",
            "archivos_analizados": 0,
            "total_apariciones": 0,
            "resultados_por_archivo": {},
            "ranking_global": {}
        }


# Middleware para asegurar Content-Type correcto para archivos estáticos
@app.middleware("http")
async def fix_content_type_middleware(request: Request, call_next):
    """Middleware para corregir Content-Type de archivos estáticos"""
    response = await call_next(request)
    
    # Solo aplicar a rutas de archivos estáticos, NO a rutas de la aplicación
    if request.url.path.startswith('/sentencias/') and request.url.path.endswith('.pdf'):
        response.headers["Content-Type"] = "application/pdf"
        response.headers["X-Content-Type-Options"] = "nosniff"
        response.headers["Cache-Control"] = "no-cache, no-store, must-revalidate"
        response.headers["Pragma"] = "no-cache"
        response.headers["Expires"] = "0"
        # Agregar header adicional para evitar interpretación como JavaScript
        response.headers["Content-Disposition"] = "inline"
    
    return response

# Ruta específica para interceptar DEMANDA.pdf directamente
@app.get("/DEMANDA.pdf")
async def servir_demanda_pdf():
    """Sirve específicamente el archivo DEMANDA.pdf con Content-Type correcto"""
    try:
        archivo_path = SENTENCIAS_DIR / "DEMANDA.pdf"
        
        if not archivo_path.exists():
            raise HTTPException(status_code=404, detail="Archivo DEMANDA.pdf no encontrado")
        
        # Headers específicos para evitar interpretación como JavaScript
        headers = {
            "Content-Type": "application/pdf",
            "X-Content-Type-Options": "nosniff",
            "Cache-Control": "no-cache, no-store, must-revalidate",
            "Pragma": "no-cache",
            "Expires": "0",
            "Content-Disposition": "inline; filename=DEMANDA.pdf"
        }
        
        return FileResponse(
            path=str(archivo_path),
            media_type="application/pdf",
            filename="DEMANDA.pdf",
            headers=headers
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error sirviendo DEMANDA.pdf: {e}")
        raise HTTPException(status_code=500, detail=f"Error interno: {str(e)}")

# Ruta específica para servir archivos desde sentencias/ con Content-Type correcto
@app.get("/sentencias/{nombre_archivo}")
async def servir_archivo_sentencias(nombre_archivo: str):
    """Sirve archivos desde el directorio sentencias/ con Content-Type correcto"""
    try:
        archivo_path = SENTENCIAS_DIR / nombre_archivo
        
        if not archivo_path.exists():
            raise HTTPException(status_code=404, detail="Archivo no encontrado")
        
        # Determinar el tipo MIME basado en la extensión
        if archivo_path.suffix.lower() == '.pdf':
            media_type = "application/pdf"
        elif archivo_path.suffix.lower() == '.txt':
            media_type = "text/plain; charset=utf-8"
        else:
            media_type = "application/octet-stream"
        
        # Configurar headers específicos
        headers = {
            "Content-Type": media_type,
            "X-Content-Type-Options": "nosniff",
            "Cache-Control": "no-cache, no-store, must-revalidate",
            "Pragma": "no-cache",
            "Expires": "0"
        }
        
        return FileResponse(
            path=str(archivo_path),
            media_type=media_type,
            filename=nombre_archivo,
            headers=headers
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error sirviendo archivo {nombre_archivo}: {e}")
        raise HTTPException(status_code=500, detail=f"Error interno: {str(e)}")

if __name__ == "__main__":
    import uvicorn
    
    def safe_print(text: str) -> None:
        try:
            print(text)
        except Exception:
            try:
                # Remove non-encodable characters for Windows consoles
                print(text.encode("cp1252", "ignore").decode("cp1252"))
            except Exception:
                try:
                    print(text.encode("utf-8", "ignore").decode("utf-8"))
                except Exception:
                    print("Inicio de servidor FastAPI")
    
    safe_print("🚀 Iniciando Analizador de Sentencias IPP/INSS...")
    safe_print(f"📁 Directorio de sentencias: {SENTENCIAS_DIR}")
    safe_print(f"🤖 IA disponible: {'✅ Sí' if ANALIZADOR_IA_DISPONIBLE else '❌ No'}")
    safe_print(f"🌐 URL: http://localhost:8000")
    safe_print(f"📚 Documentación: http://localhost:8000/docs")
    
    uvicorn.run(
        app, 
        host="0.0.0.0", 
        port=8000,
        log_level="info"
    )
