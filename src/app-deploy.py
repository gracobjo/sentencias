#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
Analizador de Sentencias IPP/INSS - AplicaciÃ³n FastAPI
AplicaciÃ³n robusta para anÃ¡lisis de documentos legales con modelo de IA pre-entrenado
"""

import os
import json
import re
import uuid
import shutil
import logging
from pathlib import Path
from threading import Lock
from typing import Optional, Dict, List, Any
from datetime import datetime
from io import BytesIO

from fastapi import FastAPI, Request, File, UploadFile, Form, HTTPException
from fastapi.templating import Jinja2Templates
from fastapi.staticfiles import StaticFiles
from fastapi.responses import HTMLResponse, JSONResponse, RedirectResponse, PlainTextResponse
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
# Imports de docx comentados para evitar problemas en despliegue
# from docx import Document
# from docx.shared import Inches, Pt
# from docx.enum.text import WD_ALIGN_PARAGRAPH
# from docx.enum.style import WD_STYLE_TYPE

# Configurar logging
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s'
)
logger = logging.getLogger(__name__)

# Cache para anÃ¡lisis
ANALISIS_CACHE = {}
CACHE_TIMESTAMP = None
CACHE_DURATION = 300  # 5 minutos en segundos

# Crear aplicaciÃ³n FastAPI con configuraciÃ³n robusta
app = FastAPI(
    title="Analizador de Sentencias IPP/INSS",
    description="API robusta para anÃ¡lisis inteligente de documentos legales",
    version="2.0.0",
    docs_url="/docs",
    redoc_url="/redoc"
)

# Configurar CORS para desarrollo
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)

# Configurar templates y archivos estÃ¡ticos
templates = Jinja2Templates(directory="src/templates")

# Configurar archivos estÃ¡ticos con tipos MIME correctos
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse
import mimetypes

# Configurar tipos MIME para archivos PDF
mimetypes.add_type('application/pdf', '.pdf')

# Montar archivos estÃ¡ticos con configuraciÃ³n personalizada
app.mount("/static", StaticFiles(directory="src/static"), name="static")

# Montar directorio de sentencias con configuraciÃ³n especÃ­fica para PDFs
class PDFStaticFiles(StaticFiles):
    def __init__(self, *args, **kwargs):
        super().__init__(*args, **kwargs)
    
    def get_response(self, path: str, scope):
        response = super().get_response(path, scope)
        if path.endswith('.pdf'):
            response.headers["Content-Type"] = "application/pdf"
            response.headers["X-Content-Type-Options"] = "nosniff"
            response.headers["Cache-Control"] = "no-cache, no-store, must-revalidate"
            response.headers["Pragma"] = "no-cache"
            response.headers["Expires"] = "0"
        return response

app.mount("/sentencias", PDFStaticFiles(directory="sentencias"), name="sentencias")

# ConfiguraciÃ³n de directorios
BASE_DIR = Path(__file__).parent.parent
SENTENCIAS_DIR = BASE_DIR / "sentencias"
UPLOADS_DIR = BASE_DIR / "uploads"
MODELS_DIR = BASE_DIR / "models"
LOGS_DIR = BASE_DIR / "logs"
FRASES_FILE = BASE_DIR / "models" / "frases_clave.json"
frases_lock = Lock()

# Crear directorios necesarios
for directory in [SENTENCIAS_DIR, UPLOADS_DIR, MODELS_DIR, LOGS_DIR]:
    directory.mkdir(exist_ok=True)

# ConfiguraciÃ³n de archivos permitidos
ALLOWED_EXTENSIONS = {'.pdf', '.txt', '.doc', '.docx'}
MAX_FILE_SIZE = 50 * 1024 * 1024  # 50MB

# Frases por defecto en caso de que no exista el archivo o sea invÃ¡lido
DEFAULT_FRASES_CLAVE: Dict[str, List[str]] = {
    "procedimiento_legal": ["procedente", "desestimamos", "estimamos"],
}


def load_frases_clave() -> Dict[str, List[str]]:
    """Carga el JSON de frases clave de disco con validaciÃ³n bÃ¡sica."""
    try:
        if not FRASES_FILE.exists():
            logger.warning(f"Archivo de frases no encontrado en {FRASES_FILE}")
            return DEFAULT_FRASES_CLAVE
        with open(FRASES_FILE, "r", encoding="utf-8") as f:
            data = json.load(f)
        if not isinstance(data, dict):
            logger.warning("Formato invÃ¡lido en frases_clave.json (no es dict)")
            return DEFAULT_FRASES_CLAVE
        # Normalizar valores a listas de strings Ãºnicas
        normalizado: Dict[str, List[str]] = {}
        for categoria, frases in data.items():
            if not isinstance(categoria, str):
                continue
            if isinstance(frases, list):
                solo_texto = [str(x).strip() for x in frases if str(x).strip()]
                # quitar duplicados preservando orden
                seen = set()
                unicos: List[str] = []
                for s in solo_texto:
                    if s.lower() not in seen:
                        seen.add(s.lower())
                        unicos.append(s)
                # Incluir tambiÃ©n categorÃ­as vacÃ­as (permitimos crear primero y rellenar despuÃ©s)
                normalizado[categoria] = unicos
        return normalizado or DEFAULT_FRASES_CLAVE
    except Exception as e:
        logger.error(f"Error cargando frases_clave.json: {e}")
        return DEFAULT_FRASES_CLAVE


def save_frases_clave(data: Dict[str, List[str]]) -> None:
    """Guarda el JSON de frases clave de forma atÃ³mica y segura."""
    if not isinstance(data, dict):
        raise ValueError("El payload debe ser un objeto {categoria: [frases]}")
    for categoria, frases in data.items():
        if not isinstance(categoria, str) or not isinstance(frases, list):
            raise ValueError("Formato invÃ¡lido: claves str, valores lista de str")
    MODELS_DIR.mkdir(exist_ok=True)
    temp_path = FRASES_FILE.with_suffix(".tmp")
    with frases_lock:
        with open(temp_path, "w", encoding="utf-8") as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        os.replace(temp_path, FRASES_FILE)

# Importar el analizador de IA (asumiendo que ya estÃ¡ entrenado)
try:
    logger.info("ðŸ” Intentando cargar mÃ³dulo de IA...")
    
    # Verificar que los archivos del modelo existan
    models_dir = Path("models")
    if not models_dir.exists():
        raise ImportError("Directorio 'models' no existe")
    
    modelo_file = models_dir / "modelo_legal.pkl"
    if not modelo_file.exists():
        raise ImportError(f"Archivo del modelo no existe: {modelo_file}")
    
    logger.info(f"âœ… Archivos del modelo encontrados: {list(models_dir.iterdir())}")
    
    # Intentar importar el mÃ³dulo
    from src.backend.analisis import AnalizadorLegal
    logger.info("âœ… MÃ³dulo backend.analisis importado")
    
    # Intentar crear una instancia para verificar que funciona
    analizador_test = AnalizadorLegal()
    logger.info("âœ… AnalizadorLegal creado exitosamente")
    
    ANALIZADOR_IA_DISPONIBLE = True
    logger.info("âœ… MÃ³dulo de IA cargado correctamente")
    
except Exception as e:
    ANALIZADOR_IA_DISPONIBLE = False
    logger.error(f"âŒ Error cargando mÃ³dulo de IA: {e}")
    logger.error(f"âŒ Tipo de error: {type(e).__name__}")
    import traceback
    logger.error(f"âŒ Traceback completo: {traceback.format_exc()}")
    logger.info("Se usarÃ¡ anÃ¡lisis bÃ¡sico como fallback")
    
    # Intentar cargar al menos el analizador bÃ¡sico
    try:
        from src.backend.analisis_basico import AnalizadorBasico
        analizador_basico = AnalizadorBasico()
        logger.info("âœ… Analizador bÃ¡sico cargado como fallback")
    except Exception as e2:
        logger.error(f"âŒ Error cargando analizador bÃ¡sico: {e2}")
        analizador_basico = None


def extraer_texto_pdf(ruta: str) -> str:
    """Lee archivos PDF y extrae el texto"""
    try:
        # Intentar importar PyPDF2
        try:
            import PyPDF2
        except ImportError:
            logger.warning("PyPDF2 no estÃ¡ instalado. Instala: pip install PyPDF2")
            return "Error: PyPDF2 no estÃ¡ instalado. Ejecuta: pip install PyPDF2"
        
        texto = ""
        with open(ruta, 'rb') as archivo:
            lector = PyPDF2.PdfReader(archivo)
            
            for pagina in lector.pages:
                texto += pagina.extract_text() + "\n"
            
            return texto.strip()
            
    except Exception as e:
        logger.error(f"Error leyendo PDF {ruta}: {e}")
        return ""

def generar_analisis_discrepancias_basico(ruta_archivo: str, resultado_base: dict) -> dict:
    """Generar anÃ¡lisis bÃ¡sico de discrepancias mÃ©dicas-legales"""
    try:
        # Leer contenido del archivo
        contenido = ""
        if ruta_archivo.endswith('.pdf'):
            contenido = extraer_texto_pdf(ruta_archivo)
        else:
            with open(ruta_archivo, 'r', encoding='utf-8') as f:
                contenido = f.read()
        
        if not contenido or len(contenido.strip()) == 0:
            logger.warning(f"Contenido vacÃ­o o no se pudo leer: {ruta_archivo}")
            return {
                "discrepancias_detectadas": [],
                "evidencia_favorable": [],
                "argumentos_juridicos": [],
                "recomendaciones_defensa": [],
                "contradicciones_internas": [],
                "puntuacion_discrepancia": 0,
                "probabilidad_ipp": 0.0,
                "resumen_ejecutivo": "No se pudo leer el contenido del archivo"
            }
        
        # AnÃ¡lisis bÃ¡sico de discrepancias
        discrepancias_detectadas = []
        evidencia_favorable = []
        argumentos_juridicos = []
        recomendaciones_defensa = []
        contradicciones_internas = []
        
        # Detectar patrones bÃ¡sicos de discrepancias
        contenido_lower = contenido.lower()
        
        # Patrones de evidencia favorable para IPP
        patrones_evidencia = [
            "rotura completa",
            "rotura de espesor completo",
            "cirugÃ­a reconstructiva", 
            "limitaciÃ³n activa",
            "fuerza insuficiente",
            "atrofia muscular",
            "discinesia escapular",
            "manguito rotador",
            "supraespinoso",
            "artropatÃ­a",
            "retracciÃ³n fibrilar",
            "tenopatÃ­a severa",
            "anclajes",
            "tornillos corkscrew",
            "flexiÃ³n activa",
            "abducciÃ³n activa",
            "balance muscular",
            "fuerza de garra",
            "prÃ¡cticamente nulo desarrollo de fuerza",
            "insuficiente para vencer la gravedad"
        ]
        
        # Detectar evidencia favorable
        for patron in patrones_evidencia:
            if patron in contenido_lower:
                evidencia_favorable.append({
                    "tipo": "evidencia_estructural",
                    "descripcion": f"Evidencia encontrada: {patron}",
                    "relevancia": "ALTA",
                    "argumento": f"PatrÃ³n '{patron}' sugiere gravedad de la lesiÃ³n que excede LPNI"
                })
        
        # Detectar discrepancias especÃ­ficas
        if "lesiones permanentes no incapacitantes" in contenido_lower or "lpni" in contenido_lower:
            if evidencia_favorable:
                discrepancias_detectadas.append({
                    "tipo": "clasificacion_inadecuada",
                    "descripcion": "ClasificaciÃ³n como LPNI incompatible con evidencia de gravedad",
                    "severidad": "ALTA",
                    "argumento_juridico": "La evidencia objetiva sugiere limitaciÃ³n funcional permanente superior al 33%"
                })
        
        # Detectar contradicciones internas
        if "no presenta limitaciÃ³n importante" in contenido_lower and ("limitaciÃ³n activa" in contenido_lower or "fuerza insuficiente" in contenido_lower):
            contradicciones_internas.append({
                "tipo": "contradiccion_interna",
                "descripcion": "ContradicciÃ³n entre conclusiÃ³n y hallazgos objetivos",
                "severidad": "MEDIA"
            })
        
        # Generar argumentos jurÃ­dicos especÃ­ficos
        if evidencia_favorable:
            argumentos_juridicos.append({
                "titulo": "Fundamento Legal - Art. 194.2 LGSS",
                "contenido": "DisminuciÃ³n â‰¥33% en el rendimiento normal de la profesiÃ³n habitual",
                "fuerza": "ALTA"
            })
            argumentos_juridicos.append({
                "titulo": "Evidencia Estructural",
                "contenido": "Lesiones anatÃ³micas permanentes que impiden recuperaciÃ³n funcional completa",
                "fuerza": "ALTA"
            })
            argumentos_juridicos.append({
                "titulo": "LimitaciÃ³n Funcional Objetiva",
                "contenido": "Diferencia entre movilidad pasiva y activa indica incapacidad laboral",
                "fuerza": "MEDIA"
            })
        
        # Generar recomendaciones especÃ­ficas
        recomendaciones_defensa.append({
            "titulo": "Estrategia de Defensa Principal",
            "contenido": "Enfocar la defensa en la evidencia objetiva y contradicciones del informe",
            "prioridad": "ALTA",
            "acciones": [
                "Preparar argumentos basados en el Art. 194.2 LGSS",
                "Documentar todas las limitaciones funcionales objetivas",
                "Presentar evidencia de duraciÃ³n prolongada del proceso",
                "Destacar las contradicciones internas del informe mÃ©dico",
                "Solicitar peritaje biomecÃ¡nico complementario"
            ]
        })
        
        if contradicciones_internas:
            recomendaciones_defensa.append({
                "titulo": "Explotar Contradicciones",
                "contenido": "Las contradicciones internas del informe debilitan la conclusiÃ³n de LPNI",
                "prioridad": "MEDIA",
                "acciones": [
                    "SeÃ±alar la discrepancia entre hallazgos y conclusiÃ³n",
                    "Cuestionar la validez de la evaluaciÃ³n mÃ©dica",
                    "Solicitar revisiÃ³n por perito independiente"
                ]
            })
        
        # Calcular mÃ©tricas mejoradas
        puntuacion_base = len(evidencia_favorable) * 5
        puntuacion_discrepancias = len(discrepancias_detectadas) * 15
        puntuacion_contradicciones = len(contradicciones_internas) * 10
        
        puntuacion_discrepancia = min(100, puntuacion_base + puntuacion_discrepancias + puntuacion_contradicciones)
        
        # Calcular probabilidad IPP basada en evidencia especÃ­fica
        probabilidad_base = len(evidencia_favorable) * 0.08
        if len(discrepancias_detectadas) > 0:
            probabilidad_base += 0.3
        if len(contradicciones_internas) > 0:
            probabilidad_base += 0.2
            
        probabilidad_ipp = min(0.95, probabilidad_base)
        
        # Generar resumen ejecutivo mejorado
        conclusion = "ALTA PROBABILIDAD DE IPP" if probabilidad_ipp > 0.6 else "MEDIA PROBABILIDAD DE IPP" if probabilidad_ipp > 0.3 else "BAJA PROBABILIDAD DE IPP"
        icono = "âœ…" if probabilidad_ipp > 0.6 else "âš ï¸" if probabilidad_ipp > 0.3 else "âŒ"
        
        resumen_ejecutivo = f"""
ANÃLISIS DE DISCREPANCIAS MÃ‰DICAS-LEGALES

ðŸ“Š RESUMEN EJECUTIVO:
â€¢ Discrepancias detectadas: {len(discrepancias_detectadas)}
â€¢ Evidencia favorable: {len(evidencia_favorable)} elementos
â€¢ Contradicciones internas: {len(contradicciones_internas)}
â€¢ PuntuaciÃ³n de discrepancia: {puntuacion_discrepancia}/100
â€¢ Probabilidad de IPP: {probabilidad_ipp*100:.1f}%

{icono} CONCLUSIÃ“N: {conclusion}
{'La evidencia disponible respalda firmemente la calificaciÃ³n de IPP.' if probabilidad_ipp > 0.6 else 'La evidencia sugiere una posible calificaciÃ³n de IPP que merece revisiÃ³n.' if probabilidad_ipp > 0.3 else 'La evidencia disponible no respalda claramente la calificaciÃ³n de IPP.'}

ðŸ” PUNTOS CLAVE PARA LA DEFENSA:
â€¢ Evidencia estructural de gravedad que excede LPNI
â€¢ Limitaciones funcionales objetivas documentadas
â€¢ Contradicciones internas en el informe mÃ©dico
â€¢ Fundamentos legales sÃ³lidos (Art. 194.2 LGSS)
"""
        
        return {
            "discrepancias_detectadas": discrepancias_detectadas,
            "evidencia_favorable": evidencia_favorable,
            "argumentos_juridicos": argumentos_juridicos,
            "recomendaciones_defensa": recomendaciones_defensa,
            "contradicciones_internas": contradicciones_internas,
            "puntuacion_discrepancia": puntuacion_discrepancia,
            "probabilidad_ipp": probabilidad_ipp,
            "resumen_ejecutivo": resumen_ejecutivo.strip()
        }
        
    except Exception as e:
        logger.error(f"Error generando anÃ¡lisis de discrepancias: {e}")
        return {
            "discrepancias_detectadas": [],
            "evidencia_favorable": [],
            "argumentos_juridicos": [],
            "recomendaciones_defensa": [],
            "contradicciones_internas": [],
            "puntuacion_discrepancia": 0,
            "probabilidad_ipp": 0.0,
            "resumen_ejecutivo": "Error en el anÃ¡lisis de discrepancias"
        }


class AnalizadorBasico:
    """Analizador bÃ¡sico como fallback cuando no hay IA disponible"""
    
    def __init__(self):
        self.frases_clave = {}
        self.cargar_frases_desde_modelo()

    def cargar_frases_desde_modelo(self) -> None:
        """Carga frases clave desde models/frases_clave.json o usa un conjunto por defecto."""
        try:
            datos = load_frases_clave()
            if isinstance(datos, dict) and datos:
                self.frases_clave = datos
            else:
                self.frases_clave = DEFAULT_FRASES_CLAVE
                logger.warning("Usando frases por defecto; archivo vacÃ­o o invÃ¡lido")
        except Exception as e:
            logger.warning(f"No se pudo cargar frases_clave.json: {e}. Usando valores por defecto")
            self.frases_clave = DEFAULT_FRASES_CLAVE
    
    def analizar_documento(self, ruta_archivo: str, nombre_original: str = "") -> Dict[str, Any]:
        """Analiza un documento usando mÃ©todos bÃ¡sicos"""
        try:
            # Leer contenido del archivo
            contenido = self._leer_archivo(ruta_archivo)
            if not contenido:
                return self._crear_resultado_error("No se pudo leer el contenido del archivo")
            
            # AnÃ¡lisis bÃ¡sico
            frases_encontradas = self._contar_frases_clave(contenido, nombre_original)
            prediccion = self._prediccion_basica(contenido)
            argumentos = self._extraer_argumentos(contenido)
            insights = self._generar_insights(prediccion, frases_encontradas)
            
            # Calcular tiempo de procesamiento
            tiempo_inicio = getattr(self, '_tiempo_inicio', None)
            tiempo_procesamiento = None
            if tiempo_inicio:
                tiempo_procesamiento = f"{(datetime.now() - tiempo_inicio).total_seconds():.2f}s"
            
            return {
                "archivo": ruta_archivo,
                "nombre_archivo": nombre_original or Path(ruta_archivo).name,
                "procesado": True,
                "texto_extraido": contenido[:1000] + "..." if len(contenido) > 1000 else contenido,
                "longitud_texto": len(contenido),
                "prediccion": prediccion,
                "argumentos": argumentos,
                "frases_clave": frases_encontradas,
                "resumen_inteligente": self._generar_resumen(prediccion, frases_encontradas),
                "insights_juridicos": insights,
                "archivo_id": Path(ruta_archivo).name,
                "tipo_documento": "sentencia",
                "extract_entities": False,
                "analyze_arguments": False,
                "modelo_ia": False,
                "total_frases_clave": sum(datos["total"] for datos in frases_encontradas.values()),
                "timestamp": datetime.now().isoformat(),
                "tiempo_procesamiento": tiempo_procesamiento or "N/A",
                "ruta_archivo": ruta_archivo
            }
            
        except Exception as e:
            logger.error(f"Error en anÃ¡lisis bÃ¡sico: {e}")
            return self._crear_resultado_error(f"Error en anÃ¡lisis: {str(e)}")
    
    def _leer_archivo(self, ruta: str) -> Optional[str]:
        """Lee el contenido de un archivo con manejo de errores"""
        try:
            if ruta.endswith('.txt'):
                # Intentar diferentes encodings
                for encoding in ['utf-8', 'latin-1', 'cp1252']:
                    try:
                        with open(ruta, 'r', encoding=encoding) as f:
                            return f.read().strip()
                    except UnicodeDecodeError:
                        continue
                return None
            else:
                # Para otros formatos, usar funciÃ³n de lectura genÃ©rica
                return self._leer_archivo_generico(ruta)
        except Exception as e:
            logger.error(f"Error leyendo archivo {ruta}: {e}")
            return None
    
    def _leer_archivo_generico(self, ruta: str) -> Optional[str]:
        """Lee archivos de diferentes formatos"""
        try:
            # Por ahora solo manejamos texto, pero aquÃ­ se podrÃ­a extender
            # para PDF, DOC, etc. usando librerÃ­as como PyPDF2, python-docx
            return "Contenido del archivo no disponible en formato de texto"
        except Exception as e:
            logger.error(f"Error leyendo archivo genÃ©rico {ruta}: {e}")
            return None
    
    def _contar_frases_clave(self, texto: str, nombre_archivo: str) -> Dict[str, Any]:
        """Cuenta las ocurrencias de frases clave"""
        if not texto:
            return {}
        
        resultados = {}
        for categoria, variantes in self.frases_clave.items():
            total = 0
            ocurrencias = []
            
            for variante in variantes:
                # Permitir espacios/guiones/underscores intercambiables entre palabras
                flexible = re.escape(variante)
                flexible = flexible.replace("\\ ", "\\s+")
                flexible = flexible.replace("\\_", "[\\s_\\-_]+")
                flexible = flexible.replace("\\-", "[\\s_\\-_]+")
                patron = re.compile(flexible, re.IGNORECASE)
                matches = patron.finditer(texto)
                
                for match in matches:
                    total += 1
                    start_pos = match.start()
                    end_pos = match.end()
                    
                    # Obtener contexto (100 caracteres antes y despuÃ©s)
                    context_start = max(0, start_pos - 100)
                    context_end = min(len(texto), end_pos + 100)
                    contexto = texto[context_start:context_end]
                    
                    # Marcar la frase encontrada
                    frase_encontrada = texto[start_pos:end_pos]
                    contexto_marcado = contexto.replace(frase_encontrada, f"**{frase_encontrada}**")
                    
                    ocurrencias.append({
                        "frase": variante,
                        "posicion": start_pos,
                        "contexto": contexto_marcado,
                        "linea": texto[:start_pos].count('\n') + 1,
                        "archivo": nombre_archivo
                    })
            
            if total > 0:
                # Obtener frases Ãºnicas encontradas
                frases_encontradas = list(set([oc["frase"] for oc in ocurrencias]))
                
                resultados[categoria] = {
                    "total": total,
                    "ocurrencias": ocurrencias,
                    "frases": frases_encontradas
                }
        
        return resultados
    
    def _prediccion_basica(self, texto: str) -> Dict[str, Any]:
        """PredicciÃ³n bÃ¡sica basada en palabras clave"""
        texto_lower = texto.lower()
        
        palabras_positivas = [
            "estimamos", "estimamos procedente", "procedente", "favorable",
            "accedemos", "concedemos", "reconocemos", "declaramos procedente",
            "estimamos fundada", "fundada", "estimamos parcialmente"
        ]
        
        palabras_negativas = [
            "desestimamos", "desestimamos la reclamaciÃ³n", "desfavorable",
            "no procedente", "rechazamos", "denegamos", "estimamos infundada",
            "infundada", "desestimamos parcialmente"
        ]
        
        positivas = sum(1 for palabra in palabras_positivas if palabra in texto_lower)
        negativas = sum(1 for palabra in palabras_negativas if palabra in texto_lower)
        
        total = positivas + negativas
        if total == 0:
            return {
                "es_favorable": True,
                "confianza": 0.5,
                "interpretacion": "Neutral - No hay indicadores claros"
            }
        
        confianza = min(0.95, max(0.1, (positivas + negativas) / (total + 1)))
        es_favorable = positivas > negativas
        
        if es_favorable:
            interpretacion = f"Favorable - {positivas} indicadores positivos encontrados"
        else:
            interpretacion = f"Desfavorable - {negativas} indicadores negativos encontrados"
        
        return {
            "es_favorable": es_favorable,
            "confianza": confianza,
            "interpretacion": interpretacion
        }
    
    def _extraer_argumentos(self, texto: str) -> List[Dict[str, Any]]:
        """Extrae argumentos legales del texto"""
        argumentos = []
        
        patrones = [
            r"por\s+(?:lo\s+)?que\s+([^.]*?\.)",
            r"fundamentos?\s+(?:de\s+)?(?:derecho|derecho\s+por\s+lo\s+que)\s+([^.]*?\.)",
            r"considerando\s+que\s+([^.]*?\.)",
            r"vistos\s+([^.]*?\.)",
            r"resultando\s+([^.]*?\.)"
        ]
        
        for patron in patrones:
            matches = re.finditer(patron, texto, re.IGNORECASE)
            for match in matches:
                argumento = match.group(1).strip()
                if len(argumento) > 20:
                    argumentos.append({
                        "tipo": "argumento_legal",
                        "texto": argumento,
                        "posicion": match.start(),
                        "confianza": 0.7,
                        "categoria": "fundamento_juridico"
                    })
        
        return argumentos
    
    def _generar_insights(self, prediccion: Dict, frases_clave: Dict) -> List[str]:
        """Genera insights jurÃ­dicos bÃ¡sicos"""
        insights = []
        
        if prediccion["es_favorable"]:
            insights.append("El documento presenta argumentos sÃ³lidos a favor del caso.")
            insights.append("La resoluciÃ³n favorece al reclamante.")
        else:
            insights.append("El documento presenta argumentos que pueden ser desfavorables.")
            insights.append("La resoluciÃ³n no favorece al reclamante.")
        
        if frases_clave:
            categorias = list(frases_clave.keys())
            insights.append(f"Se identificaron {len(categorias)} categorÃ­as de frases clave.")
            insights.append(f"Las categorÃ­as mÃ¡s relevantes son: {', '.join(categorias[:3])}.")
        
        insights.append(f"Confianza del anÃ¡lisis: {prediccion['confianza']:.1%}")
        
        return insights
    
    def _generar_resumen(self, prediccion: Dict, frases_clave: Dict) -> str:
        """Genera un resumen inteligente del anÃ¡lisis"""
        if prediccion["es_favorable"]:
            base = "AnÃ¡lisis favorable del documento legal. "
        else:
            base = "AnÃ¡lisis desfavorable del documento legal. "
        
        if frases_clave:
            total_frases = sum(datos["total"] for datos in frases_clave.values())
            base += f"Se identificaron {total_frases} frases clave relevantes. "
        
        base += f"Confianza del anÃ¡lisis: {prediccion['confianza']:.1%}."
        
        return base
    
    def _crear_resultado_error(self, mensaje: str) -> Dict[str, Any]:
        """Crea un resultado de error estandarizado"""
        return {
            "error": mensaje,
            "procesado": False,
            "nombre_archivo": "archivo_desconocido",
            "prediccion": {"es_favorable": False, "confianza": 0.0},
            "frases_clave": {},
            "argumentos": [],
            "insights_juridicos": [f"Error: {mensaje}"],
            "longitud_texto": 0,
            "total_frases_clave": 0
        }


# Instanciar analizador bÃ¡sico
analizador_basico = AnalizadorBasico()


# ====== MODELOS Pydantic para CRUD ======
class FrasesPayload(BaseModel):
    categorias: Dict[str, List[str]]

class CategoriaPayload(BaseModel):
    nombre: str
    frases: List[str] = []

class FrasePayload(BaseModel):
    categoria: str
    frase: str

class RenameCategoryPayload(BaseModel):
    old_name: str
    new_name: str

class UpdatePhrasePayload(BaseModel):
    categoria: str
    old_frase: str
    new_frase: str


class DeleteDocumentPayload(BaseModel):
    nombre_archivo: str


# ====== ENDPOINTS CRUD DE FRASES CLAVE ======
@app.get("/api/frases")
async def listar_frases():
    """Lista todas las categorÃ­as y frases clave actuales."""
    datos = load_frases_clave()
    return {"categorias": datos}


@app.post("/api/frases")
async def reemplazar_frases(payload: FrasesPayload):
    """Reemplaza completamente el set de frases clave."""
    save_frases_clave(payload.categorias)
    analizador_basico.cargar_frases_desde_modelo()
    return {"status": "ok"}


@app.post("/api/frases/categoria")
async def crear_categoria(payload: CategoriaPayload):
    datos = load_frases_clave()
    nombre = payload.nombre.strip()
    if not nombre:
        raise HTTPException(status_code=400, detail="Nombre de categorÃ­a requerido")
    if nombre in datos:
        raise HTTPException(status_code=409, detail="La categorÃ­a ya existe")
    datos[nombre] = [s.strip() for s in payload.frases if s and s.strip()]
    save_frases_clave(datos)
    analizador_basico.cargar_frases_desde_modelo()
    return {"status": "ok", "categoria": nombre}


@app.delete("/api/frases/categoria/{nombre}")
async def eliminar_categoria(nombre: str):
    datos = load_frases_clave()
    if nombre not in datos:
        raise HTTPException(status_code=404, detail="CategorÃ­a no encontrada")
    datos.pop(nombre)
    save_frases_clave(datos)
    analizador_basico.cargar_frases_desde_modelo()
    return {"status": "ok"}


@app.patch("/api/frases/categoria")
async def renombrar_categoria(payload: RenameCategoryPayload):
    datos = load_frases_clave()
    old_name = payload.old_name.strip()
    new_name = payload.new_name.strip()
    if not old_name or not new_name:
        raise HTTPException(status_code=400, detail="Nombres requeridos")
    if old_name not in datos:
        raise HTTPException(status_code=404, detail="CategorÃ­a original no encontrada")
    if new_name in datos and new_name != old_name:
        raise HTTPException(status_code=409, detail="La categorÃ­a destino ya existe")
    if new_name == old_name:
        return {"status": "ok", "categoria": new_name}
    datos[new_name] = datos.pop(old_name)
    save_frases_clave(datos)
    analizador_basico.cargar_frases_desde_modelo()
    return {"status": "ok", "categoria": new_name}


@app.post("/api/frases/frase")
async def agregar_frase(payload: FrasePayload):
    datos = load_frases_clave()
    categoria = payload.categoria
    frase = payload.frase.strip()
    if not frase:
        raise HTTPException(status_code=400, detail="Frase requerida")
    if categoria not in datos:
        datos[categoria] = []
    # evitar duplicados case-insensitive
    if frase.lower() not in [f.lower() for f in datos[categoria]]:
        datos[categoria].append(frase)
    save_frases_clave(datos)
    analizador_basico.cargar_frases_desde_modelo()
    return {"status": "ok"}


@app.delete("/api/frases/frase")
async def eliminar_frase(payload: FrasePayload):
    datos = load_frases_clave()
    categoria = payload.categoria
    frase = payload.frase.strip()
    if categoria not in datos:
        raise HTTPException(status_code=404, detail="CategorÃ­a no encontrada")
    datos[categoria] = [f for f in datos[categoria] if f.lower() != frase.lower()]
    save_frases_clave(datos)
    analizador_basico.cargar_frases_desde_modelo()
    return {"status": "ok"}


@app.patch("/api/frases/frase")
async def actualizar_frase(payload: UpdatePhrasePayload):
    datos = load_frases_clave()
    categoria = payload.categoria
    old_frase = payload.old_frase.strip()
    new_frase = payload.new_frase.strip()
    if not new_frase:
        raise HTTPException(status_code=400, detail="Nueva frase requerida")
    if categoria not in datos:
        raise HTTPException(status_code=404, detail="CategorÃ­a no encontrada")
    frases = datos[categoria]
    indices = [i for i, f in enumerate(frases) if f.lower() == old_frase.lower()]
    if not indices:
        raise HTTPException(status_code=404, detail="Frase original no encontrada")
    # Evitar duplicado de destino
    if any(f.lower() == new_frase.lower() for f in frases):
        # si ya existe exacta, eliminar la vieja
        frases = [f for f in frases if f.lower() != old_frase.lower()]
    else:
        # reemplazar el primero y eliminar duplicados del resto
        idx = indices[0]
        frases[idx] = new_frase
        # limpiar duplicados case-insensitive preservando orden
        seen = set()
        dedup: List[str] = []
        for f in frases:
            key = f.lower()
            if key not in seen:
                seen.add(key)
                dedup.append(f)
        frases = dedup
    datos[categoria] = frases
    save_frases_clave(datos)
    analizador_basico.cargar_frases_desde_modelo()
    return {"status": "ok"}


def validar_archivo(archivo: UploadFile) -> Dict[str, Any]:
    """Valida un archivo subido"""
    errores = []
    
    # Validar nombre
    if not archivo.filename:
        errores.append("No se seleccionÃ³ ningÃºn archivo")
    
    # Validar extensiÃ³n
    if archivo.filename:
        extension = Path(archivo.filename).suffix.lower()
        if extension not in ALLOWED_EXTENSIONS:
            errores.append(f"ExtensiÃ³n no permitida: {extension}. Permitidas: {', '.join(ALLOWED_EXTENSIONS)}")
    
    # Validar tamaÃ±o
    if archivo.size and archivo.size > MAX_FILE_SIZE:
        errores.append(f"Archivo demasiado grande: {archivo.size / (1024*1024):.1f}MB. MÃ¡ximo: {MAX_FILE_SIZE / (1024*1024)}MB")
    
    return {
        "valido": len(errores) == 0,
        "errores": errores
    }


@app.get("/", response_class=HTMLResponse)
async def pagina_principal(request: Request):
    """PÃ¡gina principal de la aplicaciÃ³n"""
    try:
        # Analizar sentencias existentes
        resultado = analizar_sentencias_existentes()
        
        # DEBUG: Log del resultado
        logger.info(f"Resultado de analizar_sentencias_existentes: {resultado}")
        logger.info(f"ranking_global keys: {list(resultado.get('ranking_global', {}).keys())}")
        logger.info(f"ranking_global length: {len(resultado.get('ranking_global', {}))}")
        
        return templates.TemplateResponse("index.html", {
            "request": request,
            "archivos_analizados": resultado.get("archivos_analizados", 0),
            "total_apariciones": resultado.get("total_apariciones", 0),
            "resultados_por_archivo": resultado.get("resultados_por_archivo", {}),
            "ranking_global": resultado.get("ranking_global", {}),
            "ia_disponible": ANALIZADOR_IA_DISPONIBLE
        })
    except Exception as e:
        logger.error(f"Error en pÃ¡gina principal: {e}")
        return templates.TemplateResponse("index.html", {
            "request": request,
            "error": f"Error al cargar la pÃ¡gina: {str(e)}",
            "ia_disponible": ANALIZADOR_IA_DISPONIBLE
        })


@app.get("/subir", response_class=HTMLResponse)
async def formulario_subida(request: Request):
    """Formulario para subir documentos"""
    return templates.TemplateResponse("subir.html", {
        "request": request,
        "ia_disponible": ANALIZADOR_IA_DISPONIBLE
    })


@app.post("/upload")
async def subir_documento(
    file: UploadFile = File(...),
    document_type: str = Form("sentencia"),
    extract_entities: bool = Form(True),
    analyze_arguments: bool = Form(True)
):
    """Procesa la subida de un documento"""
    try:
        # Validar archivo
        validacion = validar_archivo(file)
        if not validacion["valido"]:
            raise HTTPException(status_code=400, detail="; ".join(validacion["errores"]))
        
        # Generar nombre Ãºnico
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        unique_id = str(uuid.uuid4())[:8]
        extension = Path(file.filename).suffix
        nuevo_nombre = f"{document_type}_{timestamp}_{unique_id}{extension}"
        
        # Guardar archivo
        ruta_archivo = UPLOADS_DIR / nuevo_nombre
        with open(ruta_archivo, "wb") as buffer:
            shutil.copyfileobj(file.file, buffer)
        
        logger.info(f"Archivo subido: {nuevo_nombre}")
        
        # Analizar documento
        if ANALIZADOR_IA_DISPONIBLE:
            try:
                analizador = AnalizadorLegal()
                resultado = analizador.analizar_documento(str(ruta_archivo))
                resultado["modelo_ia"] = True
                logger.info("AnÃ¡lisis con IA completado")
            except Exception as e:
                logger.warning(f"Fallback a anÃ¡lisis bÃ¡sico: {e}")
                resultado = analizador_basico.analizar_documento(str(ruta_archivo), file.filename)
                resultado["modelo_ia"] = False
        else:
            resultado = analizador_basico.analizar_documento(str(ruta_archivo), file.filename)
            resultado["modelo_ia"] = False
        
        # Agregar metadatos
        resultado.update({
            "nombre_archivo": file.filename,
            "archivo_id": nuevo_nombre,
            "tipo_documento": document_type,
            "extract_entities": extract_entities,
            "analyze_arguments": analyze_arguments,
            "timestamp": timestamp,
            "ruta_archivo": str(ruta_archivo)
        })
        
        # Mover a carpeta de sentencias si es apropiado
        logger.info(f"ðŸ“ Tipo de documento: {document_type}")
        logger.info(f"ðŸ“ Archivo guardado en: {ruta_archivo}")
        
        if document_type == "sentencia":
            destino = SENTENCIAS_DIR / nuevo_nombre
            logger.info(f"ðŸ“ Moviendo archivo a: {destino}")
            shutil.move(str(ruta_archivo), str(destino))
            resultado["ruta_archivo"] = str(destino)
            logger.info(f"âœ… Archivo movido exitosamente a sentencias/")
        else:
            logger.info(f"ðŸ“ Archivo permanece en uploads/ (tipo: {document_type})")
        
        # Log del resultado final
        logger.info(f"ðŸ“‹ RESULTADO FINAL:")
        logger.info(f"  - Archivo ID: {resultado.get('archivo_id')}")
        logger.info(f"  - Nombre original: {resultado.get('nombre_archivo')}")
        logger.info(f"  - Tipo documento: {resultado.get('tipo_documento')}")
        logger.info(f"  - Ruta final: {resultado.get('ruta_archivo')}")
        
        return JSONResponse(content=resultado)
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error procesando archivo: {e}")
        raise HTTPException(status_code=500, detail=f"Error interno: {str(e)}")


@app.get("/resultado/{archivo_id}", response_class=HTMLResponse)
async def mostrar_resultados(request: Request, archivo_id: str):
    """Muestra los resultados del anÃ¡lisis"""
    try:
        # Buscar archivo
        ruta_archivo = UPLOADS_DIR / archivo_id
        if not ruta_archivo.exists():
            ruta_archivo = SENTENCIAS_DIR / archivo_id
        
        if not ruta_archivo.exists():
            raise HTTPException(status_code=404, detail="Archivo no encontrado")
        
        # Analizar si es necesario
        if archivo_id.startswith(("sentencia_", "demanda_", "informe_")):
            if ANALIZADOR_IA_DISPONIBLE:
                try:
                    analizador = AnalizadorLegal()
                    resultado = analizador.analizar_documento(str(ruta_archivo))
                    resultado["modelo_ia"] = True
                except Exception as e:
                    logger.warning(f"Fallback a anÃ¡lisis bÃ¡sico: {e}")
                    resultado = analizador_basico.analizar_documento(str(ruta_archivo), archivo_id)
                    resultado["modelo_ia"] = False
            else:
                resultado = analizador_basico.analizar_documento(str(ruta_archivo), archivo_id)
                resultado["modelo_ia"] = False
        else:
            # Usar anÃ¡lisis existente
            resultado = {
                "nombre_archivo": archivo_id,
                "procesado": True,
                "prediccion": {"es_favorable": True, "confianza": 0.8},
                "resumen_inteligente": "AnÃ¡lisis del documento existente.",
                "argumentos": [],
                "frases_clave": {},
                "insights_juridicos": ["Documento analizado previamente."],
                "longitud_texto": 0,
                "total_frases_clave": 0,
                "modelo_ia": False
            }
        
        return templates.TemplateResponse("resultado.html", {
            "request": request,
            "resultado": resultado
        })
        
    except Exception as e:
        logger.error(f"Error mostrando resultados: {e}")
        raise HTTPException(status_code=500, detail=f"Error: {str(e)}")


@app.get("/ver/{archivo_id}", response_class=HTMLResponse)
async def ver_archivo(request: Request, archivo_id: str, highlight: str = None, pos: int = None, index: int = None):
    """Muestra el contenido completo de un archivo con frases clave resaltadas y opcionalmente resalta una apariciÃ³n especÃ­fica"""
    try:
        logger.info(f"ðŸ” Procesando solicitud para archivo: {archivo_id}")
        logger.info(f"ðŸ“‹ ParÃ¡metros: highlight={highlight}, pos={pos}, index={index}")
        
        # Buscar el archivo en ambos directorios
        candidatos = [SENTENCIAS_DIR / archivo_id, UPLOADS_DIR / archivo_id]
        archivo_path = None
        
        for candidato in candidatos:
            if candidato.exists():
                archivo_path = candidato
                break
        
        if not archivo_path:
            logger.error(f"âŒ Archivo no encontrado en ningÃºn directorio: {archivo_id}")
            raise HTTPException(status_code=404, detail=f"Archivo '{archivo_id}' no encontrado")
        
        logger.info(f"âœ… Archivo encontrado: {archivo_path}")
        
        # Analizar el archivo para obtener frases clave
        if ANALIZADOR_IA_DISPONIBLE:
            try:
                analizador = AnalizadorLegal()
                resultado = analizador.analizar_documento(str(archivo_path))
            except Exception as e:
                logger.warning(f"Fallback a anÃ¡lisis bÃ¡sico: {e}")
                resultado = analizador_basico.analizar_documento(str(archivo_path), archivo_id)
        else:
            resultado = analizador_basico.analizar_documento(str(archivo_path), archivo_id)
        
        if not resultado.get("procesado"):
            raise HTTPException(status_code=500, detail="No se pudo procesar el archivo")
        
        # FunciÃ³n para limpiar contenido HTML mal formado
        def limpiar_contenido_html(texto: str) -> str:
            """Limpia contenido HTML mal formado y caracteres especiales"""
            if not texto:
                return ""
            
            import re
            
            # PRIMERA PASADA: Eliminar completamente todos los fragmentos HTML malformados
            # Patrones especÃ­ficos que aparecen en el texto
            patrones_malformados = [
                r'categoria="lesiones_permanentes" title="Click para ver detalles de lesiones_permanentes">lesionesass="frase-resaltada frase-lesiones" data-',
                r'lesionesass="frase-resaltada frase-lesiones" data-categoria="lesiones_permanentes" title="Click para ver detalles de lesiones_permanentes">lesiones',
                r'frase-resaltada frase-lesiones"',
                r'Click para ver detalles de lesiones_permanentes',
                r'indemnizaciÃ³nfrase-resaltada frase-prestaciones" data-categoria="prestaciones" title="Click para ver detalles de prestaciones">indemnizaciÃ³n',
                r'accidente de trabajoesaltada frase-accidente" data-categoria="accidente_laboral" title="Click para ver detalles de accidente_laboral">accidente de trabajo',
                r'INSSn class="frase-resaltada frase-inss" data-categoria="inss" title="Click para ver detalles de inss">INSS',
                r'reclamaciÃ³n="frase-resaltada frase-reclamacion" data-categoria="reclamacion_administrativa" title="Click para ver detalles de reclamacion_administrativa">reclamaciÃ³n',
                r'EVIan class="frase-resaltada frase-inss" data-categoria="inss" title="Click para ver detalles de inss">EVI',
                r'fundamento jurÃ­dicoresaltada frase-fundamentos" data-categoria="fundamentos_juridicos" title="Click para ver detalles de fundamentos_juridicos">fundamento jurÃ­dico',
                r'Seguridad Socialse-resaltada frase-inss" data-categoria="inss" title="Click para ver detalles de inss">Seguridad Social',
                r'Instituto Nacional-resaltada frase-inss" data-categoria="inss" title="Click para ver detalles de inss">Instituto Nacional',
                r'Instituto Nacional de la Seguridad Socialdata-categoria="inss" title="Click para ver detalles de inss">Instituto Nacional de la Seguridad Social',
                r'estimamosss="frase-resaltada frase-procedimiento" data-categoria="procedimiento_legal" title="Click para ver detalles de procedimiento_legal">estimamos'
            ]
            
            for patron in patrones_malformados:
                texto = re.sub(patron, '', texto, flags=re.IGNORECASE)
            
            # SEGUNDA PASADA: Eliminar todas las etiquetas HTML restantes
            texto = re.sub(r'<[^>]*>', '', texto)
            
            # TERCERA PASADA: Eliminar caracteres de escape HTML
            texto = re.sub(r'&[a-zA-Z0-9#]+;', ' ', texto)
            
            # CUARTA PASADA: Limpiar caracteres extraÃ±os y fragmentos de etiquetas
            texto = re.sub(r'[>]+', '', texto)
            texto = re.sub(r'[<]+', '', texto)
            texto = re.sub(r'frase-[a-zA-Z]+"', '', texto)
            texto = re.sub(r'onclick="[^"]*"', '', texto)
            texto = re.sub(r'title="[^"]*"', '', texto)
            texto = re.sub(r'resaltada', '', texto)
            texto = re.sub(r'data-categoria="[^"]*"', '', texto)
            texto = re.sub(r'class="[^"]*"', '', texto)
            texto = re.sub(r'style="[^"]*"', '', texto)
            
            # QUINTA PASADA: Limpiar fragmentos adicionales comunes
            texto = re.sub(r'esaltada frase-[a-zA-Z]+"', '', texto)
            texto = re.sub(r'lesionesass=', '', texto)
            texto = re.sub(r'frase-[a-zA-Z]+" data-categoria=', '', texto)
            texto = re.sub(r'Click para ver detalles de', '', texto)
            texto = re.sub(r'frase-resaltada frase-[a-zA-Z]+', '', texto)
            
            # SEXTA PASADA: Limpiar espacios mÃºltiples y normalizar
            texto = re.sub(r'\s+', ' ', texto)
            
            return texto.strip()
        
        # Obtener el texto original SIN procesamiento de resaltado
        texto_original = resultado.get("texto_extraido", "")
        
        # Limpiar completamente el contenido - SOLO TEXTO PLANO
        contenido_limpio = limpiar_contenido_html(texto_original)
        
        # Limpieza adicional mÃ¡s agresiva para eliminar cualquier resto de HTML
        import re
        contenido_limpio = re.sub(r'<[^>]*>', '', contenido_limpio)  # Eliminar cualquier etiqueta HTML restante
        contenido_limpio = re.sub(r'&[a-zA-Z0-9#]+;', ' ', contenido_limpio)  # Eliminar entidades HTML
        contenido_limpio = re.sub(r'[<>]', '', contenido_limpio)  # Eliminar caracteres < y >
        contenido_limpio = re.sub(r'\s+', ' ', contenido_limpio)  # Normalizar espacios
        contenido_limpio = contenido_limpio.strip()
        
        logger.info(f"ðŸ§¹ Contenido limpio (backend, primeros 500 chars): {contenido_limpio[:500]}...")
        
        # Preparar datos para el template - SIN resaltado automÃ¡tico
        datos_archivo = {
            "nombre": archivo_id,
            "contenido": contenido_limpio,
            "longitud": resultado.get("longitud_texto", 0),
            "frases_clave": {},  # DESHABILITAR resaltado automÃ¡tico
            "prediccion": resultado.get("prediccion", {}),
            "argumentos": resultado.get("argumentos", []),
            "insights": resultado.get("insights_juridicos", []),
            "total_frases": resultado.get("total_frases_clave", 0),
            "highlight_info": {
                "frase": highlight,
                "posicion": pos,
                "index": index
            } if highlight else None,
            "resaltado_deshabilitado": True  # Flag para indicar que el resaltado estÃ¡ deshabilitado
        }
        
        logger.info(f"ðŸ“„ Devolviendo template archivo.html para {archivo_id}")
        logger.info(f"ðŸ“Š Datos del archivo: procesado={datos_archivo.get('procesado')}, frases_clave={len(datos_archivo.get('frases_clave', {}))}")
        
        return templates.TemplateResponse("archivo.html", {
            "request": request,
            "archivo": datos_archivo
        })
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error mostrando archivo {archivo_id}: {e}")
        raise HTTPException(status_code=500, detail=f"Error interno: {str(e)}")


@app.get("/api/analizar")
async def api_analizar():
    """Endpoint API para anÃ¡lisis"""
    try:
        resultado = analizar_sentencias_existentes()
        # Asegurar que siempre devuelve los campos esperados
        if "error" not in resultado:
            resultado["archivos_analizados"] = resultado.get("archivos_analizados", 0)
            resultado["total_apariciones"] = resultado.get("total_apariciones", 0)
            resultado["resultados_por_archivo"] = resultado.get("resultados_por_archivo", {})
        return resultado
    except Exception as e:
        logger.error(f"Error en API: {e}")
        return {
            "error": f"Error al analizar: {str(e)}",
            "archivos_analizados": 0,
            "total_apariciones": 0,
            "resultados_por_archivo": {}
        }

@app.post("/api/limpiar-cache")
async def limpiar_cache():
    """Endpoint para limpiar el cachÃ© de anÃ¡lisis"""
    global CACHE_TIMESTAMP, ANALISIS_CACHE
    ANALISIS_CACHE = {}
    CACHE_TIMESTAMP = None
    logger.info("ðŸ—‘ï¸ CachÃ© limpiado")
    return JSONResponse(content={"mensaje": "CachÃ© limpiado correctamente"})


@app.get("/api/analisis-predictivo")
async def api_analisis_predictivo():
    """Endpoint API para anÃ¡lisis predictivo e inteligente de resoluciones"""
    try:
        # Importar funciones del mÃ³dulo de anÃ¡lisis predictivo
        try:
            from src.backend.analisis_predictivo import (
                realizar_analisis_predictivo,
                generar_insights_juridicos,
                identificar_patrones_favorables,
                extraer_factores_clave,
                generar_recomendaciones,
                calcular_confianza_analisis
            )
        except ImportError as e:
            logger.error(f"Error importando funciones de anÃ¡lisis predictivo: {e}")
            raise HTTPException(status_code=500, detail=f"Error interno del servidor: {str(e)}")
        
        # Obtener datos base
        resultado_base = analizar_sentencias_existentes()
        
        # Realizar anÃ¡lisis predictivo avanzado
        analisis_predictivo = realizar_analisis_predictivo(resultado_base)
        
        # Generar insights y recomendaciones
        insights = generar_insights_juridicos(resultado_base, analisis_predictivo)
        
        # Crear respuesta estructurada para UX mejorada
        return {
            "status": "success",
            "timestamp": datetime.now().isoformat(),
            "resumen_ejecutivo": {
                "total_documentos": resultado_base.get("archivos_analizados", 0),
                "total_frases_clave": resultado_base.get("total_apariciones", 0),
                "categorias_identificadas": len(resultado_base.get("ranking_global", {})),
                "confianza_analisis": calcular_confianza_analisis(resultado_base)
            },
            "analisis_predictivo": analisis_predictivo,
            "insights_juridicos": insights,
            "patrones_favorables": identificar_patrones_favorables(resultado_base),
            "factores_clave": extraer_factores_clave(resultado_base),
            "recomendaciones": generar_recomendaciones(resultado_base, analisis_predictivo),
            "datos_estadisticos": resultado_base,
            "metadata": {
                "modelo_ia": ANALIZADOR_IA_DISPONIBLE,
                "version_analisis": "2.0.0",
                "fecha_ultima_actualizacion": datetime.now().strftime("%Y-%m-%d %H:%M:%S")
            }
        }
        
    except Exception as e:
        logger.error(f"Error en anÃ¡lisis predictivo: {e}")
        return {
            "status": "error",
            "error": f"Error en anÃ¡lisis predictivo: {str(e)}",
            "timestamp": datetime.now().isoformat()
        }


@app.get("/analisis-predictivo")
async def pagina_analisis_predictivo(request: Request):
    """PÃ¡gina web para el anÃ¡lisis predictivo"""
    return templates.TemplateResponse("analisis_predictivo.html", {"request": request})


@app.get("/analisis-discrepancias/{archivo_id}")
async def pagina_analisis_discrepancias(request: Request, archivo_id: str):
    """PÃ¡gina web para mostrar anÃ¡lisis de discrepancias de un archivo especÃ­fico"""
    try:
        logger.info(f"ðŸ” Buscando archivo para anÃ¡lisis de discrepancias: {archivo_id}")
        
        # Buscar el archivo en ambos directorios
        archivo_path = None
        
        # Buscar en directorio sentencias (PDF y otros formatos)
        for extension in ["*.pdf", "*.txt", "*.docx"]:
            for archivo in Path("sentencias").glob(extension):
                if archivo_id in archivo.name or archivo.name in archivo_id:
                    archivo_path = archivo
                    logger.info(f"âœ… Archivo encontrado en sentencias/: {archivo}")
                    break
            if archivo_path:
                break
        
        # Si no se encuentra, buscar en directorio uploads
        if not archivo_path:
            for extension in ["*.pdf", "*.txt", "*.docx"]:
                for archivo in Path("uploads").glob(extension):
                    if archivo_id in archivo.name or archivo.name in archivo_id:
                        archivo_path = archivo
                        logger.info(f"âœ… Archivo encontrado en uploads/: {archivo}")
                        break
                if archivo_path:
                    break
        
        # BÃºsqueda mÃ¡s flexible: buscar por partes del nombre
        if not archivo_path:
            logger.info(f"ðŸ” BÃºsqueda flexible para: {archivo_id}")
            # Extraer partes del ID que podrÃ­an ser el nombre real
            partes_id = archivo_id.split('_')
            posibles_nombres = [p for p in partes_id if len(p) > 5 and not p.isdigit()]
            
            for nombre_posible in posibles_nombres:
                logger.info(f"ðŸ” Buscando archivos que contengan: {nombre_posible}")
                for extension in ["*.pdf", "*.txt", "*.docx"]:
                    for archivo in Path("sentencias").glob(extension):
                        if nombre_posible in archivo.name:
                            archivo_path = archivo
                            logger.info(f"âœ… Archivo encontrado por bÃºsqueda flexible en sentencias/: {archivo}")
                            break
                    if archivo_path:
                        break
                    
                    for archivo in Path("uploads").glob(extension):
                        if nombre_posible in archivo.name:
                            archivo_path = archivo
                            logger.info(f"âœ… Archivo encontrado por bÃºsqueda flexible en uploads/: {archivo}")
                            break
                    if archivo_path:
                        break
                if archivo_path:
                    break
        
        # BÃºsqueda final: buscar archivos que contengan 'informe' y parte del hash
        if not archivo_path and 'informe' in archivo_id:
            logger.info("ðŸ” BÃºsqueda final por archivos de informe...")
            hash_parts = [p for p in archivo_id.split('_') if len(p) == 8 and p.isalnum()]
            for hash_part in hash_parts:
                for archivo in Path("sentencias").glob("*informe*.pdf"):
                    if hash_part in archivo.name:
                        archivo_path = archivo
                        logger.info(f"âœ… Archivo encontrado por bÃºsqueda de informe en sentencias/: {archivo}")
                        break
                if archivo_path:
                    break
                    
                for archivo in Path("uploads").glob("*informe*.pdf"):
                    if hash_part in archivo.name:
                        archivo_path = archivo
                        logger.info(f"âœ… Archivo encontrado por bÃºsqueda de informe en uploads/: {archivo}")
                        break
                if archivo_path:
                    break
        
        if not archivo_path:
            logger.error(f"âŒ Archivo no encontrado: {archivo_id}")
            # Listar archivos disponibles para debug
            archivos_sentencias = list(Path("sentencias").glob("*"))
            archivos_uploads = list(Path("uploads").glob("*"))
            logger.info(f"Archivos en sentencias/: {[f.name for f in archivos_sentencias]}")
            logger.info(f"Archivos en uploads/: {[f.name for f in archivos_uploads]}")
            raise HTTPException(status_code=404, detail=f"Archivo no encontrado: {archivo_id}")
        
        # Realizar anÃ¡lisis del archivo
        logger.info(f"ðŸ”¬ Iniciando anÃ¡lisis de discrepancias para: {archivo_path}")
        try:
            if ANALIZADOR_IA_DISPONIBLE:
                try:
                    from src.backend.analisis import AnalizadorLegal
                except ImportError as e:
                    logger.error(f"Error importando AnalizadorLegal: {e}")
                    raise HTTPException(status_code=500, detail=f"Error interno del servidor: {str(e)}")
                analizador = AnalizadorLegal()
                resultado = analizador.analizar_documento(str(archivo_path))
                logger.info("âœ… AnÃ¡lisis con IA completado")
            else:
                resultado = analizador_basico.analizar_documento(str(archivo_path), archivo_path.name)
                logger.info("âœ… AnÃ¡lisis bÃ¡sico completado")
            
            # Generar anÃ¡lisis de discrepancias especÃ­fico usando el mÃ³dulo avanzado
            logger.info("ðŸ” Generando anÃ¡lisis de discrepancias avanzado...")
            try:
                from src.backend.analisis_discrepancias import AnalizadorDiscrepancias
                analizador_discrepancias = AnalizadorDiscrepancias()
                analisis_discrepancias = analizador_discrepancias.analizar_discrepancias(
                    resultado.get("texto_extraido", ""), 
                    archivo_path.name
                )
                logger.info("âœ… AnÃ¡lisis de discrepancias avanzado completado")
            except Exception as e:
                logger.warning(f"Fallback a anÃ¡lisis bÃ¡sico: {e}")
                analisis_discrepancias = generar_analisis_discrepancias_basico(str(archivo_path), resultado)
                logger.info("âœ… AnÃ¡lisis de discrepancias bÃ¡sico completado")
            
            resultado["analisis_discrepancias"] = analisis_discrepancias
            
        except Exception as e:
            logger.error(f"âŒ Error en anÃ¡lisis: {e}")
            raise HTTPException(status_code=500, detail=f"Error en anÃ¡lisis: {str(e)}")
        
        return templates.TemplateResponse("analisis_discrepancias.html", {
            "request": request,
            "resultado": resultado
        })
        
    except Exception as e:
        logger.error(f"Error en anÃ¡lisis de discrepancias: {e}")
        raise HTTPException(status_code=500, detail=f"Error procesando archivo: {str(e)}")


@app.get("/test-analisis-discrepancias/{archivo_id}")
async def test_analisis_discrepancias(archivo_id: str):
    """Endpoint de prueba para verificar que el anÃ¡lisis de discrepancias funciona"""
    try:
        logger.info(f"ðŸ§ª Test endpoint - Buscando archivo: {archivo_id}")
        
        # Buscar archivo
        archivo_path = None
        for extension in ["*.pdf", "*.txt", "*.docx"]:
            for archivo in Path("sentencias").glob(extension):
                if archivo_id in archivo.name:
                    archivo_path = archivo
                    break
            if archivo_path:
                break
        
        if not archivo_path:
            archivos_disponibles = [f.name for f in Path("sentencias").glob("*")]
            return {"error": f"Archivo no encontrado: {archivo_id}", "archivos_disponibles": archivos_disponibles}
        
        return {
            "archivo_encontrado": str(archivo_path),
            "archivo_id": archivo_id,
            "status": "ok"
        }
        
    except Exception as e:
        return {"error": str(e)}


@app.get("/listar-archivos")
async def listar_archivos_disponibles():
    """Endpoint para listar todos los archivos disponibles para anÃ¡lisis"""
    try:
        archivos_sentencias = [f.name for f in Path("sentencias").glob("*") if f.is_file()]
        archivos_uploads = [f.name for f in Path("uploads").glob("*") if f.is_file()]
        
        # DEBUG: Log detallado de archivos
        logger.info(f"ðŸ“ ARCHIVOS EN SENTENCIAS/ ({len(archivos_sentencias)}): {archivos_sentencias}")
        logger.info(f"ðŸ“ ARCHIVOS EN UPLOADS/ ({len(archivos_uploads)}): {archivos_uploads}")
        
        return {
            "archivos_sentencias": archivos_sentencias,
            "archivos_uploads": archivos_uploads,
            "total_sentencias": len(archivos_sentencias),
            "total_uploads": len(archivos_uploads),
            "directorio_actual": str(Path.cwd()),
            "directorio_sentencias": str(Path("sentencias").resolve()),
            "directorio_uploads": str(Path("uploads").resolve()),
            "existe_sentencias": Path("sentencias").exists(),
            "existe_uploads": Path("uploads").exists()
        }
        
    except Exception as e:
        logger.error(f"Error listando archivos: {e}")
        return {"error": str(e)}


@app.get("/test-sts2384")
async def test_sts2384():
    """Endpoint especÃ­fico para probar STS_2384_2025"""
    try:
        archivo_id = "STS_2384_2025"
        
        # Buscar archivo
        archivo_path = None
        for extension in ["*.pdf", "*.txt", "*.docx"]:
            for archivo in Path("sentencias").glob(extension):
                if archivo_id in archivo.name:
                    archivo_path = archivo
                    break
            if archivo_path:
                break
        
        if not archivo_path:
            for extension in ["*.pdf", "*.txt", "*.docx"]:
                for archivo in Path("uploads").glob(extension):
                    if archivo_id in archivo.name:
                        archivo_path = archivo
                        break
                if archivo_path:
                    break
        
        return {
            "archivo_id": archivo_id,
            "archivo_encontrado": str(archivo_path) if archivo_path else None,
            "existe": archivo_path is not None,
            "archivos_sentencias": [f.name for f in Path("sentencias").glob("STS*")],
            "archivos_uploads": [f.name for f in Path("uploads").glob("STS*")]
        }
        
    except Exception as e:
        return {"error": str(e)}


@app.get("/diagnostico")
async def pagina_diagnostico(request: Request):
    """PÃ¡gina de diagnÃ³stico del sistema"""
    return templates.TemplateResponse("diagnostico.html", {"request": request})


@app.post("/api/descargar-informe-discrepancias")
async def descargar_informe_discrepancias(request: Request):
    """Genera y descarga un informe completo de discrepancias en formato Word"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.enum.style import WD_STYLE_TYPE
        
        # Obtener datos del request
        datos = await request.json()
        nombre_archivo = datos.get("nombre_archivo", "archivo_desconocido")
        analisis = datos.get("analisis_discrepancias", {})
        timestamp = datos.get("timestamp", "")
        
        # Crear documento Word
        try:
            doc = Document()
            logger.info("âœ… Documento Word creado exitosamente")
        except Exception as e:
            logger.error(f"âŒ Error creando documento Word: {e}")
            raise HTTPException(status_code=500, detail=f"Error creando documento Word: {str(e)}")
        
        # TÃ­tulo principal
        titulo = doc.add_heading('ANÃLISIS DE DISCREPANCIAS MÃ‰DICAS-LEGALES', 0)
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # InformaciÃ³n del archivo
        doc.add_heading('InformaciÃ³n del Archivo', level=1)
        doc.add_paragraph(f"Archivo: {nombre_archivo}")
        doc.add_paragraph(f"Fecha de anÃ¡lisis: {timestamp}")
        doc.add_paragraph(f"MÃ©todo: AnÃ¡lisis automÃ¡tico con IA")
        
        # Resumen ejecutivo
        doc.add_heading('Resumen Ejecutivo', level=1)
        discrepancias = analisis.get("discrepancias_detectadas", [])
        evidencia = analisis.get("evidencia_favorable", [])
        puntuacion = analisis.get("puntuacion_discrepancia", 0)
        probabilidad = analisis.get("probabilidad_ipp", 0)
        
        doc.add_paragraph(f"â€¢ Discrepancias detectadas: {len(discrepancias)}")
        doc.add_paragraph(f"â€¢ Evidencia favorable: {len(evidencia)} elementos")
        doc.add_paragraph(f"â€¢ PuntuaciÃ³n discrepancia: {puntuacion}/100")
        doc.add_paragraph(f"â€¢ Probabilidad IPP: {probabilidad:.1%}")
        
        # ConclusiÃ³n
        if probabilidad >= 0.7:
            conclusion = "ALTA PROBABILIDAD DE IPP"
        elif probabilidad >= 0.5:
            conclusion = "PROBABILIDAD MEDIA DE IPP"
        else:
            conclusion = "BAJA PROBABILIDAD DE IPP"
        
        doc.add_paragraph(f"ConclusiÃ³n: {conclusion}")
        
        # Discrepancias detectadas
        if discrepancias:
            doc.add_heading('Discrepancias Detectadas', level=1)
            for i, disc in enumerate(discrepancias, 1):
                doc.add_heading(f"{i}. {disc.get('tipo', '').replace('_', ' ').title()}", level=2)
                doc.add_paragraph(f"DescripciÃ³n: {disc.get('descripcion', '')}")
                doc.add_paragraph(f"Severidad: {disc.get('severidad', '')}")
                doc.add_paragraph(f"Argumento jurÃ­dico: {disc.get('argumento_juridico', '')}")
        
        # Evidencia favorable
        if evidencia:
            doc.add_heading('Evidencia Favorable para IPP', level=1)
            for i, ev in enumerate(evidencia, 1):
                doc.add_heading(f"{i}. {ev.get('tipo', '').replace('_', ' ').title()}", level=2)
                doc.add_paragraph(f"DescripciÃ³n: {ev.get('descripcion', '')}")
                doc.add_paragraph(f"Relevancia: {ev.get('relevancia', '')}")
                doc.add_paragraph(f"Argumento: {ev.get('argumento', '')}")
        
        # Argumentos jurÃ­dicos
        argumentos = analisis.get("argumentos_juridicos", [])
        if argumentos:
            doc.add_heading('Argumentos JurÃ­dicos Generados', level=1)
            for i, arg in enumerate(argumentos, 1):
                doc.add_heading(f"{i}. {arg.get('titulo', '')}", level=2)
                doc.add_paragraph(f"Contenido: {arg.get('contenido', '')}")
                doc.add_paragraph(f"Fuerza: {arg.get('fuerza', '')}")
        
        # Recomendaciones de defensa
        recomendaciones = analisis.get("recomendaciones_defensa", [])
        if recomendaciones:
            doc.add_heading('Recomendaciones de Defensa', level=1)
            for i, rec in enumerate(recomendaciones, 1):
                doc.add_heading(f"{i}. {rec.get('titulo', '')}", level=2)
                doc.add_paragraph(f"Contenido: {rec.get('contenido', '')}")
                doc.add_paragraph(f"Prioridad: {rec.get('prioridad', '')}")
                
                acciones = rec.get('acciones', [])
                if acciones:
                    doc.add_paragraph("Acciones recomendadas:")
                    for accion in acciones:
                        doc.add_paragraph(f"â€¢ {accion}", style='List Bullet')
        
        # Contradicciones internas
        contradicciones = analisis.get("contradicciones_internas", [])
        if contradicciones:
            doc.add_heading('Contradicciones Internas Detectadas', level=1)
            for i, cont in enumerate(contradicciones, 1):
                doc.add_heading(f"{i}. ContradicciÃ³n Interna", level=2)
                doc.add_paragraph(f"DescripciÃ³n: {cont.get('descripcion', '')}")
                doc.add_paragraph(f"Texto detectado: {cont.get('texto', '')}")
                doc.add_paragraph(f"Argumento: {cont.get('argumento', '')}")
        
        # Guardar en memoria con validaciÃ³n
        from io import BytesIO
        buffer = BytesIO()
        
        try:
            doc.save(buffer)
            buffer.seek(0)
            
            # Validar que el documento se generÃ³ correctamente
            if buffer.getvalue() is None or len(buffer.getvalue()) == 0:
                raise Exception("El documento generado estÃ¡ vacÃ­o")
            
            logger.info(f"âœ… Documento Word generado exitosamente: {len(buffer.getvalue())} bytes")
            
        except Exception as e:
            logger.error(f"âŒ Error guardando documento Word: {e}")
            raise HTTPException(status_code=500, detail=f"Error guardando documento: {str(e)}")
        
        # Preparar respuesta
        from fastapi.responses import Response
        filename = f"informe_discrepancias_{nombre_archivo.replace('.pdf', '').replace('.txt', '')}.docx"
        
        return Response(
            content=buffer.getvalue(),
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            headers={
                "Content-Disposition": f"attachment; filename={filename}",
                "X-Content-Type-Options": "nosniff",
                "Cache-Control": "no-cache, no-store, must-revalidate",
                "Pragma": "no-cache",
                "Expires": "0"
            }
        )
        
    except Exception as e:
        logger.error(f"Error generando informe Word: {e}")
        raise HTTPException(status_code=500, detail=f"Error generando informe: {str(e)}")


@app.post("/api/descargar-resumen-pdf")
async def descargar_resumen_pdf(request: Request):
    """Genera y descarga un resumen en formato PDF"""
    try:
        from reportlab.lib.pagesizes import letter, A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, PageBreak
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from io import BytesIO
        
        # Obtener datos del request
        datos = await request.json()
        contenido = datos.get("contenido", "")
        nombre_archivo = datos.get("nombre_archivo", "archivo_desconocido")
        
        # Crear documento PDF en memoria
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=A4, rightMargin=72, leftMargin=72, topMargin=72, bottomMargin=18)
        
        # Estilos
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=16,
            spaceAfter=30,
            alignment=1  # Centrado
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            spaceAfter=12,
            spaceBefore=20
        )
        
        # Contenido
        story = []
        
        # Convertir HTML a texto plano para PDF
        import re
        contenido_texto = re.sub(r'<[^>]+>', '', contenido)
        contenido_texto = contenido_texto.replace('&nbsp;', ' ')
        
        # Dividir en pÃ¡rrafos
        parrafos = contenido_texto.split('\n')
        
        for parrafo in parrafos:
            parrafo = parrafo.strip()
            if not parrafo:
                continue
                
            if parrafo.startswith('ANÃLISIS DE DISCREPANCIAS'):
                story.append(Paragraph(parrafo, title_style))
            elif parrafo.startswith('Archivo:') or parrafo.startswith('Fecha de anÃ¡lisis:'):
                story.append(Paragraph(parrafo, styles['Normal']))
            elif parrafo.startswith('RESUMEN EJECUTIVO') or parrafo.startswith('DISCREPANCIAS') or parrafo.startswith('EVIDENCIA') or parrafo.startswith('ARGUMENTOS') or parrafo.startswith('RECOMENDACIONES'):
                story.append(Spacer(1, 12))
                story.append(Paragraph(parrafo, heading_style))
            elif parrafo.startswith('â€¢'):
                story.append(Paragraph(parrafo, styles['Normal']))
            else:
                story.append(Paragraph(parrafo, styles['Normal']))
            
            story.append(Spacer(1, 6))
        
        # Construir PDF
        doc.build(story)
        buffer.seek(0)
        
        # Preparar respuesta
        from fastapi.responses import Response
        return Response(
            content=buffer.getvalue(),
            media_type="application/pdf",
            headers={
                "Content-Disposition": f"attachment; filename=resumen_discrepancias_{nombre_archivo.replace('.pdf', '')}.pdf"
            }
        )
        
    except Exception as e:
        logger.error(f"Error generando PDF: {e}")
        raise HTTPException(status_code=500, detail=f"Error generando PDF: {str(e)}")


# ====== DEMANDA BASE: generaciÃ³n desde fallos y fundamentos ======
def _leer_texto_archivo_simple(path: Path) -> str:
    try:
        if path.suffix.lower() == '.txt':
            for enc in ['utf-8', 'latin-1', 'cp1252']:
                try:
                    return path.read_text(encoding=enc)
                except UnicodeDecodeError:
                    continue
            return path.read_text(errors='ignore')
        else:
            # Delegar a analizador para extraer texto
            if ANALIZADOR_IA_DISPONIBLE:
                try:
                    from src.backend.analisis import AnalizadorLegal
                except ImportError as e:
                    logger.error(f"Error importando AnalizadorLegal: {e}")
                    raise HTTPException(status_code=500, detail=f"Error interno del servidor: {str(e)}")
                analizador = AnalizadorLegal()
                res = analizador.analizar_documento(str(path))
            else:
                res = analizador_basico.analizar_documento(str(path), path.name)
            return res.get('texto_extraido', '') or ''
    except Exception:
        return ''


def _extraer_seccion(texto: str, encabezados: List[str], fin_encabezados: List[str]) -> str:
    if not texto:
        return ''
    t = texto
    import re as _re
    # Construir regex de inicio y fin
    ini = r"|".join([_re.escape(h) for h in encabezados])
    fin = r"|".join([_re.escape(h) for h in fin_encabezados])
    patron_ini = _re.compile(rf"(?i)(?:^|\n)\s*(?:{ini})\b[:\-\s]*", _re.MULTILINE)
    patron_fin = _re.compile(rf"(?i)(?:^|\n)\s*(?:{fin})\b", _re.MULTILINE)
    m = patron_ini.search(t)
    if not m:
        return ''
    start = m.end()
    m2 = patron_fin.search(t, start)
    end = m2.start() if m2 else len(t)
    seccion = t[start:end].strip()
    # Limpiar artefactos HTML simples
    seccion = _re.sub(r"<[^>]+>", " ", seccion)
    seccion = _re.sub(r"\s+", " ", seccion).strip()
    return seccion


def _generar_demanda_base_para(paths: List[Path], meta: Dict[str, Any] = None) -> Dict[str, Any]:
    documentos: List[Dict[str, Any]] = []
    for p in paths:
        texto = _leer_texto_archivo_simple(p)
        fallo = _extraer_seccion(
            texto,
            encabezados=["FALLO", "PARTE DISPOSITIVA", "RESUELVO", "RESOLVEMOS"],
            fin_encabezados=["FUNDAMENTOS", "FUNDAMENTOS DE HECHO", "HECHOS", "FUNDAMENTOS DE DERECHO", "ANTECEDENTES", "SEGUNDO", "TERCERO"]
        )
        fundamentos = _extraer_seccion(
            texto,
            encabezados=["FUNDAMENTOS DE HECHO", "HECHOS PROBADOS", "ANTECEDENTES DE HECHO"],
            fin_encabezados=["FUNDAMENTOS DE DERECHO", "PARTE DISPOSITIVA", "FALLO", "RESUELVO", "RESOLVEMOS"]
        )
        # Resumenes breves
        fundamentos_resumen = _resumir_fundamentos(texto)
        fallo_breve = (fallo or "").strip()
        if len(fallo_breve) > 400:
            fallo_breve = fallo_breve[:400].rstrip() + "â€¦"
        documentos.append({
            "nombre": p.name,
            "fallo": fallo,
            "fundamentos": fundamentos,
            "fallo_breve": fallo_breve,
            "fundamentos_resumen": fundamentos_resumen
        })

    meta = meta or {}
    nombre = meta.get("nombre", "[NOMBRE DEMANDANTE]")
    dni = meta.get("dni", "[DNI]")
    domicilio = meta.get("domicilio", "[DOMICILIO A EFECTOS]")
    letrado = meta.get("letrado", "[LETRADO/A]")
    empresa = meta.get("empresa", "[EMPRESA]")
    profesion = meta.get("profesion", "[PROFESIÃ“N HABITUAL]")
    grado_principal = meta.get("grado_principal", "Incapacidad Permanente Total")
    grado_subsidiario = meta.get("grado_subsidiario", "Incapacidad Permanente Parcial")
    base_reguladora = meta.get("base_reguladora", "[BASE REGULADORA]")
    indemnizacion_parcial = meta.get("indemnizacion_parcial", "24 mensualidades")
    mutua = meta.get("mutua", "[MUTUA]")
    
    # Campos de HECHOS
    relacion_laboral = meta.get("relacion_laboral", f"La demandante presta servicios como {profesion} para la empresa {empresa}.")
    contingencia_evolucion = meta.get("contingencia_evolucion", "[Accidente de trabajo / enfermedad comÃºn], con periodos de IT y secuelas actuales.")
    actuaciones_administrativas = meta.get("actuaciones_administrativas", "(EVI, inicio IP, audiencia, resoluciÃ³n del INSS y ReclamaciÃ³n Previa).")
    cuadro_clinico = meta.get("cuadro_clinico", "[Describir secuelas relevantes y su impacto en las tareas fundamentales].")
    
    # Plantilla base de demanda (formal)
    cuerpo = {
        "encabezado": "AL JUZGADO DE LO SOCIAL QUE POR TURNO CORRESPONDA\n\n",
        "parte": (
            f"D./DÃ±a. {nombre}, con DNI {dni}, y domicilio a efectos de notificaciones en {domicilio}, "
            f"representado/a por Letrado/a {letrado}, ante el Juzgado comparece y DICE:\n\n"
        ),
        "hechos": None,  # se compone mÃ¡s abajo
        "fundamentos_derecho": (
            "FUNDAMENTOS DE DERECHO\n"
            "I. JurisdicciÃ³n y competencia (arts. 2 y 6 LRJS).\n"
            "II. LegitimaciÃ³n activa y pasiva (LRJS).\n"
            "III. Fondo del asunto. Arts. 193 y 194 LGSS (grados de incapacidad). Art. 194.2 LGSS (IPP = 24 mensualidades).\n"
            "IV. Doctrina jurisprudencial aplicable (STS 04-07-2025, rec. 1096/2024, sobre IPP subsidiaria; TSJ Castilla y LeÃ³n sobre limitaciones en limpiadoras).\n"
            "V. Principios y derechos constitucionales (art. 24 CE tutela judicial efectiva; 9.3 CE interdicciÃ³n de la arbitrariedad).\n\n"
        ),
        "peticion": (
            "SUPLICO AL JUZGADO:\n"
            f"Primero.- Que, estimando la demanda, se declare a la actora afecta a {grado_principal} para su profesiÃ³n habitual de {profesion}, derivada de [contingencia], con derecho a la prestaciÃ³n correspondiente sobre una base reguladora de {base_reguladora}.\n"
            f"Subsidiariamente.- Que, para el caso de no apreciarse lo anterior, se declare la {grado_subsidiario}, con derecho a la indemnizaciÃ³n de {indemnizacion_parcial} de la base reguladora, a cargo de la Mutua {mutua} en caso de accidente de trabajo.\n"
            "Con expresa condena en costas a la parte demandada en los tÃ©rminos legalmente procedentes.\n\n"
        ),
    }

    # HECHOS (limpios y propios)
    cuerpo["hechos"] = (
        "HECHOS\n"
        f"1. RelaciÃ³n laboral. {relacion_laboral}\n"
        f"2. Contingencia y evoluciÃ³n. {contingencia_evolucion}\n"
        f"3. Actuaciones administrativas. {actuaciones_administrativas}\n"
        f"4. Cuadro clÃ­nico y limitaciones. {cuadro_clinico}\n\n"
    )

    # Resumen de fallos extraÃ­dos
    resumen_fallos = []
    for d in documentos:
        if d["fallo"]:
            resumen_fallos.append(f"â€” {d['nombre']}: {d['fallo']}")

    anexos = []
    for i, d in enumerate(documentos, start=1):
        if d["fallo"]:
            anexos.append(f"Documento nÂº {i}: Parte dispositiva de {d['nombre']}")

    # Jurisprudencia de apoyo (extractada)
    juris = []
    for d in documentos:
        linea = f"â€” {d['nombre']}: {d.get('fallo_breve','').strip()}"
        if d.get("fundamentos_resumen"):
            linea += "\n   Fundamentos: " + "; ".join(d["fundamentos_resumen"][:2])
        juris.append(linea.strip())

    doc_txt = (
        cuerpo["encabezado"] +
        cuerpo["parte"] +
        ("JURISPRUDENCIA DE APOYO (extracto de fallos)\n" + "\n\n".join(juris) + "\n\n" if juris else "") +
        cuerpo["hechos"] +
        cuerpo["fundamentos_derecho"] +
        cuerpo["peticion"] +
        ("ANEXO: RelaciÃ³n de documentos adjuntos\n" + "\n".join(anexos) + "\n" if anexos else "")
    )

    return {
        "documentos": documentos,
        "texto": doc_txt
    }


def _generar_demanda_docx(paths: List[Path], meta: Dict[str, Any] = None) -> BytesIO:
    """Genera un documento DOCX con formato profesional para la demanda"""
    # Obtener el texto de la demanda
    demanda_data = _generar_demanda_base_para(paths, meta)
    texto_demanda = demanda_data.get("texto", "")
    
    # Crear documento Word
    doc = Document()
    
    # Configurar estilos del documento
    styles = doc.styles
    
    # Estilo para el tÃ­tulo principal
    title_style = styles.add_style('DemandaTitle', WD_STYLE_TYPE.PARAGRAPH)
    title_font = title_style.font
    title_font.name = 'Times New Roman'
    title_font.size = Pt(16)
    title_font.bold = True
    title_style.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_style.paragraph_format.space_after = Pt(12)
    
    # Estilo para encabezados de secciÃ³n
    heading_style = styles.add_style('DemandaHeading', WD_STYLE_TYPE.PARAGRAPH)
    heading_font = heading_style.font
    heading_font.name = 'Times New Roman'
    heading_font.size = Pt(14)
    heading_font.bold = True
    heading_style.paragraph_format.space_before = Pt(12)
    heading_style.paragraph_format.space_after = Pt(6)
    
    # Estilo para texto normal
    normal_style = styles.add_style('DemandaNormal', WD_STYLE_TYPE.PARAGRAPH)
    normal_font = normal_style.font
    normal_font.name = 'Times New Roman'
    normal_font.size = Pt(12)
    normal_style.paragraph_format.space_after = Pt(6)
    normal_style.paragraph_format.line_spacing = 1.5
    
    # Estilo para texto con sangrÃ­a
    indent_style = styles.add_style('DemandaIndent', WD_STYLE_TYPE.PARAGRAPH)
    indent_font = indent_style.font
    indent_font.name = 'Times New Roman'
    indent_font.size = Pt(12)
    indent_style.paragraph_format.left_indent = Inches(0.5)
    indent_style.paragraph_format.space_after = Pt(6)
    indent_style.paragraph_format.line_spacing = 1.5
    
    # Configurar mÃ¡rgenes del documento
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)
    
    # Procesar el texto lÃ­nea por lÃ­nea
    lines = texto_demanda.split('\n')
    
    for line in lines:
        line = line.strip()
        if not line:
            # LÃ­nea vacÃ­a - agregar espacio
            doc.add_paragraph('', style='DemandaNormal')
            continue
            
        # Detectar tipo de contenido
        if line.upper().startswith('AL JUZGADO'):
            # TÃ­tulo principal
            p = doc.add_paragraph(line, style='DemandaTitle')
        elif any(line.upper().startswith(keyword) for keyword in ['HECHOS', 'FUNDAMENTOS', 'SUPLICO', 'JURISPRUDENCIA', 'ANEXO']):
            # Encabezado de secciÃ³n
            p = doc.add_paragraph(line, style='DemandaHeading')
        elif line.startswith('D./DÃ±a.') or line.startswith('D./DÃ±a'):
            # InformaciÃ³n de la parte
            p = doc.add_paragraph(line, style='DemandaNormal')
        elif line.startswith(('Primero.-', 'Segundo.-', 'Tercero.-', 'Subsidiariamente.-')):
            # Peticiones numeradas
            p = doc.add_paragraph(line, style='DemandaIndent')
        elif line.startswith(('â€”', 'â€¢', '1.', '2.', '3.', '4.')):
            # Lista con viÃ±etas
            p = doc.add_paragraph(line, style='DemandaIndent')
        elif line.startswith('I.') or line.startswith('II.') or line.startswith('III.') or line.startswith('IV.') or line.startswith('V.'):
            # NumeraciÃ³n romana
            p = doc.add_paragraph(line, style='DemandaIndent')
        else:
            # Texto normal
            p = doc.add_paragraph(line, style='DemandaNormal')
    
    # Guardar en memoria
    docx_buffer = BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)
    
    return docx_buffer


# ====== EXTRACTOR ESTRUCTURADO PARA DEMANDA ======
def _inferir_instancia_desde_texto(texto: str) -> str:
    t = (texto or '').lower()
    if 'tribunal supremo' in t:
        return 'TS'
    if 'tribunal superior de justicia' in t or 'tsj' in t:
        return 'TSJ'
    return 'otra'


def _extraer_primera_fecha(texto: str) -> Optional[str]:
    import re as _re
    for patron in [r"\b\d{1,2}[\-/\.\s]\d{1,2}[\-/\.\s]\d{2,4}\b", r"\b\d{4}[\-/\.]\d{1,2}[\-/\.]\d{1,2}\b"]:
        m = _re.search(patron, texto)
        if m:
            return m.group(0)
    return None


def _extraer_por_regex(texto: str, regex: str, group: int = 1) -> Optional[str]:
    import re as _re
    m = _re.search(regex, texto, _re.IGNORECASE)
    return m.group(group).strip() if m else None


def _resumir_fundamentos(texto: str) -> List[str]:
    seccion = _extraer_seccion(
        texto,
        ["FUNDAMENTOS DE DERECHO", "FUNDAMENTOS"],
        ["FALLO", "PARTE DISPOSITIVA", "RESUELVO", "RESOLVEMOS", "SUPLICO", "HECHOS"]
    )
    if not seccion:
        return []
    # dividir en frases y tomar las 3 primeras relevantes
    import re as _re
    frases = [f.strip() for f in _re.split(r"(?<=[\.!?])\s+", seccion) if len(f.strip()) > 30]
    return frases[:3]


@app.post("/api/extract/demanda")
async def api_extract_demanda(payload: Dict[str, Any]):
    try:
        nombres: List[str] = payload.get("nombres_archivo") or []
        if not isinstance(nombres, list) or not nombres:
            raise HTTPException(status_code=400, detail="Debe indicar 'nombres_archivo' (lista)")
        # Buscar archivos en ambos directorios
        paths = []
        for nombre in nombres:
            candidatos = [SENTENCIAS_DIR / nombre, UPLOADS_DIR / nombre]
            for candidato in candidatos:
                if candidato.exists():
                    paths.append(candidato)
                    break
        if not paths:
            raise HTTPException(status_code=404, detail="No se encontraron los archivos indicados")

        docs_out: List[Dict[str, Any]] = []
        # Sugerencias extraÃ­das dinÃ¡micamente del documento
        sugerencias: Dict[str, Any] = {
            "profesion": None, 
            "empresa": None, 
            "mutua": None, 
            "base_reguladora": None
        }

        for p in paths:
            texto = _leer_texto_archivo_simple(p)
            fallo = _extraer_seccion(texto, ["FALLO", "PARTE DISPOSITIVA", "RESUELVO", "RESOLVEMOS"], ["FUNDAMENTOS", "HECHOS", "ANTECEDENTES"]) or ""
            fundamentos_resumen = _resumir_fundamentos(texto)
            instancia = _inferir_instancia_desde_texto(texto)
            fecha = _extraer_primera_fecha(texto)
            organo = None
            # aproximaciÃ³n al Ã³rgano
            for k in ["TRIBUNAL SUPREMO", "TRIBUNAL SUPERIOR DE JUSTICIA", "AUDIENCIA", "JUZGADO DE LO SOCIAL"]:
                if k.lower() in (texto or '').lower():
                    organo = k.title()
                    break

            # Extraer sugerencias del documento
            if sugerencias["profesion"] is None:
                profesion = _extraer_por_regex(texto, r"profesi[oÃ³]n\s+habitual\s+(de|:)?\s*([A-Za-zÃÃ‰ÃÃ“ÃšÃœÃ‘Ã¡Ã©Ã­Ã³ÃºÃ¼Ã± ]+)", 2)
                if profesion:
                    sugerencias["profesion"] = profesion
            if sugerencias["empresa"] is None:
                empresa = _extraer_por_regex(texto, r"emplead[oa]\s+por\s+([A-ZÃÃ‰ÃÃ“ÃšÃœÃ‘a-zÃ¡Ã©Ã­Ã³ÃºÃ¼Ã±0-9 .,&;-]+)")
                if empresa:
                    sugerencias["empresa"] = empresa
            if sugerencias["mutua"] is None:
                mutua = _extraer_por_regex(texto, r"mutua\s+([A-ZÃÃ‰ÃÃ“ÃšÃœÃ‘][A-Za-zÃÃ‰ÃÃ“ÃšÃœÃ‘ ]+)")
                if mutua:
                    sugerencias["mutua"] = mutua
            if sugerencias["base_reguladora"] is None:
                br = _extraer_por_regex(texto, r"base\s+reguladora[^0-9]*([0-9\.,]+)")
                if br:
                    sugerencias["base_reguladora"] = br

            docs_out.append({
                "archivo": p.name,
                "instancia": instancia,
                "fecha": fecha,
                "organo": organo,
                "fallo": fallo,
                "fundamentos_resumen": fundamentos_resumen
            })

        return {"documentos": docs_out, "sugerencias_meta": sugerencias}

    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error en extracciÃ³n de demanda: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.get("/api/sugerencias-test")
async def api_sugerencias_test():
    """Endpoint de prueba para verificar las sugerencias extraÃ­das"""
    # Usar el primer documento disponible para extraer sugerencias
    archivos = list(SENTENCIAS_DIR.glob("*.pdf")) + list(SENTENCIAS_DIR.glob("*.txt"))
    if archivos:
        texto = _leer_texto_archivo_simple(archivos[0])
        sugerencias = {
            "profesion": None, 
            "empresa": None, 
            "mutua": None, 
            "base_reguladora": None
        }
        
        # Extraer sugerencias del documento
        profesion = _extraer_por_regex(texto, r"profesi[oÃ³]n\s+habitual\s+(de|:)?\s*([A-Za-zÃÃ‰ÃÃ“ÃšÃœÃ‘Ã¡Ã©Ã­Ã³ÃºÃ¼Ã± ]+)", 2)
        if profesion:
            sugerencias["profesion"] = profesion
        empresa = _extraer_por_regex(texto, r"emplead[oa]\s+por\s+([A-ZÃÃ‰ÃÃ“ÃšÃœÃ‘a-zÃ¡Ã©Ã­Ã³ÃºÃ¼Ã±0-9 .,&;-]+)")
        if empresa:
            sugerencias["empresa"] = empresa
        mutua = _extraer_por_regex(texto, r"mutua\s+([A-ZÃÃ‰ÃÃ“ÃšÃœÃ‘][A-Za-zÃÃ‰ÃÃ“ÃšÃœÃ‘ ]+)")
        if mutua:
            sugerencias["mutua"] = mutua
        br = _extraer_por_regex(texto, r"base\s+reguladora[^0-9]*([0-9\.,]+)")
        if br:
            sugerencias["base_reguladora"] = br
    else:
        sugerencias = {
            "profesion": None, 
            "empresa": None, 
            "mutua": None, 
            "base_reguladora": None
        }
    
    return {"sugerencias_meta": sugerencias}


@app.get("/api/diagnostico/modelo")
async def api_diagnostico_modelo():
    """Endpoint para diagnÃ³stico detallado del modelo IA en producciÃ³n"""
    try:
        logger.info("ðŸ” Iniciando diagnÃ³stico del modelo...")
        
        diagnostico = {
            "timestamp": datetime.now().isoformat(),
            "ia_disponible_global": ANALIZADOR_IA_DISPONIBLE,
            "modelos_disponibles": {},
            "errores": [],
            "estado_importacion": {},
            "archivos_modelo": {},
            "prueba_carga": {}
        }
        
        # 1. Verificar archivos de modelo
        models_dir = Path("models")
        if models_dir.exists():
            archivos = list(models_dir.iterdir())
            diagnostico["archivos_modelo"]["existe_directorio"] = True
            diagnostico["archivos_modelo"]["archivos"] = [f.name for f in archivos]
            diagnostico["archivos_modelo"]["tamaÃ±os"] = {f.name: f.stat().st_size for f in archivos}
        else:
            diagnostico["archivos_modelo"]["existe_directorio"] = False
            diagnostico["errores"].append("Directorio 'models' no existe")
        
        # 2. Probar importaciÃ³n del mÃ³dulo backend
        try:
            from src.backend.analisis import AnalizadorLegal
            diagnostico["estado_importacion"]["backend_analisis"] = "âœ… OK"
        except Exception as e:
            diagnostico["estado_importacion"]["backend_analisis"] = f"âŒ Error: {e}"
            diagnostico["errores"].append(f"Error importando backend.analisis: {e}")
        
        # 3. Probar carga directa del modelo
        try:
            import pickle
            with open('models/modelo_legal.pkl', 'rb') as f:
                modelo = pickle.load(f)
            diagnostico["prueba_carga"]["modelo_principal"] = "âœ… OK"
            diagnostico["prueba_carga"]["tipo_modelo"] = str(type(modelo))
            if isinstance(modelo, dict):
                diagnostico["prueba_carga"]["keys_modelo"] = list(modelo.keys())
        except Exception as e:
            diagnostico["prueba_carga"]["modelo_principal"] = f"âŒ Error: {e}"
            diagnostico["errores"].append(f"Error cargando modelo: {e}")
        
        # 4. Probar creaciÃ³n del analizador
        try:
            if "backend.analisis" in diagnostico["estado_importacion"]:
                analizador = AnalizadorLegal()
                diagnostico["prueba_carga"]["creacion_analizador"] = "âœ… OK"
                diagnostico["prueba_carga"]["modelo_ia_analizador"] = getattr(analizador, 'modelo', None) is not None
            else:
                diagnostico["prueba_carga"]["creacion_analizador"] = "âŒ No se puede crear - error de importaciÃ³n"
        except Exception as e:
            diagnostico["prueba_carga"]["creacion_analizador"] = f"âŒ Error: {e}"
            diagnostico["errores"].append(f"Error creando analizador: {e}")
        
        # 5. Probar anÃ¡lisis bÃ¡sico
        try:
            if "creacion_analizador" in diagnostico["prueba_carga"] and "âœ… OK" in diagnostico["prueba_carga"]["creacion_analizador"]:
                analizador = AnalizadorLegal()
                resultado = analizador.analizar_documento("sentencias/STS_2384_2025.pdf")
                diagnostico["prueba_carga"]["analisis_basico"] = "âœ… OK"
                diagnostico["prueba_carga"]["metodo_analisis"] = resultado.get("metodo_analisis", "desconocido")
                diagnostico["prueba_carga"]["modelo_ia_resultado"] = resultado.get("modelo_ia", False)
            else:
                diagnostico["prueba_carga"]["analisis_basico"] = "âŒ No se puede probar - error anterior"
        except Exception as e:
            diagnostico["prueba_carga"]["analisis_basico"] = f"âŒ Error: {e}"
            diagnostico["errores"].append(f"Error en anÃ¡lisis: {e}")
        
        logger.info(f"ðŸ” DiagnÃ³stico completado: {len(diagnostico['errores'])} errores encontrados")
        
        return diagnostico
        
    except Exception as e:
        logger.error(f"âŒ Error en diagnÃ³stico: {e}")
        return {
            "error": str(e),
            "timestamp": datetime.now().isoformat(),
            "ia_disponible_global": ANALIZADOR_IA_DISPONIBLE
        }


@app.get("/api/diagnostico/ia")
async def api_diagnostico_ia():
    """Endpoint para diagnÃ³stico detallado del modelo IA"""
    try:
        # Verificar modelos disponibles
        modelos = {}
        
        # Verificar modelo TF-IDF original
        tfidf_path = Path("models/modelo_legal.pkl")
        if tfidf_path.exists():
            try:
                with open(tfidf_path, 'rb') as f:
                    import pickle
                    data = pickle.load(f)
                    modelos["tfidf_original"] = {
                        "existe": True,
                        "tipo": "TF-IDF",
                        "componentes": list(data.keys()) if isinstance(data, dict) else ["modelo"]
                    }
            except Exception as e:
                modelos["tfidf_original"] = {
                    "existe": True,
                    "error": str(e)
                }
        else:
            modelos["tfidf_original"] = {"existe": False}
        
        # Verificar modelo SBERT
        sbert_path = Path("models/modelo_legal_sbert.pkl")
        if sbert_path.exists():
            try:
                with open(sbert_path, 'rb') as f:
                    import pickle
                    data = pickle.load(f)
                    modelos["sbert"] = {
                        "existe": True,
                        "encoder_name": data.get("encoder_name", "desconocido"),
                        "tipo": "SBERT" if data.get("encoder_name") != "tfidf" else "TF-IDF Fallback",
                        "componentes": list(data.keys()) if isinstance(data, dict) else ["modelo"]
                    }
            except Exception as e:
                modelos["sbert"] = {
                    "existe": True,
                    "error": str(e)
                }
        else:
            modelos["sbert"] = {"existe": False}
        
        # Verificar analizador
        try:
            from src.backend.analisis import AnalizadorLegal
            analizador = AnalizadorLegal()
            
            estado_analizador = {
                "sbert_encoder": analizador.sbert_encoder is not None,
                "sbert_clf": analizador.sbert_clf is not None,
                "vectorizador": analizador.vectorizador is not None,
                "clasificador": analizador.clasificador is not None,
                "modelo": analizador.modelo is not None
            }
            
            # Determinar quÃ© modelo se estÃ¡ usando
            if analizador.sbert_encoder is not None and analizador.sbert_clf is not None:
                if hasattr(analizador.sbert_encoder, 'encode'):
                    modelo_activo = "SBERT Real"
                else:
                    modelo_activo = "TF-IDF Fallback"
            elif analizador.vectorizador is not None and analizador.clasificador is not None:
                modelo_activo = "TF-IDF"
            else:
                modelo_activo = "AnÃ¡lisis por Reglas"
                
        except ImportError as e:
            estado_analizador = {"error": f"Error de importaciÃ³n: {str(e)}"}
            modelo_activo = "AnÃ¡lisis por Reglas"
        except Exception as e:
            estado_analizador = {"error": f"Error general: {str(e)}"}
            modelo_activo = "AnÃ¡lisis por Reglas"
        
        return {
            "modelos": modelos,
            "estado_analizador": estado_analizador,
            "modelo_activo": modelo_activo,
            "ia_disponible": ANALIZADOR_IA_DISPONIBLE,
            "timestamp": datetime.now().isoformat()
        }
        
    except Exception as e:
        return {
            "error": str(e),
            "timestamp": datetime.now().isoformat()
        }

@app.post("/api/demanda-base")
async def api_demanda_base(payload: Dict[str, Any]):
    try:
        nombres: List[str] = payload.get("nombres_archivo") or []
        if not isinstance(nombres, list) or not nombres:
            raise HTTPException(status_code=400, detail="Debe indicar 'nombres_archivo' (lista)")
        paths: List[Path] = []
        for n in nombres:
            p = SENTENCIAS_DIR / n
            if p.exists():
                paths.append(p)
        if not paths:
            raise HTTPException(status_code=404, detail="No se encontraron los archivos indicados")
        meta = payload.get("meta") if isinstance(payload.get("meta"), dict) else {}
        doc = _generar_demanda_base_para(paths, meta)
        return doc
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error generando demanda base: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/demanda-base/txt")
async def api_demanda_base_txt(payload: Dict[str, Any]):
    try:
        nombres: List[str] = payload.get("nombres_archivo") or []
        if not isinstance(nombres, list) or not nombres:
            raise HTTPException(status_code=400, detail="Debe indicar 'nombres_archivo' (lista)")
        # Buscar archivos en ambos directorios
        paths = []
        for nombre in nombres:
            candidatos = [SENTENCIAS_DIR / nombre, UPLOADS_DIR / nombre]
            for candidato in candidatos:
                if candidato.exists():
                    paths.append(candidato)
                    break
        if not paths:
            raise HTTPException(status_code=404, detail="No se encontraron los archivos indicados")
        meta = payload.get("meta") if isinstance(payload.get("meta"), dict) else {}
        doc = _generar_demanda_base_para(paths, meta)
        filename = f"demanda_base_{datetime.now().strftime('%Y%m%d_%H%M%S')}.txt"
        headers = {"Content-Disposition": f"attachment; filename={filename}"}
        return PlainTextResponse(content=doc.get("texto", ""), headers=headers)
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error generando TXT demanda base: {e}")
        raise HTTPException(status_code=500, detail=str(e))


@app.post("/api/demanda-base/docx")
async def api_demanda_base_docx(payload: Dict[str, Any]):
    """Genera una demanda en formato DOCX con formato profesional"""
    try:
        nombres: List[str] = payload.get("nombres_archivo") or []
        if not isinstance(nombres, list) or not nombres:
            raise HTTPException(status_code=400, detail="Debe indicar 'nombres_archivo' (lista)")
        # Buscar archivos en ambos directorios
        paths = []
        for nombre in nombres:
            candidatos = [SENTENCIAS_DIR / nombre, UPLOADS_DIR / nombre]
            for candidato in candidatos:
                if candidato.exists():
                    paths.append(candidato)
                    break
        if not paths:
            raise HTTPException(status_code=404, detail="No se encontraron los archivos indicados")
        meta = payload.get("meta") if isinstance(payload.get("meta"), dict) else {}
        
        # Generar documento DOCX
        docx_buffer = _generar_demanda_docx(paths, meta)
        
        # Crear nombre de archivo
        nombre_demandante = meta.get("nombre", "demandante").replace(" ", "_")
        filename = f"demanda_{nombre_demandante}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
        
        # Configurar headers para descarga
        headers = {
            "Content-Disposition": f"attachment; filename={filename}",
            "Content-Type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        }
        
        # Devolver el documento DOCX
        from fastapi.responses import Response
        return Response(
            content=docx_buffer.getvalue(),
            headers=headers,
            media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error generando DOCX demanda base: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# ====== LISTAR / ELIMINAR DOCUMENTOS ======
@app.get("/api/documentos")
async def listar_documentos():
    """Lista documentos disponibles en sentencias/ y uploads/."""
    docs = []
    for carpeta in [SENTENCIAS_DIR, UPLOADS_DIR]:
        if carpeta.exists():
            for f in carpeta.iterdir():
                if f.is_file() and f.suffix.lower() in ['.pdf', '.txt']:
                    docs.append({
                        "nombre": f.name,
                        "ruta": str(f),
                        "carpeta": 'sentencias' if carpeta == SENTENCIAS_DIR else 'uploads',
                        "tamaÃ±o": f.stat().st_size,
                        "modificado": datetime.fromtimestamp(f.stat().st_mtime).isoformat()
                    })
    return {"documentos": sorted(docs, key=lambda x: x["modificado"], reverse=True)}


@app.delete("/api/documentos")
async def eliminar_documento_api(payload: DeleteDocumentPayload):
    """Elimina un documento por nombre desde sentencias/ o uploads/."""
    nombre = payload.nombre_archivo
    if not nombre:
        raise HTTPException(status_code=400, detail="Nombre de archivo requerido")
    candidatos = [SENTENCIAS_DIR / nombre, UPLOADS_DIR / nombre]
    encontrado = None
    for p in candidatos:
        if p.exists() and p.is_file():
            encontrado = p
            break
    if not encontrado:
        raise HTTPException(status_code=404, detail="Documento no encontrado")
    try:
        encontrado.unlink()
        return {"status": "ok", "eliminado": nombre}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"No se pudo eliminar: {e}")


@app.get("/health")
async def health_check():
    """Endpoint de salud del sistema"""
    return {
        "status": "ok",
        "version": "2.0.0",
        "timestamp": datetime.now().isoformat(),
        "ia_disponible": ANALIZADOR_IA_DISPONIBLE,
        "directorios": {
            "sentencias": str(SENTENCIAS_DIR),
            "uploads": str(UPLOADS_DIR),
            "models": str(MODELS_DIR)
        }
    }


@app.get("/api/documento/{nombre_archivo}")
async def obtener_documento(nombre_archivo: str):
    """Obtiene detalles de un documento especÃ­fico"""
    try:
        # Decodificar el nombre del archivo
        nombre_decodificado = nombre_archivo
        
        # Buscar el archivo en ambos directorios
        candidatos = [SENTENCIAS_DIR / nombre_decodificado, UPLOADS_DIR / nombre_decodificado]
        archivo_path = None
        
        for candidato in candidatos:
            if candidato.exists():
                archivo_path = candidato
                break
        
        if not archivo_path:
            raise HTTPException(status_code=404, detail=f"Documento '{nombre_decodificado}' no encontrado")
        
        # Obtener informaciÃ³n del archivo
        stat = archivo_path.stat()
        tamaÃ±o = f"{stat.st_size / 1024:.1f} KB"
        fecha_modificacion = datetime.fromtimestamp(stat.st_mtime).strftime("%Y-%m-%d %H:%M:%S")
        
        # Intentar analizar el documento si no ha sido analizado
        try:
            if ANALIZADOR_IA_DISPONIBLE:
                try:
                    from src.backend.analisis import AnalizadorLegal
                except ImportError as e:
                    logger.error(f"Error importando AnalizadorLegal: {e}")
                    raise HTTPException(status_code=500, detail=f"Error interno del servidor: {str(e)}")
                analizador = AnalizadorLegal()
                resultado = analizador.analizar_documento(str(archivo_path))
            else:
                resultado = analizador_basico.analizar_documento(str(archivo_path), nombre_decodificado)
            
            if resultado.get("procesado"):
                return {
                    "nombre": nombre_decodificado,
                    "procesado": True,
                    "tamaÃ±o": tamaÃ±o,
                    "fecha_procesamiento": fecha_modificacion,
                    "frases_clave": resultado.get("frases_clave", {}),
                    "total_frases": sum(info["total"] for info in resultado.get("frases_clave", {}).values()),
                    "resumen": resultado.get("resumen", ""),
                    "tipo_documento": resultado.get("tipo_documento", "desconocido")
                }
            else:
                return {
                    "nombre": nombre_decodificado,
                    "procesado": False,
                    "tamaÃ±o": tamaÃ±o,
                    "fecha_procesamiento": fecha_modificacion,
                    "error": resultado.get("error", "Error desconocido en el procesamiento"),
                    "frases_clave": {},
                    "total_frases": 0
                }
                
        except Exception as e:
            logger.error(f"Error analizando documento {nombre_decodificado}: {e}")
            return {
                "nombre": nombre_decodificado,
                "procesado": False,
                "tamaÃ±o": tamaÃ±o,
                "fecha_procesamiento": fecha_modificacion,
                "error": f"Error en el anÃ¡lisis: {str(e)}",
                "frases_clave": {},
                "total_frases": 0
            }
            
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error obteniendo documento {nombre_archivo}: {e}")
        raise HTTPException(status_code=500, detail=f"Error interno: {str(e)}")


# Ruta de compatibilidad: redirige a la vista de archivo
@app.get("/documento/{nombre_archivo}")
async def redirigir_documento(nombre_archivo: str):
    return RedirectResponse(url=f"/archivo/{nombre_archivo}")



def analizar_sentencias_existentes() -> Dict[str, Any]:
    """Analiza las sentencias existentes en la carpeta con cachÃ©"""
    global CACHE_TIMESTAMP, ANALISIS_CACHE
    
    # Verificar si el cachÃ© es vÃ¡lido
    if (CACHE_TIMESTAMP and 
        (datetime.now() - CACHE_TIMESTAMP).total_seconds() < CACHE_DURATION and 
        ANALISIS_CACHE):
        logger.info("ðŸ“‹ Usando anÃ¡lisis desde cachÃ©")
        return ANALISIS_CACHE
    
    try:
        logger.info(f"ðŸ” Iniciando anÃ¡lisis de sentencias existentes en: {SENTENCIAS_DIR}")
        
        if not SENTENCIAS_DIR.exists():
            logger.warning(f"âŒ La carpeta '{SENTENCIAS_DIR}' no existe")
            resultado = {
                "error": f"La carpeta '{SENTENCIAS_DIR}' no existe",
                "archivos_analizados": 0,
                "total_apariciones": 0,
                "resultados_por_archivo": {},
                "ranking_global": {}
            }
            ANALISIS_CACHE = resultado
            CACHE_TIMESTAMP = datetime.now()
            return resultado
        
        # Buscar archivos de texto y PDF en ambos directorios
        archivos_soportados = []
        
        # Buscar en directorio sentencias
        if SENTENCIAS_DIR.exists():
            archivos_soportados.extend([f for f in SENTENCIAS_DIR.iterdir() 
                                      if f.suffix.lower() in ['.txt', '.pdf']])
        
        # Buscar en directorio uploads
        if UPLOADS_DIR.exists():
            archivos_soportados.extend([f for f in UPLOADS_DIR.iterdir() 
                                      if f.suffix.lower() in ['.txt', '.pdf']])
        
        logger.info(f"ðŸ“ Archivos encontrados: {[f.name for f in archivos_soportados]}")
        
        if not archivos_soportados:
            logger.warning(f"âŒ No se encontraron archivos .txt o .pdf en '{SENTENCIAS_DIR}' ni '{UPLOADS_DIR}'")
            return {
                "error": f"No se encontraron archivos .txt o .pdf en '{SENTENCIAS_DIR}' ni '{UPLOADS_DIR}'",
                "archivos_analizados": 0,
                "total_apariciones": 0,
                "resultados_por_archivo": {},
                "ranking_global": {}
            }
        
        resultados_por_archivo = {}
        ranking_global = {}
        total_apariciones = 0
        
        for archivo in archivos_soportados:
            try:
                logger.info(f"ðŸ” Analizando archivo: {archivo.name}")
                
                # Marcar tiempo de inicio para este archivo
                tiempo_inicio = datetime.now()
                
                # Usar el analizador de IA si estÃ¡ disponible, sino el bÃ¡sico
                if ANALIZADOR_IA_DISPONIBLE:
                    logger.info(f"ðŸ¤– Usando analizador de IA para: {archivo.name}")
                    try:
                        from src.backend.analisis import AnalizadorLegal
                    except ImportError as e:
                        logger.error(f"Error importando AnalizadorLegal: {e}")
                        raise HTTPException(status_code=500, detail=f"Error interno del servidor: {str(e)}")
                    analizador = AnalizadorLegal()
                    analizador._tiempo_inicio = tiempo_inicio
                    resultado = analizador.analizar_documento(str(archivo))
                else:
                    logger.info(f"ðŸ”§ Usando analizador bÃ¡sico para: {archivo.name}")
                    analizador_basico._tiempo_inicio = tiempo_inicio
                    resultado = analizador_basico.analizar_documento(str(archivo), archivo.name)
                
                logger.info(f"ðŸ“Š Resultado para {archivo.name}: procesado={resultado.get('procesado')}")
                
                if resultado.get("procesado"):
                    # Calcular total de frases clave
                    frases_clave = resultado.get("frases_clave", {})
                    total_frases = sum(datos.get("total", 0) for datos in frases_clave.values())
                    
                    # Agregar campos adicionales al resultado
                    resultado["total_frases"] = total_frases
                    resultado["timestamp"] = tiempo_inicio.strftime("%Y-%m-%d %H:%M:%S")
                    resultado["categorias_encontradas"] = len(frases_clave)
                    
                    resultados_por_archivo[archivo.name] = resultado
                    
                    logger.info(f"ðŸ”‘ Frases clave encontradas en {archivo.name}: {list(frases_clave.keys())}")
                    logger.info(f"ðŸ“Š Total frases en {archivo.name}: {total_frases}")
                    
                    for categoria, datos in frases_clave.items():
                        if categoria in ranking_global:
                            ranking_global[categoria]["total"] += datos["total"]
                            ranking_global[categoria]["ocurrencias"].extend(datos["ocurrencias"])
                            logger.info(f"ðŸ“ˆ Actualizando {categoria}: total={ranking_global[categoria]['total']}")
                        else:
                            ranking_global[categoria] = {
                                "total": datos["total"],
                                "ocurrencias": datos["ocurrencias"]
                            }
                            logger.info(f"ðŸ†• Nueva categorÃ­a {categoria}: total={datos['total']}")
                        total_apariciones += datos["total"]
                else:
                    logger.warning(f"âš ï¸ Archivo {archivo.name} no se pudo procesar")
                    resultados_por_archivo[archivo.name] = {"error": "No se pudo procesar"}
                    
            except Exception as e:
                logger.error(f"âŒ Error analizando {archivo}: {e}")
                resultados_por_archivo[archivo.name] = {"error": f"Error: {str(e)}"}
        
        # Ordenar ranking
        ranking_ordenado = dict(sorted(ranking_global.items(), key=lambda x: x[1]["total"], reverse=True))
        
        logger.info(f"ðŸ“Š RESUMEN FINAL:")
        logger.info(f"  - Archivos analizados: {len(archivos_soportados)}")
        logger.info(f"  - Total apariciones: {total_apariciones}")
        logger.info(f"  - CategorÃ­as encontradas: {list(ranking_ordenado.keys())}")
        logger.info(f"  - Ranking global: {ranking_ordenado}")
        
        resultado = {
            "archivos_analizados": len(archivos_soportados),
            "total_apariciones": total_apariciones,
            "resultados_por_archivo": resultados_por_archivo,
            "ranking_global": ranking_ordenado
        }
        
        # Guardar en cachÃ©
        ANALISIS_CACHE = resultado
        CACHE_TIMESTAMP = datetime.now()
        logger.info("ðŸ’¾ AnÃ¡lisis guardado en cachÃ©")
        
        return resultado
        
    except Exception as e:
        logger.error(f"âŒ Error analizando sentencias: {e}")
        return {
            "error": f"Error al analizar sentencias: {str(e)}",
            "archivos_analizados": 0,
            "total_apariciones": 0,
            "resultados_por_archivo": {},
            "ranking_global": {}
        }


# Middleware para asegurar Content-Type correcto para archivos estÃ¡ticos
@app.middleware("http")
async def fix_content_type_middleware(request: Request, call_next):
    """Middleware para corregir Content-Type de archivos estÃ¡ticos"""
    response = await call_next(request)
    
    # Solo aplicar a rutas de archivos estÃ¡ticos, NO a rutas de la aplicaciÃ³n
    if request.url.path.startswith('/sentencias/') and request.url.path.endswith('.pdf'):
        response.headers["Content-Type"] = "application/pdf"
        response.headers["X-Content-Type-Options"] = "nosniff"
        response.headers["Cache-Control"] = "no-cache, no-store, must-revalidate"
        response.headers["Pragma"] = "no-cache"
        response.headers["Expires"] = "0"
        # Agregar header adicional para evitar interpretaciÃ³n como JavaScript
        response.headers["Content-Disposition"] = "inline"
    
    return response

# Ruta especÃ­fica para interceptar DEMANDA.pdf directamente
@app.get("/DEMANDA.pdf")
async def servir_demanda_pdf():
    """Sirve especÃ­ficamente el archivo DEMANDA.pdf con Content-Type correcto"""
    try:
        archivo_path = SENTENCIAS_DIR / "DEMANDA.pdf"
        
        if not archivo_path.exists():
            raise HTTPException(status_code=404, detail="Archivo DEMANDA.pdf no encontrado")
        
        # Headers especÃ­ficos para evitar interpretaciÃ³n como JavaScript
        headers = {
            "Content-Type": "application/pdf",
            "X-Content-Type-Options": "nosniff",
            "Cache-Control": "no-cache, no-store, must-revalidate",
            "Pragma": "no-cache",
            "Expires": "0",
            "Content-Disposition": "inline; filename=DEMANDA.pdf"
        }
        
        return FileResponse(
            path=str(archivo_path),
            media_type="application/pdf",
            filename="DEMANDA.pdf",
            headers=headers
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error sirviendo DEMANDA.pdf: {e}")
        raise HTTPException(status_code=500, detail=f"Error interno: {str(e)}")

# Ruta especÃ­fica para servir archivos desde sentencias/ con Content-Type correcto
@app.get("/sentencias/{nombre_archivo}")
async def servir_archivo_sentencias(nombre_archivo: str):
    """Sirve archivos desde el directorio sentencias/ con Content-Type correcto"""
    try:
        archivo_path = SENTENCIAS_DIR / nombre_archivo
        
        if not archivo_path.exists():
            raise HTTPException(status_code=404, detail="Archivo no encontrado")
        
        # Determinar el tipo MIME basado en la extensiÃ³n
        if archivo_path.suffix.lower() == '.pdf':
            media_type = "application/pdf"
        elif archivo_path.suffix.lower() == '.txt':
            media_type = "text/plain; charset=utf-8"
        else:
            media_type = "application/octet-stream"
        
        # Configurar headers especÃ­ficos
        headers = {
            "Content-Type": media_type,
            "X-Content-Type-Options": "nosniff",
            "Cache-Control": "no-cache, no-store, must-revalidate",
            "Pragma": "no-cache",
            "Expires": "0"
        }
        
        return FileResponse(
            path=str(archivo_path),
            media_type=media_type,
            filename=nombre_archivo,
            headers=headers
        )
        
    except HTTPException:
        raise
    except Exception as e:
        logger.error(f"Error sirviendo archivo {nombre_archivo}: {e}")
        raise HTTPException(status_code=500, detail=f"Error interno: {str(e)}")

if __name__ == "__main__":
    import uvicorn
    
    def safe_print(text: str) -> None:
        try:
            print(text)
        except Exception:
            try:
                # Remove non-encodable characters for Windows consoles
                print(text.encode("cp1252", "ignore").decode("cp1252"))
            except Exception:
                try:
                    print(text.encode("utf-8", "ignore").decode("utf-8"))
                except Exception:
                    print("Inicio de servidor FastAPI")
    
    safe_print("ðŸš€ Iniciando Analizador de Sentencias IPP/INSS...")
    safe_print(f"ðŸ“ Directorio de sentencias: {SENTENCIAS_DIR}")
    safe_print(f"ðŸ¤– IA disponible: {'âœ… SÃ­' if ANALIZADOR_IA_DISPONIBLE else 'âŒ No'}")
    safe_print(f"ðŸŒ URL: http://localhost:8000")
    safe_print(f"ðŸ“š DocumentaciÃ³n: http://localhost:8000/docs")
    
    uvicorn.run(
        app, 
        host="0.0.0.0", 
        port=8000,
        log_level="info"
    )
